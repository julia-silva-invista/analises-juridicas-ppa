# -*- coding: utf-8 -*-
"""Geração do Dossiê PPA Invista em Word — modelo Parecer_Invista_PPA_v2_Atualizada 3.0."""

import io
import json
import os
import re
import shutil
import tempfile
from copy import deepcopy
from datetime import date as _date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import Table as _DocxTable
from docx.text.paragraph import Paragraph as _DocxParagraph

from google.genai import types

from utils import _retry

# ── Paleta Invista (extraída do template v3.0) ────────────────────────────
_LARANJA   = "E8440A"   # headers de tabela, títulos
_CINZA     = "F5F5F3"   # coluna-label (embargo/recurso), linhas TOTAL
_BRANCO    = "FFFFFF"   # células de valor
_DESTAQUE  = "FEF0EB"   # caixa de considerações
_BORDA     = "E8E8E8"   # bordas suaves
_TXT       = "555555"   # texto escuro padrão
_TXT_MUTE  = "AAAAAA"   # notas/placeholder

_LOGO_PATH = Path(__file__).parent / "assets" / "invista_logo.png"
_TEMPLATE_PATH = Path(__file__).parent / "assets" / "Parecer_Invista_PPA_v2_Atualizada_3.0.docx"

# Siglas que devem conservar a grafia técnica mesmo quando o restante do
# texto vier integralmente em caixa alta.
_SIGLAS_PRESERVADAS = {
    "AC", "AL", "AM", "AP", "AR", "BA", "BACENJUD", "CAC", "CC", "CCB",
    "CE", "CNAE", "CNJ", "CNPJ", "CPC", "CPF", "CPR", "CPRF", "CRI", "CTN", "DF", "ES",
    "GO", "ID", "IDPJ", "INFOJUD", "INPC", "IPCA", "MA", "MG", "MS", "MT",
    "OAB", "OJ", "PA", "PB", "PE", "PI", "PPA", "PR", "RENAJUD", "RG", "RJ",
    "RN", "RO", "RR", "RS", "SA", "SAT", "SC", "SE", "SERASAJUD",
    "SISBAJUD", "SNIPER", "SOP", "SP", "STF", "STJ", "TJSP", "TO", "TRF",
    "TST", "UF", "VM", "VP",
}

# Regra de citação compartilhada entre Checklist RJ, Checklist de Créditos e Dossiê PPA
# (checklist_rj.py importa esta constante — mantém as 3 fontes de extração JSON com o
# mesmo formato de referência, revisado para incluir ID/Mov./Evento/fls. JUNTO com a
# página, e indicar de qual PDF quando mais de um foi analisado).
REGRA_CITACAO_PADRAO = """\
═══ REGRA — REFERÊNCIA PROCESSUAL (identificador + página, sempre juntos) ═══
Toda referência deve trazer DOIS elementos juntos, separados por " | ":
  1. O identificador de movimentação do sistema do tribunal daquele processo — "ID <nº>"
     (Projudi/outro), "Mov. <nº>" (PJe), "Evento <nº>" (Eproc) ou "fls. <nº>" (eSAJ/físico).
  2. O número da folha onde a informação aparece, no formato "fl. <nº>".
Formato-padrão: "(ID 188753786 | fl. 135)".
Se MAIS DE 1 PDF foi analisado nesta análise, acrescente ao final de onde veio, preferindo
o número do processo daquele PDF quando identificável no próprio texto, ou o nome do arquivo
como alternativa: "(ID 188753786 | fl. 135 do pdf 0001259-66.1996.8.11.0041)". Com um único
PDF analisado, NUNCA acrescente "do pdf ..." — é informação redundante e polui a leitura.
NUNCA cite "texto bruto", "PARTE i/N" ou qualquer marcador técnico de divisão de arquivo como
se fosse referência processual — são só limites internos de divisão, sem significado jurídico.
Antes de responder, confira mentalmente se o identificador e a página citados são exatamente
os que aparecem no trecho de onde a informação foi extraída — referências erradas (ID ou
página incorretos) já causaram problemas antes; na dúvida, prefira omitir a referência a
citar uma incorreta.
Campos que NUNCA levam referência (só o valor puro, sem parênteses): número do processo,
vara, nome do credor, nome do exequente, nome do executado.
═══ REGRA — SIGLAS DE INSTRUMENTO EM MAIÚSCULA ═══
Sempre que citar um instrumento de crédito por sigla (CCB, CRI, CPR, CPRF, CCI, NCE, CDCA,
LCA ou similar), escreva a sigla inteira em MAIÚSCULA, mesmo que o restante da frase esteja
em title case/minúsculas.
═══ REGRA — STATUS PROCESSUAL LEGÍVEL (nunca "movimento" cru) ═══
Campos de status/situação processual NUNCA devem reproduzir o texto literal de um andamento
("movimento", "certidão", etc.) — escreva uma descrição curta e legível do estado real:
"Suspenso" (sem caixa alta) quando o processo/recurso estiver de fato suspenso; senão, descreva
a fase concreta, ex.: "Discussão de Embargos à Execução" (só se os embargos ainda não tiverem
transitado em julgado), "Em fase de busca de bens", "Leilão agendado". Se um recurso já
transitou em julgado, informe isso explicitamente no status.
"""

REGRA_COMPLETUDE_PADRAO = """\
═══ REGRA — COMPLETUDE ENTRE RELATÓRIO, DOSSIÊ E CHECKLISTS ═══
Quando houver RELATÓRIO CONSOLIDADO, ele é a fonte principal para todo fato que já esteja
descrito nele. Ao transportar uma informação do relatório para um campo do JSON:
- reproduza o conteúdo completo, sem resumir, encurtar ou parafrasear;
- preserve ressalvas, fatos posteriores, decisões, datas, índices, taxas e TODAS as referências
  relacionadas àquela informação;
- se o relatório registrar uma situação original e uma alteração posterior, mantenha ambas no
  mesmo campo, na mesma ordem lógica;
- use o texto bruto apenas para acrescentar fatos ausentes do relatório. Nunca substitua uma
  formulação completa do relatório por uma versão menor encontrada no texto bruto.
═══ FIM DA REGRA DE COMPLETUDE ═══
"""

_RE_MARCADOR_REFERENCIA = re.compile(
    r"(?:\bID\s+(?:n[.º°o]\s*)?\d"
    r"|\bMov(?:imento)?\.?\s+(?:n[.º°o]\s*)?\d"
    r"|\bEvento\s+(?:n[.º°o]\s*)?\d"
    r"|\bfls?\.?\s+(?:\d|s\.?\s*/\s*n\.?)"
    r"|\bp(?:\.|\u00e1g(?:ina)?\.?)\s*\d"
    r"|refer[êé]ncia\s+não\s+localizada)",
    re.IGNORECASE,
)
_RE_GRUPO_PARENTESES = re.compile(r"\(([^()\n]+)\)")
_RE_REFERENCIA_PURA = re.compile(
    r"^\s*(?:"
    r"(?:ID\s+(?:n[.º°o]\s*)?\d[\d./-]*)"
    r"|(?:Mov(?:imento)?\.?\s+(?:n[.º°o]\s*)?\d[\d./-]*)"
    r"|(?:Evento\s+(?:n[.º°o]\s*)?\d[\d./-]*)"
    r"|(?:fls?\.?\s+(?:\d[\d./-]*|s\.?\s*/\s*n\.?))"
    r"|(?:p(?:\.|\u00e1g(?:ina)?\.?)\s*\d[\d./-]*)"
    r")"
    r"(?:\s*(?:\||;|,)\s*(?:"
    r"ID\s+(?:n[.º°o]\s*)?\d[\d./-]*"
    r"|Mov(?:imento)?\.?\s+(?:n[.º°o]\s*)?\d[\d./-]*"
    r"|Evento\s+(?:n[.º°o]\s*)?\d[\d./-]*"
    r"|fls?\.?\s+(?:\d[\d./-]*|s\.?\s*/\s*n\.?)"
    r"|p(?:\.|\u00e1g(?:ina)?\.?)\s*\d[\d./-]*"
    r"))*"
    r"(?:\s+do\s+pdf\s+.+)?\s*$",
    re.IGNORECASE,
)


def _segmentar_referencias(texto: str):
    """Separa referências para que só o conteúdo entre parênteses fique em itálico."""
    valor = str(texto if texto is not None else "")
    if _RE_REFERENCIA_PURA.fullmatch(valor):
        return [(valor, True)]

    segmentos = []
    inicio = 0
    for match in _RE_GRUPO_PARENTESES.finditer(valor):
        interior = match.group(1)
        if not _RE_MARCADOR_REFERENCIA.search(interior):
            continue
        if match.start() > inicio:
            segmentos.append((valor[inicio:match.start()], False))
        segmentos.extend([
            ("(", False),
            (interior, True),
            (")", False),
        ])
        inicio = match.end()
    if inicio < len(valor):
        segmentos.append((valor[inicio:], False))
    return segmentos or [(valor, False)]
_PARTICULAS_NOME = {
    "a", "as", "à", "às", "ao", "aos", "com", "da", "das", "de", "do",
    "dos", "e", "em", "na", "nas", "no", "nos", "para", "pela", "pelas",
    "pelo", "pelos", "por", "sob", "sobre", "x",
}
_CAMPOS_NARRATIVOS = {
    "analise",
    "andamentos_resumo",
    "atividades_secundarias",
    "ato",
    "consideracoes_gerais",
    "credito_propositura",
    "criterio_sat",
    "decisao_recorrida",
    "descricao",
    "detalhamento",
    "evidencias",
    "fraude_execucao",
    "fundamentacao",
    "garantia",
    "ma_fe_insolvencia",
    "memoria_indices",
    "memoria_ponderacoes",
    "observacoes",
    "outras",
    "plan_ponderacoes",
    "resumo",
    "risco_juridico",
    "situacao_produtiva",
    "status",
    "status_processo",
    "tese",
    "teses_principais",
    "texto",
    "upside",
}
_RE_PALAVRA = re.compile(
    r"[A-Za-zÀ-ÖØ-öø-ÿ]+(?:[.'’][A-Za-zÀ-ÖØ-öø-ÿ]+)*\.?"
)
# Trecho de 2+ palavras em caixa alta seguidas (ex.: "MIRANDA LIMA E LOBO
# ADVOGADOS" dentro de uma frase mista) — usado quando a string INTEIRA não
# bate o limiar de 72% maiúsculas (uma célula que mistura nome em caixa alta
# com texto de referência em minúsculo, por exemplo), mas ainda assim tem um
# trecho isolado que merece a mesma normalização.
_RE_TRECHO_MAIUSCULO = re.compile(
    r"[A-ZÀ-Ö][A-ZÀ-Ö.'’]*\.?(?:\s+[A-ZÀ-Ö][A-ZÀ-Ö.'’]*\.?)+"
)


def _titlecase_palavra(match, indice_ref) -> str:
    """Normaliza uma palavra pra Title Case, preservando siglas e não
    capitalizando partículas de nome (de/da/dos/e/...) fora da 1ª posição."""
    palavra = match.group(0)
    chave = re.sub(r"[^A-Za-zÀ-ÖØ-öø-ÿ]", "", palavra).upper()
    minuscula = palavra.lower()
    if chave in _SIGLAS_PRESERVADAS:
        resultado = palavra.upper()
    elif indice_ref[0] > 0 and minuscula.rstrip(".") in _PARTICULAS_NOME:
        resultado = minuscula
    else:
        resultado = minuscula[:1].upper() + minuscula[1:]
    indice_ref[0] += 1
    return resultado


def _normalizar_trechos_maiusculos(texto: str) -> str:
    """Normaliza só os TRECHOS em caixa alta (2+ palavras seguidas) dentro de
    uma string mista — usada quando a string inteira não bate o limiar de
    maiúsculas de `_normalizar_caixa_alta`/`_normalizar_texto_narrativo`
    (ex.: célula que combina "MIRANDA LIMA E LOBO ADVOGADOS" com texto de
    referência em minúsculo — a string toda cai abaixo do limiar, mas o nome
    embutido continua merecendo a mesma normalização)."""
    def _sub_trecho(match):
        indice_ref = [0]
        return _RE_PALAVRA.sub(lambda m: _titlecase_palavra(m, indice_ref), match.group(0))

    resultado = _RE_TRECHO_MAIUSCULO.sub(_sub_trecho, texto)
    resultado = re.sub(r"\be-CAC\b", "e-CAC", resultado, flags=re.IGNORECASE)
    return re.sub(r"\br\s*\$\s*", "R$ ", resultado, flags=re.IGNORECASE)


def _normalizar_caixa_alta(texto) -> str:
    """
    Converte texto predominantemente em caixa alta para capitalização natural.

    Nomes como ``JULIA DE OLIVEIRA`` tornam-se ``Julia de Oliveira``. Siglas
    jurídicas e financeiras conhecidas permanecem em caixa alta. Textos que
    já possuem capitalização normal não são alterados.
    """
    valor = str(texto if texto is not None else "")
    if valor.strip().casefold() == "e-cac":
        return "e-CAC"
    letras = [c for c in valor if c.isalpha()]
    if len(letras) < 2:
        return valor
    proporcao_maiusculas = sum(c.isupper() for c in letras) / len(letras)
    if proporcao_maiusculas < 0.72:
        # A string inteira não é predominantemente maiúscula (ex.: mistura
        # nome em caixa alta com referência/conector em minúsculo) — ainda
        # assim normaliza qualquer trecho isolado que esteja em caixa alta.
        return _normalizar_trechos_maiusculos(valor)

    indice_ref = [0]
    resultado = _RE_PALAVRA.sub(lambda m: _titlecase_palavra(m, indice_ref), valor)
    resultado = re.sub(r"\be-CAC\b", "e-CAC", resultado, flags=re.IGNORECASE)
    return re.sub(r"\br\s*\$\s*", "R$ ", resultado, flags=re.IGNORECASE)


def _normalizar_texto_narrativo(texto) -> str:
    """Converte caixa alta em capitalização de frase, preservando siglas."""
    valor = str(texto if texto is not None else "")
    letras = [c for c in valor if c.isalpha()]
    if len(letras) < 2:
        return valor
    proporcao_maiusculas = sum(c.isupper() for c in letras) / len(letras)
    if proporcao_maiusculas < 0.72:
        # Mesmo fallback de _normalizar_caixa_alta: a frase toda não é
        # predominantemente maiúscula, mas um nome/entidade embutido pode
        # estar — normaliza só esse trecho (Title Case, não minúsculo de
        # frase, já que é um nome próprio, não o início de uma sentença).
        return _normalizar_trechos_maiusculos(valor)

    inicio_frase = True
    fim_anterior = 0

    def _ajustar(match):
        nonlocal inicio_frase, fim_anterior
        entre = valor[fim_anterior:match.start()]
        if re.search(r"(?:^|[.!?]\s+|\n+)$", entre):
            inicio_frase = True
        palavra = match.group(0)
        chave = re.sub(r"[^A-Za-zÀ-ÖØ-öø-ÿ]", "", palavra).upper()
        if chave in _SIGLAS_PRESERVADAS:
            resultado = palavra.upper()
        else:
            resultado = palavra.lower()
            if inicio_frase:
                resultado = resultado[:1].upper() + resultado[1:]
        inicio_frase = False
        fim_anterior = match.end()
        return resultado

    resultado = _RE_PALAVRA.sub(_ajustar, valor)
    resultado = re.sub(r"\be-CAC\b", "e-CAC", resultado, flags=re.IGNORECASE)
    return re.sub(r"\br\s*\$\s*", "R$ ", resultado, flags=re.IGNORECASE)


def _normalizar_dados(valor, chave_atual=""):
    """Aplica a capitalização natural recursivamente aos dados extraídos."""
    if isinstance(valor, dict):
        return {
            chave: _normalizar_dados(item, chave)
            for chave, item in valor.items()
        }
    if isinstance(valor, list):
        return [_normalizar_dados(item, chave_atual) for item in valor]
    if isinstance(valor, str):
        if chave_atual in _CAMPOS_NARRATIVOS:
            return _normalizar_texto_narrativo(valor)
        return _normalizar_caixa_alta(valor)
    return valor


# ══════════════════════════════════════════════════════════════════════════
# Helpers de célula (baixo nível)
# ══════════════════════════════════════════════════════════════════════════

def _set_bg(cell, hex6: str):
    tcp = cell._tc.get_or_add_tcPr()
    for old in tcp.findall(qn("w:shd")):
        tcp.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcp.append(shd)


def _set_borders(cell, color=_BORDA, sz=4):
    tcp = cell._tc.get_or_add_tcPr()
    for old in tcp.findall(qn("w:tcBorders")):
        tcp.remove(old)
    be = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        be.append(b)
    tcp.append(be)


def _cell_pad(cell, top=50, left=90, bottom=50, right=90):
    tcp = cell._tc.get_or_add_tcPr()
    for old in tcp.findall(qn("w:tcMar")):
        tcp.remove(old)
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tcp.append(mar)


def _apply_font(run, bold, size, color, italic):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    try:
        run.font.color.rgb = RGBColor.from_string(color)
    except Exception:
        pass


def _adicionar_runs_formatados(paragraph, texto, aplicar_formato):
    """Adiciona texto preservando quebras e isolando referências em runs itálicos."""
    linhas = str(texto if texto is not None else "").split("\n")
    for indice, linha in enumerate(linhas):
        normalizada = _normalizar_caixa_alta(linha)
        for trecho, eh_referencia in _segmentar_referencias(normalizada):
            if not trecho:
                continue
            run = paragraph.add_run(trecho)
            aplicar_formato(run, eh_referencia)
        if indice < len(linhas) - 1:
            run_quebra = paragraph.add_run()
            aplicar_formato(run_quebra, False)
            run_quebra.add_break()


def _write(cell, text, bold=False, size=9, color=_TXT, italic=False, align=None, valign=True):
    cell.text = ""
    if valign:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    _adicionar_runs_formatados(
        p,
        text,
        lambda run, eh_ref: _apply_font(run, bold, size, color, italic or eh_ref),
    )


# ══════════════════════════════════════════════════════════════════════════
# Helpers de linha / tabela (médio nível)
# ══════════════════════════════════════════════════════════════════════════

def _widths(row, ws):
    for i, w in enumerate(ws):
        row.cells[i].width = Cm(w)


def _lock_widths(table, ws):
    """Faz a largura de coluna pedida valer de verdade.

    python-docx só respeita w:tcW por célula se a tabela não estiver em
    autofit — sem isso, o Word redistribui as colunas quase igualmente na
    hora de renderizar, ignorando a proporção pedida em `ws`."""
    table.autofit = False
    for i, w in enumerate(ws):
        table.columns[i].width = Cm(w)


def _orange_header(table, labels, ws=None):
    """Linha de header laranja com texto branco em negrito."""
    row = table.add_row()
    for i, lab in enumerate(labels):
        c = row.cells[i]
        _set_bg(c, _LARANJA); _set_borders(c); _cell_pad(c)
        _write(c, lab, bold=True, size=9, color=_BRANCO)
    if ws:
        _widths(row, ws)
    return row


def _span_header(table, text):
    """Header laranja que abrange todas as colunas (ex.: DADOS DA MEMÓRIA DE CÁLCULO)."""
    row = table.add_row()
    n = len(row.cells)
    cell = row.cells[0]
    for i in range(1, n):
        cell = cell.merge(row.cells[i])
    _set_bg(cell, _LARANJA); _set_borders(cell); _cell_pad(cell)
    _write(cell, text, bold=True, size=9, color=_BRANCO)
    return row


def _kv_table(doc, header, rows, w_label=5.9, w_value=11.0):
    """Tabela 2 col estilo 'header laranja' (CAMPO|INFORMAÇÃO) + labels/valores brancos."""
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _lock_widths(t, [w_label, w_value])
    if header:
        if isinstance(header, str):
            _span_header(t, header)
        else:
            _orange_header(t, header, [w_label, w_value])
    for label, value in rows:
        r = t.add_row()
        cl, cv = r.cells[0], r.cells[1]
        _set_bg(cl, _BRANCO); _set_borders(cl); _cell_pad(cl)
        _write(cl, label, bold=True, color=_TXT)
        _set_bg(cv, _BRANCO); _set_borders(cv); _cell_pad(cv)
        _write(cv, value, color=_TXT, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _widths(r, [w_label, w_value])
        _nao_dividir_linha(r)
    return t


def _kv_label_table(doc, rows, w_label=5.9, w_value=11.0):
    """Tabela 2 col com coluna-label cinza (embargo/recurso), sem header laranja."""
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _lock_widths(t, [w_label, w_value])
    for label, value in rows:
        r = t.add_row()
        cl, cv = r.cells[0], r.cells[1]
        _set_bg(cl, _CINZA); _set_borders(cl); _cell_pad(cl)
        _write(cl, label, bold=True, color=_TXT)
        _set_bg(cv, _BRANCO); _set_borders(cv); _cell_pad(cv)
        _write(cv, value, color=_TXT, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _widths(r, [w_label, w_value])
        _nao_dividir_linha(r)
    return t


def _grid_table(doc, headers, rows, ws, total_row=None, min_rows=1):
    """Tabela multi-coluna: header laranja + linhas brancas (+ TOTAL cinza opcional)."""
    t = doc.add_table(rows=0, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _lock_widths(t, ws)
    _orange_header(t, headers, ws)
    data = list(rows)
    while len(data) < min_rows:
        data.append([""] * len(headers))
    for row_data in data:
        r = t.add_row()
        for i in range(len(headers)):
            c = r.cells[i]
            val = row_data[i] if i < len(row_data) else ""
            _set_bg(c, _BRANCO); _set_borders(c); _cell_pad(c)
            _write(c, val, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _widths(r, ws)
    if total_row:
        r = t.add_row()
        for i in range(len(headers)):
            c = r.cells[i]
            val = total_row[i] if i < len(total_row) else ""
            _set_bg(c, _CINZA); _set_borders(c); _cell_pad(c)
            _write(c, val, bold=True)
        _widths(r, ws)
    return t


# ══════════════════════════════════════════════════════════════════════════
# Helpers de parágrafo (títulos, notas, texto)
# ══════════════════════════════════════════════════════════════════════════

def _para(doc, text, bold=False, size=9.5, color=_TXT, italic=False,
          before=0, after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    _adicionar_runs_formatados(
        p,
        text,
        lambda run, eh_ref: _apply_font(run, bold, size, color, italic or eh_ref),
    )
    return p


def _sec_title(doc, text):
    return _para(doc, text, bold=True, size=13, color=_LARANJA, before=14, after=5)


def _sub_orange(doc, text):
    return _para(doc, text, bold=True, size=11, color=_LARANJA, before=10, after=4)


def _sub_gray(doc, text):
    return _para(doc, text, bold=True, size=10, color=_TXT, before=7, after=3)


def _note(doc, text):
    return _para(
        doc, text, size=8, color=_TXT_MUTE, italic=True, before=2, after=6,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _body(doc, text):
    return _para(
        doc, text, size=9.5, color=_TXT, before=0, after=5,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _guidance(doc, text):
    """Caixa de orientação com fundo FEF0EB (célula única)."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = t.rows[0].cells[0]
    _set_bg(c, _DESTAQUE); _set_borders(c, color="F3D9CF"); _cell_pad(c, top=90, bottom=90)
    _write(
        c, text, size=9, color=_TXT, italic=True,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    c.width = Cm(16.9)
    return t


def _spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)


# ══════════════════════════════════════════════════════════════════════════
# Extração de dados via Gemini
# ══════════════════════════════════════════════════════════════════════════

_PROMPT_DOSSIE = REGRA_CITACAO_PADRAO + REGRA_COMPLETUDE_PADRAO + """
Com base no material jurídico abaixo (extração completa do(s) processo(s) e, quando houver, relatório
consolidado de apoio), extraia os dados para preencher um dossiê no formato JSON.
Extraia APENAS o que está explicitamente no material. Use string vazia "" para o que não constar.
NÃO invente números, nomes ou datas. Responda SOMENTE com o JSON, sem texto adicional.
As regras de referência/siglas/status acima (antes deste parágrafo) valem para TODOS os campos abaixo.

═══ REGRA 2 — CAMPOS QUE DEVEM SER PREENCHIDOS ═══
- Índices de Correção do Contrato (ind_*): fonte é SEMPRE o LASTRO/TÍTULO EXECUTIVO em si — o
  contrato, CCB, CPR, CPRF, duplicata etc. que disciplina a dívida. Esses instrumentos costumam
  estar digitalizados/escaneados no INÍCIO do processo (aplique OCR visual se necessário) — NÃO
  extraia esses campos de petições, decisões ou outras peças processuais, mesmo que elas
  mencionem os índices de passagem; a fonte é sempre a cláusula contratual que os disciplina.
  Cada campo (ind_cm/ind_jr/ind_jm/ind_multa/ind_cap) é o ÍNDICE OU A TAXA que o contrato prevê
  — nunca um valor em R$. Exemplos corretos: "IPCA" (ind_cm), "1% ao mês" (ind_jr), "1% ao mês"
  (ind_jm), "2% sobre o débito" (ind_multa), "capitalização mensal/anual" (ind_cap). Se o contrato
  não disciplinar algum desses itens, deixe "" — nunca infira um valor calculado como se fosse o
  índice.
- Planilha Inicial (plan_*): fonte é a PRIMEIRA memória de cálculo/planilha que instruiu a petição
  inicial (a que embasou o valor da causa). Mesma regra de conteúdo dos campos ind_*: cada campo
  (plan_cm/plan_jr/plan_multa/plan_cap) deve trazer o ÍNDICE OU A TAXA que essa planilha efetivamente
  aplicou no cálculo (ex.: "INPC", "1,5% ao mês", "10% sobre o débito") — NUNCA o valor em R$
  resultante do cálculo (isso NÃO é o que se quer aqui). "plan_ponderacoes" é o único campo de
  texto livre deste bloco (ex.: alguma ressalva ou observação sobre a planilha).
- Última Memória de Cálculo (memoria_*): a memória de débito GENUINAMENTE mais recente JUNTADA
  aos autos — confira a data de juntada de cada memória de cálculo encontrada no processo e use
  a de data MAIS RECENTE, não a que aparecer por último na leitura ou a mais citada. Aqui sim
  "memoria_total" é o valor em R$ total atualizado (é o objetivo deste bloco: qual o débito
  atualizado segundo o cálculo mais recente) e "memoria_indices" descreve os índices/taxas que
  essa memória aplicou.
- Liste TODOS os embargos à execução, TODAS as exceções/objeções de pré-executividade e TODOS os
  recursos que aparecerem (não apenas o primeiro). No campo "tipo" de cada defesa informe
  "Embargos à Execução" ou "Exceção de Pré-Executividade".

═══ REGRA 3 — CITAÇÃO (uma entrada por executado) ═══
Gere UMA entrada por executado da execução. Se o executado FOI citado, informe modalidade (AR/OJ/
Edital/Precatória), data e fls. da citação. Se NÃO foi citado, ponha modalidade="Pendente" e liste no
campo "data" as datas das tentativas e no campo "fls" as páginas correspondentes (na mesma linha).

═══ REGRA 4 — PENHORAS E ANDAMENTOS ═══
- Constrições: no campo "status", se a penhora/constrição foi DEFERIDA, inclua a DATA da decisão de
  deferimento. Ex.: "Penhora deferida em 15/03/2024 (fls. 120)", "Sisbajud deferido em 02/2024 (Mov. 88)".
- andamentos: liste TODOS os andamentos processuais que constarem no material E no relatório de apoio —
  NÃO resuma, NÃO omita e NÃO limite a quantidade. Cada andamento com data, descrição objetiva e
  referência (Mov./fls./Evento/ID).

═══ REGRA 5 — FORMATAÇÃO DO TEXTO ═══
- Nunca escreva nomes de pessoas físicas ou jurídicas integralmente em caixa alta.
- Use capitalização natural: "Julia de Oliveira Bernardo da Silva", "Empresa Agrícola Ltda.".
- Preserve em caixa alta somente siglas técnicas, como CPF, CNPJ, CNJ, OAB, SAT, SOP, IDPJ, VM e VP.
- Redija títulos, descrições, teses, status e andamentos com maiúsculas e minúsculas normais.

═══ REGRA 6 — QUADROS DO TEMPLATE OFICIAL ═══
- Preencha a visão consolidada dos ativos e as teses de recuperação sempre que a informação constar
  no material. Esses dados serão inseridos nos quadros já existentes do template; não invente outro layout.
- Em cada item narrativo, diferencie fato comprovado, indício, fundamento jurídico, risco e conclusão.
- Para ativos, informe matrícula, proprietário, ônus, fração atingível, VM, VP e saldo quando disponíveis.
- Não repita o mesmo fato em vários campos. Preserve a referência processual de cada afirmação.

{
  "nome_caso": "identificador curto (ex: BASF x São Lourenço)",
  "data_analise": "DD/MM/AAAA ou vazio",
  "advogada_responsavel": "nome com capitalização natural ou vazio",
  "exequentes": "Nome · CNPJ ... (fls.)",
  "executados": "Nome1 · CPF/CNPJ ...; Nome2 ... (fls.)",
  "sat_total": "R$ ... (fls.)",
  "total_atingivel_vm": "R$ ... ou vazio",
  "total_atingivel_vp": "R$ ... ou vazio",
  "teses_principais": "síntese objetiva das teses principais (com refs)",
  "passivo_fiscal": "", "passivo_trabalhista": "", "passivo_civel": "",
  "passivo_total": "",
  "risco_juridico": "resumo dos principais riscos (com refs)",
  "consideracoes_gerais": "2-4 parágrafos de análise geral (separe parágrafos com \\n)",
  "visao_consolidada_ativos": [
    {"tese": "Penhora Direta | IDPJ | Fraude à Execução | outra", "vm": "R$ ...", "vp": "R$ ...", "onus": "R$ ...", "observacoes": ""}
  ],
  "ativos": [
    {
      "tese": "Penhora Direta | IDPJ | Fraude à Execução | outra",
      "matricula": "Matrícula nº ...", "comarca": "Cidade/UF", "proprietario_atual": "",
      "tipo_ativo": "apartamento, galpão, terreno etc.", "descricao": "", "area": "",
      "situacao_produtiva": "", "liquidez": "", "fracao_atingivel": "",
      "onus_vigentes": "", "vm": "R$ ...", "vp": "R$ ...", "onus_total": "R$ ...",
      "saldo": "R$ ...", "observacoes": ""
    }
  ],
  "teses_recuperacao": {
    "penhora_direta": {
      "analise": [{"titulo": "", "texto": "fato, fundamento e conclusão", "referencia": "fls./Mov./ID/Evento"}]
    },
    "idpj": {
      "resumo": [{"titulo": "", "texto": "", "referencia": ""}],
      "empresa_alvo": {
        "razao_social": "", "cnpj": "", "cnae_principal": "", "atividades_secundarias": "",
        "socio_atual": "", "endereco_fiscal": "", "fundamentacao": "", "credito_propositura": ""
      },
      "cronologia": [{"data": "", "ato": "", "detalhamento": "", "referencia": ""}],
      "evidencias": [{"titulo": "", "texto": "", "referencia": ""}],
      "upside": [{"titulo": "", "texto": "", "referencia": ""}]
    },
    "fraude_execucao": {
      "resumo": [{"titulo": "", "texto": "", "referencia": ""}],
      "ma_fe_insolvencia": [{"titulo": "", "texto": "", "referencia": ""}]
    },
    "outras": [{"titulo": "", "texto": "", "referencia": ""}]
  },
  "creditos": [
    {
      "id": "Crédito [Credor]",
      "numero_processo": "0000000-00.0000.0.00.0000",
      "vara_comarca": "1ª Vara Cível — Cidade/UF",
      "exequente_info": "Nome · Adv: Escritório | OAB",
      "executados_info": "Nome · Adv: Escritório | OAB",
      "data_distribuicao": "DD/MM/AAAA (fls.)",
      "sop": "R$ ... (fls.)", "sat": "R$ ... (fls.)", "criterio_sat": "ex: INPC + 12% a.a JM (fls.)",
      "honorarios": "ex: 10% (fls.)",
      "lastro": "CCB nº / Contrato nº (fls.)", "data_emissao": "DD/MM/AAAA (fls.)", "data_vencimento": "DD/MM/AAAA (fls.)",
      "assinaturas": "Nomes (fls.)", "garantia": "descrição da garantia (fls.)",
      "status_processo": "ex: em fase de penhora (Mov.)",
      "ind_cm": "índice do CONTRATO/lastro, ex: IPCA (fls.)", "ind_jr": "taxa do CONTRATO, ex: 1% ao mês (fls.)", "ind_jm": "taxa do CONTRATO, ex: 1% ao mês (fls.)", "ind_multa": "percentual do CONTRATO, ex: 2% (fls.)", "ind_cap": "ex: mensal/anual, conforme o CONTRATO (fls.)",
      "plan_cm": "índice aplicado na PLANILHA INICIAL, ex: INPC (fls.) — nunca o valor em R$", "plan_jr": "taxa aplicada na planilha, ex: 1,5% ao mês (fls.)", "plan_multa": "percentual aplicado na planilha, ex: 10% (fls.)", "plan_cap": "ex: mensal/anual (fls.)", "plan_ponderacoes": "(fls.)",
      "memoria_data_juntada": "DD/MM/AAAA (fls.)", "memoria_total": "R$ ... (fls.)", "memoria_data_base": "DD/MM/AAAA",
      "memoria_indices": "índices aplicados (fls.)", "memoria_ponderacoes": "",
      "citacoes": [{"executado": "Nome", "modalidade": "AR/OJ/Edital ou Pendente", "data": "data citação OU datas das tentativas", "fls": "fls. citação OU fls. das tentativas"}],
      "embargos": [{"tipo": "Embargos à Execução | Exceção de Pré-Executividade", "embargante": "", "data_dist": "(fls.)", "tese": "", "andamentos_resumo": "(com refs)", "status": ""}],
      "recursos": [{"recorrente": "", "decisao_recorrida": "", "data_dist": "(fls.)", "tese": "", "andamentos_resumo": "(com refs)", "status": ""}],
      "constricoes": [{"tipo": "", "descricao": "(fls.)", "valor": "R$ ... (fls.)", "status": "Ativa | Deferida em DD/MM/AAAA (fls.)"}],
      "andamentos": [{"data": "DD/MM/AAAA", "descricao": "descrição objetiva", "fls": "Mov./fls./Evento/ID"}]
    }
  ]
}

MATERIAL:
"""


def _montar_material_prioritario(
    relatorio: str,
    texto_bruto: str,
    limite_relatorio: int = 200_000,
    limite_bruto: int = 700_000,
) -> str:
    """Relatório primeiro; extração bruta apenas complementa fatos ausentes."""
    rel = (relatorio or "").strip()[:limite_relatorio]
    bruto = (texto_bruto or "").strip()[:limite_bruto]
    partes = []
    if rel:
        partes.append(
            "═══ RELATÓRIO CONSOLIDADO (FONTE PRINCIPAL; PRESERVAR TEXTO COMPLETO) ═══\n"
            + rel
        )
    if bruto:
        partes.append(
            "═══ TEXTO BRUTO EXTRAÍDO (FONTE COMPLEMENTAR) ═══\n" + bruto
        )
    return "\n\n".join(partes)


def _extrair_dados(texto_completo: str, relatorio: str, client, model: str) -> dict:
    prompt = _PROMPT_DOSSIE + _montar_material_prioritario(relatorio, texto_completo)
    try:
        config = types.GenerateContentConfig(response_mime_type="application/json")

        def _fn():
            return client.models.generate_content(
                model=model,
                contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
                config=config,
            ).text

        raw = _retry(_fn, tentativas=3, espera_base=10)
        raw = re.sub(r"^```[a-z]*\n?", "", raw.strip(), flags=re.IGNORECASE)
        raw = re.sub(r"\n?```$", "", raw.strip())
        return json.loads(raw)
    except Exception:
        return {}


_PROMPT_COMPLETAR = REGRA_COMPLETUDE_PADRAO + """\
Abaixo está um dossiê em JSON (parcialmente preenchido) e o RELATÓRIO consolidado do caso.
Compare TODOS os campos com o relatório e devolva a versão mais completa de cada informação.
Regras:
- Se um campo estiver vazio, preencha-o com o que constar no relatório.
- Se um campo já tiver valor, mas o relatório trouxer uma versão MAIS COMPLETA, substitua a versão
  curta pela formulação completa do relatório, sem resumir ou parafrasear.
- Preserve todas as ressalvas, alterações posteriores e referências. NÃO apague informação existente.
- NÃO invente — use somente o que constar no relatório ou no JSON atual.
- Mantenha EXATAMENTE a mesma estrutura, chaves e lista de créditos do JSON atual.
- Inclua a referência da fonte (fls./Mov./ID/Evento) quando o relatório trouxer.
- Se um andamento/embargo/recurso constar no relatório e faltar no JSON, ACRESCENTE à lista.
Responda SOMENTE com o JSON completo.

JSON ATUAL:
{json_atual}

RELATÓRIO:
{relatorio}
"""


def _contar_referencias_texto(valor: str) -> int:
    return len(_RE_MARCADOR_REFERENCIA.findall(str(valor or "")))


def _mesclar_mais_completo(atual, candidato):
    """Aceita ampliações da segunda passada sem permitir que ela encurte dados."""
    if isinstance(atual, dict):
        novo = candidato if isinstance(candidato, dict) else {}
        return {
            chave: _mesclar_mais_completo(valor, novo.get(chave))
            for chave, valor in atual.items()
        } | {
            chave: valor for chave, valor in novo.items() if chave not in atual
        }
    if isinstance(atual, list):
        novo = candidato if isinstance(candidato, list) else []
        return [
            _mesclar_mais_completo(
                atual[indice] if indice < len(atual) else None,
                novo[indice] if indice < len(novo) else None,
            )
            for indice in range(max(len(atual), len(novo)))
        ]
    if isinstance(atual, str):
        base = atual.strip()
        opcao = candidato.strip() if isinstance(candidato, str) else ""
        if not opcao:
            return atual
        if not base:
            return candidato
        base_cf, opcao_cf = base.casefold(), opcao.casefold()
        if base_cf == opcao_cf or opcao_cf in base_cf:
            return atual
        if base_cf in opcao_cf:
            return candidato
        refs_base = _contar_referencias_texto(base)
        refs_opcao = _contar_referencias_texto(opcao)
        if refs_opcao > refs_base or (refs_opcao == refs_base and len(opcao) > len(base)):
            return candidato
        return atual
    if atual in (None, "") and candidato not in (None, ""):
        return candidato
    return atual


def _completar_com_relatorio(dados: dict, relatorio: str, client, model: str) -> dict:
    """Passada final: amplia campos curtos com a versão completa do relatório."""
    rel = (relatorio or "").strip()
    if not dados or not dados.get("creditos") or not rel:
        return dados
    try:
        prompt = _PROMPT_COMPLETAR.format(
            json_atual=json.dumps(dados, ensure_ascii=False)[:150_000],
            relatorio=rel[:150_000],
        )
        config = types.GenerateContentConfig(response_mime_type="application/json")

        def _fn():
            return client.models.generate_content(
                model=model,
                contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
                config=config,
            ).text

        raw = _retry(_fn, tentativas=2, espera_base=10)
        raw = re.sub(r"^```[a-z]*\n?", "", raw.strip(), flags=re.IGNORECASE)
        raw = re.sub(r"\n?```$", "", raw.strip())
        completo = json.loads(raw)
        # A mesclagem recursiva impede que a segunda passada apague ou encurte dados.
        if isinstance(completo, dict) and len(completo.get("creditos") or []) >= len(dados.get("creditos") or []):
            return _mesclar_mais_completo(dados, completo)
        return dados
    except Exception:
        return dados


# ══════════════════════════════════════════════════════════════════════════
# Construção do documento
# ══════════════════════════════════════════════════════════════════════════

def _montar_cabecalho_rodape(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.4)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    # Logo no cabeçalho de página (repete)
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    if _LOGO_PATH.exists():
        try:
            hp.add_run().add_picture(str(_LOGO_PATH), width=Cm(3.2))
        except Exception:
            pass

    # Rodapé: confidencialidade + página
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Confidencial  ·  Uso interno  ·  Time PPA  ·  Invista")
    _apply_font(r, bold=False, size=8, color=_TXT_MUTE, italic=False)


def _build_doc_programatico_obsoleto(dados: dict) -> str:
    dados = _normalizar_dados(dados or {})
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9.5)

    _montar_cabecalho_rodape(doc)

    hoje = _date.today().strftime("%d/%m/%Y")

    # ── Título ────────────────────────────────────────────────────────────
    _para(doc, "ANÁLISE DE CRÉDITO", bold=True, size=21, color=_LARANJA, before=2, after=0)
    _para(doc, "PPA Invista", bold=False, size=11, color=_TXT_MUTE, before=0, after=8)

    # ── Info box ─────────────────────────────────────────────────────────
    _kv_table(doc, ["CAMPO", "INFORMAÇÃO"], [
        ("Nome do Caso",           dados.get("nome_caso", "")),
        ("Data da Análise",        dados.get("data_analise") or hoje),
        ("Advogada Responsável",   ""),
    ])
    _spacer(doc)

    # ══ 0. VISÃO GERAL DO CASO ═══════════════════════════════════════════
    _sec_title(doc, "0. VISÃO GERAL DO CASO")

    pf = dados.get("passivo_fiscal", "")
    pt = dados.get("passivo_trabalhista", "")
    pc = dados.get("passivo_civel", "")
    passivo_cell = (
        f"Fiscal: {pf or 'R$ [___]'}\n"
        f"Trabalhista: {pt or 'R$ [___]'}\n"
        f"Cível: {pc or 'R$ [___]'}\n"
        f"TOTAL: R$ [___]"
    )
    _kv_table(doc, ["INDICADOR", "VALOR / INFORMAÇÃO"], [
        ("Exequente(s)",                 dados.get("exequentes", "")),
        ("Executado(s)",                 dados.get("executados", "")),
        ("SAT",                          dados.get("sat_total", "R$")),
        ("Total atingível mapeado — VM", "R$ [___]"),
        ("Total atingível mapeado — VP", "R$ [___]"),
        ("Passivo total identificado",   passivo_cell),
        ("Tese(s) principal(is)",        ""),
        ("Risco Jurídico",               dados.get("risco_juridico", "")),
    ])
    _spacer(doc)

    # Visão consolidada dos ativos
    _sub_orange(doc, "VISÃO CONSOLIDADA DOS ATIVOS")
    _note(doc, "Quadro-resumo dos ativos atingíveis por tese. Os valores já descontam os ônus incidentes sobre os ativos.")
    _grid_table(
        doc,
        ["TESE", "VM (R$)", "VP (R$)", "ÔNUS (R$)", "OBSERVAÇÕES"],
        [["Penhora Direta", "", "", "", ""],
         ["IDPJ", "", "", "", ""],
         ["Fraude à Execução", "", "", "", ""],
         ["Pauliana", "", "", "", ""]],
        [3.5, 2.6, 2.6, 2.6, 5.6],
        total_row=["TOTAL GERAL", "", "", "", ""],
    )
    _spacer(doc)

    # Visão geral dos atingíveis (3 tabelas)
    _sub_orange(doc, "VISÃO GERAL DOS ATINGÍVEIS")
    cols_at = ["ATIVO / MATRÍCULA", "TIPO DE ATIVO", "ÁREA / HA", "SIT. PRODUTIVA",
               "LIQUIDEZ", "% ATING.", "ÔNUS (R$)", "OBSERVAÇÕES"]
    ws_at = [2.6, 2.1, 1.6, 2.2, 1.9, 1.5, 2.0, 3.0]
    for tese in ["Penhora Direta", "IDPJ", "Fraude à Execução"]:
        _sub_gray(doc, tese)
        _grid_table(doc, cols_at, [], ws_at, min_rows=3)
        _spacer(doc, pts=2)

    # Principais considerações
    _sub_orange(doc, "PRINCIPAIS CONSIDERAÇÕES")
    cons = (dados.get("consideracoes_gerais") or "").strip()
    if cons:
        for para in cons.split("\n"):
            if para.strip():
                _body(doc, para.strip())
    else:
        _guidance(doc, "Registre aqui: visão geral da operação · principais riscos jurídicos e "
                       "patrimoniais · oportunidades identificadas · pontos críticos para aquisição · "
                       "impressões iniciais e contexto relevante do caso.")

    # ══ 1. VISÃO JURÍDICA ════════════════════════════════════════════════
    _sec_title(doc, "1. VISÃO JURÍDICA")

    creditos = dados.get("creditos") or []
    if not creditos:
        _body(doc, "Nenhum crédito identificado no relatório.")
    for idx, cred in enumerate(creditos, 1):
        cid = cred.get("id") or f"Crédito {idx}"
        _sub_orange(doc, f"1.{idx} {cid}")

        # Dados do Processo
        _sub_gray(doc, "Dados do Processo")
        _kv_table(doc, ["CAMPO", "INFORMAÇÃO"], [
            ("Número do processo",            cred.get("numero_processo", "")),
            ("Vara / Comarca",                cred.get("vara_comarca", "")),
            ("Exequente",                     cred.get("exequente_info", "")),
            ("Executado(s)",                  cred.get("executados_info", "")),
            ("Data de distribuição",          cred.get("data_distribuicao", "")),
            ("SOP",                           cred.get("sop", "R$")),
            ("SAT",                           cred.get("sat", "R$")),
            ("Critério de atualização do SAT", cred.get("criterio_sat", "")),
            ("Honorários",                    cred.get("honorarios", "")),
            ("Lastro / Instrumento",          cred.get("lastro", "")),
            ("Data de Emissão",               cred.get("data_emissao", "")),
            ("Data do Vencimento",            cred.get("data_vencimento", "")),
            ("Assinaturas",                   cred.get("assinaturas", "")),
            ("Garantia",                      cred.get("garantia", "")),
            ("Status do Processo",            cred.get("status_processo", "")),
        ])
        _spacer(doc, pts=2)

        # Índices de Correção do Contrato
        _sub_gray(doc, "Índices de Correção do Contrato")
        _kv_table(doc, ["CRITÉRIO", "PARÂMETRO"], [
            ("Correção monetária",   cred.get("ind_cm", "")),
            ("Juros remuneratórios", cred.get("ind_jr", "")),
            ("Juros moratórios",     cred.get("ind_jm", "")),
            ("Multa moratória",      cred.get("ind_multa", "")),
            ("Capitalização",        cred.get("ind_cap", "")),
        ])
        _spacer(doc, pts=2)

        # Planilha Inicial
        _sub_gray(doc, "Planilha Inicial")
        _kv_table(doc, ["CRITÉRIO", "PARÂMETRO"], [
            ("Correção monetária",   cred.get("plan_cm", "")),
            ("Juros remuneratórios", cred.get("plan_jr", "")),
            ("Multa moratória",      cred.get("plan_multa", "")),
            ("Capitalização",        cred.get("plan_cap", "")),
            ("Ponderações",          cred.get("plan_ponderacoes", "")),
        ])
        _spacer(doc, pts=2)

        # Última Memória de Cálculo
        _sub_gray(doc, "Última Memória de Cálculo")
        _kv_table(doc, "DADOS DA MEMÓRIA DE CÁLCULO", [
            ("Data da Juntada",   cred.get("memoria_data_juntada", "")),
            ("Total Atualizado",  cred.get("memoria_total", "")),
            ("Data-base",         cred.get("memoria_data_base", "")),
            ("Índices aplicados", cred.get("memoria_indices", "")),
            ("Ponderações",       cred.get("memoria_ponderacoes", "")),
        ])
        _spacer(doc, pts=2)

        # Citação
        _sub_gray(doc, "Citação")
        citacoes = cred.get("citacoes") or []
        _grid_table(
            doc,
            ["EXECUTADO", "MODALIDADE", "DATA", "FLS."],
            [[c.get("executado", ""), c.get("modalidade", ""), c.get("data", ""), c.get("fls", "")]
             for c in citacoes],
            [5.5, 4.4, 3.5, 3.5],
            min_rows=1,
        )
        _spacer(doc, pts=2)

        # Embargos à Execução e/ou Exceção de Pré-Executividade
        _sub_gray(doc, "Embargos e/ou Exceção")
        embargos = cred.get("embargos") or []
        if not embargos:
            embargos = [{}]
        _cont_def = {}
        for emb in embargos:
            tipo = (emb.get("tipo") or "").strip() or "Embargos à Execução"
            _cont_def[tipo] = _cont_def.get(tipo, 0) + 1
            _para(doc, f"{tipo} nº {_cont_def[tipo]}", bold=True, size=9.5, color=_TXT, before=3, after=2)
            _kv_label_table(doc, [
                ("Embargante / Excipiente", emb.get("embargante", "")),
                ("Data da Distribuição",  emb.get("data_dist", "")),
                ("Tese",                  emb.get("tese", "")),
                ("Principais Andamentos", emb.get("andamentos_resumo", "")),
                ("Status Atual",          emb.get("status", "")),
            ])
        _note(doc, "⊕ Replicar o quadro acima para cada embargo / exceção adicional identificado.")

        # Recursos
        _sub_gray(doc, "Recursos")
        recursos = cred.get("recursos") or []
        if not recursos:
            recursos = [{}]
        for ri, rec in enumerate(recursos, 1):
            _para(doc, f"Recurso nº {ri}", bold=True, size=9.5, color=_TXT, before=3, after=2)
            _kv_label_table(doc, [
                ("Recorrente",            rec.get("recorrente", "")),
                ("Decisão Recorrida",     rec.get("decisao_recorrida", "")),
                ("Data da Distribuição",  rec.get("data_dist", "")),
                ("Tese do Recurso",       rec.get("tese", "")),
                ("Principais Andamentos", rec.get("andamentos_resumo", "")),
                ("Status Atual",          rec.get("status", "")),
            ])
        _note(doc, "⊕ Replicar o quadro acima para cada recurso adicional identificado.")

        # Constrições Vigentes
        _sub_gray(doc, "Constrições Vigentes")
        constricoes = cred.get("constricoes") or []
        _grid_table(
            doc,
            ["TIPO", "DESCRIÇÃO", "VALOR (R$)", "STATUS"],
            [[c.get("tipo", ""), c.get("descricao", ""), c.get("valor", ""), c.get("status", "")]
             for c in constricoes],
            [4.0, 5.9, 3.0, 4.0],
            min_rows=1,
        )
        _note(doc, "Status possíveis: Ativa · Suspensa · Levantada · Contestada · Em discussão · "
                   "Aguardando avaliação · Aguardando expropriação")

        # Principais Andamentos Processuais
        _sub_gray(doc, "Principais Andamentos Processuais")
        andamentos = cred.get("andamentos") or []
        _grid_table(
            doc,
            ["DATA", "DESCRIÇÃO", "FLS. / EVENTO"],
            [[a.get("data", ""), a.get("descricao", ""), a.get("fls", "")] for a in andamentos],
            [2.6, 11.3, 3.0],
            min_rows=1,
        )
        _spacer(doc)

    # ══ 2. TESES DE RECUPERAÇÃO ══════════════════════════════════════════
    _sec_title(doc, "2. TESES DE RECUPERAÇÃO")

    cols_mat = ["Mat.", "Comarca", "Proprietário Atual", "Descrição do Imóvel", "Ônus Vigentes",
                "Fração", "VM (R$)", "VP (R$)", "Ônus Total (R$)", "Saldo (R$)"]
    ws_mat = [1.3, 1.7, 2.0, 2.5, 1.9, 1.2, 1.4, 1.4, 1.7, 1.4]

    # 2.1 Penhora Direta
    _sub_orange(doc, "2.1 Penhora Direta")
    _sub_gray(doc, "a. Ponderações e Observações Gerais")
    _guidance(doc, "Descrever aqui: situação geral dos imóveis localizados, gravames identificados, "
                   "bloqueios de matrícula, imóveis não localizados, pesquisas pendentes e demais "
                   "observações relevantes.")
    _sub_gray(doc, "b. Matrículas Mapeadas")
    _grid_table(doc, cols_mat, [], ws_mat, total_row=["TOTAL", "", "", "", "", "", "", "", "", ""], min_rows=2)
    _spacer(doc)

    # 2.2 IDPJ
    _sub_orange(doc, "2.2 IDPJ")
    _sub_gray(doc, "a. Resumo da Tese")
    _guidance(doc, "Fundamento principal · Histórico dos fatos · Objetivo do IDPJ · Principais elementos "
                   "probatórios identificados (procurações, representação conjunta, declarações em processo, "
                   "confusão patrimonial, etc.).")
    _sub_gray(doc, "b. Empresa-Alvo")
    _kv_table(doc, ["CAMPO", "DADO"], [
        ("Razão social da empresa-alvo",        ""),
        ("CNPJ",                                ""),
        ("Fundamentação da tese",               "Simulação / Desvio de finalidade / Confusão patrimonial"),
        ("Crédito mais adequado para propositura", ""),
    ])
    _spacer(doc, pts=2)
    _sub_gray(doc, "c. Cronologia Societária — Atos Relevantes")
    _grid_table(doc, ["DATA", "ATO", "DETALHAMENTO"], [], [2.6, 4.5, 9.8], min_rows=2)
    _spacer(doc, pts=2)
    _sub_gray(doc, "d. Evidências de Controle Informal")
    _guidance(doc, "Espaço destinado à consolidação dos principais elementos probatórios que sustentam a "
                   "tese: fatos relevantes identificados durante a investigação, documentos, prints, imagens, "
                   "pesquisas patrimoniais, alterações societárias, procurações, manifestações processuais e "
                   "demais evidências que demonstrem a viabilidade da medida.")
    _sub_gray(doc, "e. Ativos Atingíveis via IDPJ  —  ~R$ [___] (VM) / ~R$ [___] (VP)")
    _grid_table(doc, cols_mat, [], ws_mat, total_row=["TOTAL", "", "", "", "", "", "", "", "", ""], min_rows=2)
    _spacer(doc, pts=2)
    _sub_gray(doc, "f. Upside Identificado (se houver)")
    _guidance(doc, "Descrever eventual upside — ex: recebíveis futuros de compra e venda em andamento, "
                   "participações societárias, créditos a receber, outros ativos. Indicar valor estimado e "
                   "prazo de realização.")
    _spacer(doc)

    # 2.3 Fraude à Execução
    _sub_orange(doc, "2.3 Fraude à Execução")
    _sub_gray(doc, "a. Resumo da Tese")
    _guidance(doc, "Descrever os negócios jurídicos impugnáveis (doações, transferências, usufruto) "
                   "praticados no curso ou após o ajuizamento da execução. Indicar a data do ajuizamento e a "
                   "data de cada ato para enquadramento da tese. Identificar qual crédito é mais adequado "
                   "para a propositura.")
    _sub_gray(doc, "b. Má-fé e Insolvência")
    _guidance(doc, "Registrar os elementos que demonstram a má-fé do adquirente (consilium fraudis) e a "
                   "insolvência do devedor decorrente do ato — pressupostos da fraude à execução / ação pauliana.")
    _sub_gray(doc, "c. Ativos Atingíveis via Fraude à Execução  —  ~R$ [___] (VM) / ~R$ [___] (VP)")
    _grid_table(doc, cols_mat, [], ws_mat, total_row=["TOTAL", "", "", "", "", "", "", "", "", ""], min_rows=2)
    _spacer(doc)

    # 2.4 Outras teses
    _sub_orange(doc, "2.4 Outras teses")
    _guidance(doc, "Registrar outras teses de recuperação eventualmente aplicáveis ao caso.")
    _spacer(doc)

    # ══ 3. PASSIVO ═══════════════════════════════════════════════════════
    _sec_title(doc, "3. PASSIVO")
    _note(doc, "Mapear o passivo integral dos devedores e das empresas-alvo. O passivo é fator "
               "determinante na avaliação do risco real de recuperação.")

    # 3.1 Passivo Fiscal
    _sub_orange(doc, "3.1 Passivo Fiscal")
    _sub_gray(doc, "e-CAC")
    _grid_table(doc, ["NOME", "CPF / CNPJ", "SALDO e-CAC (R$)"], [], [7.0, 5.0, 4.9], min_rows=2)
    _spacer(doc, pts=2)
    _sub_gray(doc, "Execuções Fiscais")
    cols_exec = ["Nº CNJ", "Vinculado a", "Distribuição", "Exequente", "Valor Causa (R$)", "SAT Est. (R$)", "Status"]
    ws_exec = [3.0, 2.2, 2.0, 2.8, 2.3, 2.3, 2.3]
    _grid_table(doc, cols_exec, [], ws_exec, min_rows=2)
    _spacer(doc)

    # 3.2 Passivo Trabalhista
    _sub_orange(doc, "3.2 Passivo Trabalhista")
    cols_trab = ["Nº CNJ", "Vinculado a", "Distribuição", "Autor", "Valor Causa (R$)", "SAT Est. (R$)", "Status"]
    _grid_table(doc, cols_trab, [], ws_exec, min_rows=2)
    _spacer(doc)

    # 3.3 Passivo Cível (Terceiros)
    _sub_orange(doc, "3.3 Passivo Cível (Terceiros)")
    cols_civ = ["Nº CNJ", "Vinculado a", "Distribuição", "Exequente", "Valor Causa (R$)", "SAT Est. (R$)", "Obs."]
    _grid_table(doc, cols_civ, [], ws_exec, min_rows=2)
    _spacer(doc, pts=2)

    # Principais Credores
    _sub_gray(doc, "Principais Credores")
    _grid_table(
        doc,
        ["CREDOR", "NATUREZA DO CRÉDITO", "VALOR ESTIMADO (R$)", "GARANTIA / PREFERÊNCIA", "STATUS", "OBSERVAÇÕES"],
        [], [3.0, 2.9, 2.5, 3.0, 2.1, 3.4], min_rows=2,
    )
    _note(doc, "Pontos de atenção sobre o passivo: hipotecas preferentes, créditos arrematados por "
               "terceiros, passivo oculto potencial, ações de reintegração, litígios relevantes, etc.")
    _spacer(doc)

    # ══ 5. PENDÊNCIAS ════════════════════════════════════════════════════
    _sec_title(doc, "5. PENDÊNCIAS")
    _note(doc, "Campo destinado ao registro de pendências relacionadas à pesquisa de bens, obtenção de "
               "escrituras, avaliações de imóveis e demais diligências ainda necessárias para a conclusão "
               "da análise do caso.")
    _grid_table(doc, ["PENDÊNCIA", "PROTOCOLO", "PRAZO / STATUS"], [], [8.9, 4.0, 4.0], min_rows=2)
    _spacer(doc)

    # ══ 6. ELABORAÇÃO E REVISÃO ══════════════════════════════════════════
    _sec_title(doc, "6. ELABORAÇÃO E REVISÃO")
    _grid_table(
        doc,
        ["FUNÇÃO", "NOME", "DATA"],
        [["Advogada(o) Responsável pela Análise", "", ""],
         ["Advogada(o) Revisor", "", ""]],
        [8.4, 5.5, 3.0],
    )
    _spacer(doc)

    # Rodapé de confidencialidade (corpo)
    _para(doc,
          "Documento confidencial. Uso interno exclusivo do Time PPA — Invista. "
          "Vedada a reprodução ou distribuição a terceiros sem autorização prévia.",
          size=8, color=_TXT_MUTE, italic=True, before=8, after=0,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Salvar ────────────────────────────────────────────────────────────
    nome = re.sub(r"[^\w\s-]", "", dados.get("nome_caso", "") or "caso").strip().replace(" ", "_")[:40] or "caso"
    caminho = os.path.join(tempfile.gettempdir(), f"Dossie_PPA_{nome}.docx")
    doc.save(caminho)
    return caminho


def _nao_dividir_linha(row):
    """Evita que uma linha curta seja fragmentada entre páginas."""
    trp = row._tr.get_or_add_trPr()
    if trp.find(qn("w:cantSplit")) is None:
        trp.append(OxmlElement("w:cantSplit"))


def _repetir_cabecalho(row):
    trp = row._tr.get_or_add_trPr()
    if trp.find(qn("w:tblHeader")) is None:
        el = OxmlElement("w:tblHeader")
        el.set(qn("w:val"), "true")
        trp.append(el)


def _substituir_texto_paragrafo(paragraph, texto, justificar=False):
    """Troca somente o conteúdo, preservando o pPr e o rPr do template."""
    valor = _normalizar_caixa_alta(texto)
    rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        rpr = deepcopy(paragraph.runs[0]._r.rPr)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    def _aplicar_template(run, eh_referencia):
        if rpr is not None:
            if run._r.rPr is not None:
                run._r.remove(run._r.rPr)
            run._r.insert(0, deepcopy(rpr))
        if eh_referencia:
            run.italic = True

    _adicionar_runs_formatados(paragraph, valor, _aplicar_template)
    paragraph.paragraph_format.keep_together = True
    if justificar or len(valor.strip()) >= 80:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def _substituir_texto_celula(cell, texto, justificar=None):
    """Preenche uma célula sem reconstruir sua formatação ou geometria."""
    valor = _normalizar_caixa_alta(texto)
    paragraphs = list(cell.paragraphs)
    p = paragraphs[0]
    for extra in paragraphs[1:]:
        cell._tc.remove(extra._p)
    if justificar is None:
        justificar = len(valor.strip()) >= 80 or "\n" in valor
    _substituir_texto_paragrafo(p, valor, justificar=justificar)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return cell


def _chave_rotulo(texto):
    return re.sub(r"\s+", " ", str(texto or "")).strip().casefold()


def _preencher_tabela_chave_valor(table, valores):
    mapa = {_chave_rotulo(k): v for k, v in valores.items()}
    preenchidas = set()
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        chave = _chave_rotulo(row.cells[0].text)
        if chave in mapa and chave not in preenchidas:
            _substituir_texto_celula(row.cells[1], mapa[chave])
            _nao_dividir_linha(row)
            preenchidas.add(chave)


def _ajustar_linhas_dados(table, quantidade, minimo=1):
    """Clona ou remove apenas linhas de dados, preservando o cabeçalho real."""
    alvo = max(int(quantidade), minimo)
    while len(table.rows) - 1 < alvo:
        modelo = table.rows[-1]._tr
        table._tbl.append(deepcopy(modelo))
    while len(table.rows) - 1 > alvo:
        table._tbl.remove(table.rows[-1]._tr)
    if table.rows:
        _repetir_cabecalho(table.rows[0])
        _nao_dividir_linha(table.rows[0])


def _preencher_tabela_grade(table, registros, minimo=1):
    dados = list(registros or [])
    _ajustar_linhas_dados(table, len(dados), minimo=minimo)
    colunas = len(table.columns)
    for indice, row in enumerate(table.rows[1:]):
        valores = dados[indice] if indice < len(dados) else []
        for coluna in range(colunas):
            valor = valores[coluna] if coluna < len(valores) else ""
            _substituir_texto_celula(row.cells[coluna], valor)
        _nao_dividir_linha(row)


def _manter_tabela_inteira(table):
    """Mantém quadros jurídicos curtos na mesma página."""
    linhas = list(table.rows)
    for indice, row in enumerate(linhas):
        _nao_dividir_linha(row)
        manter_proxima = indice < len(linhas) - 1
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = manter_proxima
                paragraph.paragraph_format.keep_together = True


def _texto_elemento(elemento, doc):
    if elemento.tag == qn("w:p"):
        return _DocxParagraph(elemento, doc).text.strip()
    return ""


def _indice_paragrafo(elementos, doc, prefixo, inicio=0):
    alvo = prefixo.casefold()
    for indice in range(inicio, len(elementos)):
        texto = _texto_elemento(elementos[indice], doc)
        if texto.casefold().startswith(alvo):
            return indice
    raise ValueError(f"Bloco obrigatório não localizado no template: {prefixo}")


def _primeira_tabela(elementos, doc):
    for elemento in elementos:
        if elemento.tag == qn("w:tbl"):
            return _DocxTable(elemento, doc)
    raise ValueError("Tabela esperada não localizada no bloco clonado.")


def _expandir_quadros_repetidos(
    elementos, doc, prefixo_inicio, itens, titulo_item, preencher_item
):
    """Replica o quadro original até a nota ⊕, sem criar outro componente."""
    inicio = _indice_paragrafo(elementos, doc, prefixo_inicio)
    fim = None
    for indice in range(inicio + 1, len(elementos)):
        if _texto_elemento(elementos[indice], doc).startswith("⊕"):
            fim = indice
            break
    if fim is None:
        raise ValueError(f"Nota de replicação não localizada após {prefixo_inicio}.")
    if prefixo_inicio.casefold().startswith("recurso"):
        _substituir_texto_paragrafo(
            _DocxParagraph(elementos[fim], doc),
            "⊕ Replicar o quadro acima para cada recurso adicional identificado.",
        )
    padrao = [deepcopy(el) for el in elementos[inicio:fim]]
    registros = list(itens or [{}])
    novos = []
    for numero, item in enumerate(registros, 1):
        grupo = [deepcopy(el) for el in padrao]
        pidx = _indice_paragrafo(grupo, doc, prefixo_inicio)
        titulo = _substituir_texto_paragrafo(
            _DocxParagraph(grupo[pidx], doc), titulo_item(item, numero)
        )
        for elemento in grupo[pidx:]:
            if elemento.tag == qn("w:tbl"):
                break
            if elemento.tag == qn("w:p"):
                paragraph = _DocxParagraph(elemento, doc)
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True
        preencher_item(_primeira_tabela(grupo, doc), item)
        novos.extend(grupo)
    elementos[inicio:fim] = novos


def _preencher_embargo(table, embargo):
    tipo = (embargo.get("tipo") or "").casefold()
    if "exce" in tipo and len(table.rows) >= 3:
        _substituir_texto_celula(table.rows[2].cells[0], "Tese da Exceção")
    _preencher_tabela_chave_valor(table, {
        "Embargante": embargo.get("embargante", ""),
        "Embargante / Excipiente": embargo.get("embargante", ""),
        "Data da Distribuição": embargo.get("data_dist", ""),
        "Tese dos Embargos": embargo.get("tese", ""),
        "Tese da Exceção": embargo.get("tese", ""),
        "Principais Andamentos": embargo.get("andamentos_resumo", ""),
        "Status Atual": embargo.get("status", ""),
    })
    _manter_tabela_inteira(table)


def _preencher_recurso(table, recurso):
    _preencher_tabela_chave_valor(table, {
        "Recorrente": recurso.get("recorrente", ""),
        "Decisão Recorrida": recurso.get("decisao_recorrida", ""),
        "Data da Distribuição": recurso.get("data_dist", ""),
        "Tese do Recurso": recurso.get("tese", ""),
        "Principais Andamentos": recurso.get("andamentos_resumo", ""),
        "Status Atual": recurso.get("status", ""),
    })
    _manter_tabela_inteira(table)


def _iterar_elementos_com_titulo(elementos, doc):
    titulo = ""
    for elemento in elementos:
        if elemento.tag == qn("w:p"):
            texto = _DocxParagraph(elemento, doc).text.strip()
            if texto:
                titulo = texto
        elif elemento.tag == qn("w:tbl"):
            yield titulo, _DocxTable(elemento, doc)


def _preencher_bloco_credito(elementos, doc, credito, indice_credito):
    """Preenche uma cópia do bloco 1.1 e replica defesas/recursos existentes."""
    cid = credito.get("id") or f"Crédito {indice_credito}"
    titulo_idx = _indice_paragrafo(elementos, doc, "1.1 Crédito")
    titulo_credito = _substituir_texto_paragrafo(
        _DocxParagraph(elementos[titulo_idx], doc),
        f"1.{indice_credito} {cid}",
    )
    if indice_credito > 1:
        titulo_credito.paragraph_format.page_break_before = True

    embargos = credito.get("embargos") or []
    _expandir_quadros_repetidos(
        elementos,
        doc,
        "Embargo nº 1",
        embargos,
        lambda item, n: (
            f"{(item.get('tipo') or 'Embargo à Execução')} nº {n}"
        ),
        _preencher_embargo,
    )
    recursos = credito.get("recursos") or []
    _expandir_quadros_repetidos(
        elementos,
        doc,
        "Recurso nº 1",
        recursos,
        lambda item, n: f"Recurso nº {n}",
        _preencher_recurso,
    )

    for titulo, table in _iterar_elementos_com_titulo(elementos, doc):
        chave = titulo.casefold()
        if chave == "dados do processo":
            _preencher_tabela_chave_valor(table, {
                "Número do processo": credito.get("numero_processo", ""),
                "Vara / Comarca": credito.get("vara_comarca", ""),
                "Exequente": credito.get("exequente_info", ""),
                "Executado(s)": credito.get("executados_info", ""),
                "Data de distribuição": credito.get("data_distribuicao", ""),
                "SOP": credito.get("sop", ""),
                "SAT": credito.get("sat", ""),
                "Critério de atualização do SAT": credito.get("criterio_sat", ""),
                "Honorários": credito.get("honorarios", ""),
                "Lastro / Instrumento": credito.get("lastro", ""),
                "Data de Emissão": credito.get("data_emissao", ""),
                "Data do Vencimento": credito.get("data_vencimento", ""),
                "Assinaturas": credito.get("assinaturas", ""),
                "Garantia": credito.get("garantia", ""),
                "Status do Processo": credito.get("status_processo", ""),
            })
        elif chave == "índices de correção do contrato":
            _preencher_tabela_chave_valor(table, {
                "Correção monetária": credito.get("ind_cm", ""),
                "Juros remuneratórios": credito.get("ind_jr", ""),
                "Juros moratórios": credito.get("ind_jm", ""),
                "Multa moratória": credito.get("ind_multa", ""),
                "Capitalização": credito.get("ind_cap", ""),
            })
        elif chave == "planilha inicial":
            _preencher_tabela_chave_valor(table, {
                "Correção monetária": credito.get("plan_cm", ""),
                "Juros remuneratórios": credito.get("plan_jr", ""),
                "Multa moratória": credito.get("plan_multa", ""),
                "Capitalização": credito.get("plan_cap", ""),
                "Ponderações": credito.get("plan_ponderacoes", ""),
            })
        elif chave == "última memória de cálculo":
            _preencher_tabela_chave_valor(table, {
                "Data da Juntada": credito.get("memoria_data_juntada", ""),
                "Total Atualizado": credito.get("memoria_total", ""),
                "Data-base": credito.get("memoria_data_base", ""),
                "Índices aplicados": credito.get("memoria_indices", ""),
                "Ponderações": credito.get("memoria_ponderacoes", ""),
            })
        elif chave == "citação":
            _preencher_tabela_grade(
                table,
                [[
                    c.get("executado", ""),
                    c.get("modalidade", ""),
                    c.get("data", ""),
                    c.get("fls", ""),
                ] for c in (credito.get("citacoes") or [])],
                minimo=1,
            )
        elif chave == "constrições vigentes":
            _preencher_tabela_grade(
                table,
                [[
                    c.get("tipo", ""),
                    c.get("descricao", ""),
                    c.get("valor", ""),
                    c.get("status", ""),
                ] for c in (credito.get("constricoes") or [])],
                minimo=1,
            )
        elif chave == "principais andamentos processuais":
            _preencher_tabela_grade(
                table,
                [[
                    a.get("data", ""),
                    a.get("descricao", ""),
                    a.get("fls", ""),
                ] for a in (credito.get("andamentos") or [])],
                minimo=1,
            )


def _preencher_creditos_template(doc, creditos):
    body = doc.element.body
    elementos = list(body.iterchildren())
    inicio = _indice_paragrafo(elementos, doc, "1.1 Crédito")
    placeholder = _indice_paragrafo(elementos, doc, "1.2 Crédito", inicio + 1)
    secao_dois = _indice_paragrafo(elementos, doc, "2. TESES DE RECUPERAÇÃO", placeholder + 1)
    padrao = [deepcopy(el) for el in elementos[inicio:placeholder]]
    ancora = elementos[secao_dois]
    for elemento in elementos[inicio:secao_dois]:
        body.remove(elemento)

    lista = list(creditos or [{}])
    for numero, credito in enumerate(lista, 1):
        bloco = [deepcopy(el) for el in padrao]
        _preencher_bloco_credito(bloco, doc, credito, numero)
        for elemento in bloco:
            ancora.addprevious(elemento)


def _texto_analise(valor) -> str:
    """Consolida itens jurídicos no quadro narrativo original do template."""
    if not valor:
        return ""
    if isinstance(valor, str):
        return valor.strip()
    if isinstance(valor, dict):
        valor = [valor]
    trechos = []
    for item in valor:
        if not isinstance(item, dict):
            if str(item).strip():
                trechos.append(str(item).strip())
            continue
        titulo = str(item.get("titulo") or "").strip()
        texto = str(
            item.get("texto")
            or item.get("descricao")
            or item.get("detalhamento")
            or item.get("analise")
            or ""
        ).strip()
        referencia = str(item.get("referencia") or item.get("fonte") or "").strip()
        trecho = f"{titulo}: {texto}" if titulo and texto else (titulo or texto)
        if referencia and referencia.casefold() not in trecho.casefold():
            trecho = f"{trecho} ({referencia})" if trecho else referencia
        if trecho:
            trechos.append(trecho)
    return "\n\n".join(trechos)


def _valor_monetario(texto):
    valor = str(texto or "")
    match = re.search(r"-?\d[\d.\s]*(?:,\d{1,2})?", valor)
    if not match:
        return None
    numero = match.group(0).replace(" ", "").replace(".", "").replace(",", ".")
    try:
        return float(numero)
    except ValueError:
        return None


def _somar_moeda(registros, campo):
    valores = [
        numero
        for numero in (_valor_monetario(item.get(campo)) for item in registros)
        if numero is not None
    ]
    if not valores:
        return ""
    return _fmt_valor_br(sum(valores))


def _ajustar_linhas_com_total(table, quantidade):
    """Ajusta linhas sem perder a linha TOTAL estilizada do template."""
    alvo = max(int(quantidade), 1)
    while len(table.rows) - 2 < alvo:
        modelo = table.rows[-2]._tr
        table.rows[-1]._tr.addprevious(deepcopy(modelo))
    while len(table.rows) - 2 > alvo:
        table._tbl.remove(table.rows[-2]._tr)
    _repetir_cabecalho(table.rows[0])
    _nao_dividir_linha(table.rows[0])


def _preencher_resumo_ativos(table, registros, dados):
    itens = list(registros or [])
    _ajustar_linhas_com_total(table, len(itens))
    for indice, row in enumerate(table.rows[1:-1]):
        item = itens[indice] if indice < len(itens) else {}
        valores = [
            item.get("tese", ""),
            item.get("vm", ""),
            item.get("vp", ""),
            item.get("onus", ""),
            item.get("observacoes", ""),
        ]
        for coluna, valor in enumerate(valores):
            _substituir_texto_celula(row.cells[coluna], valor)
        _nao_dividir_linha(row)
    total = table.rows[-1]
    totais = [
        "TOTAL GERAL",
        dados.get("total_atingivel_vm") or _somar_moeda(itens, "vm"),
        dados.get("total_atingivel_vp") or _somar_moeda(itens, "vp"),
        _somar_moeda(itens, "onus"),
        "",
    ]
    for coluna, valor in enumerate(totais):
        _substituir_texto_celula(total.cells[coluna], valor)
    _nao_dividir_linha(total)


def _filtrar_ativos(ativos, tese):
    alvo = tese.casefold()
    return [
        item for item in (ativos or [])
        if alvo in str(item.get("tese") or "").casefold()
    ]


def _preencher_ativos_resumidos(table, ativos):
    _preencher_tabela_grade(
        table,
        [[
            item.get("matricula", ""),
            item.get("tipo_ativo") or item.get("descricao", ""),
            item.get("area", ""),
            item.get("situacao_produtiva", ""),
            item.get("liquidez", ""),
            item.get("fracao_atingivel", ""),
            item.get("onus_vigentes") or item.get("onus_total", ""),
            item.get("observacoes", ""),
        ] for item in (ativos or [])],
        minimo=1,
    )


def _compactar_celula(cell, tamanho, centralizar=False):
    _cell_pad(cell, top=30, left=25, bottom=30, right=25)
    for paragraph in cell.paragraphs:
        if centralizar:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(tamanho)


def _adicionar_coluna(table, header_texto: str):
    """Insere uma coluna nova ao final da tabela via XML — python-docx não tem API de alto
    nível pra isso. Clona o último <w:gridCol> e a última <w:tc> de cada linha (cabeçalho,
    dados, total), zerando o texto da célula clonada mas preservando a formatação/largura."""
    tbl = table._tbl
    ultima_col = tbl.tblGrid.gridCol_lst[-1]
    tbl.tblGrid.append(deepcopy(ultima_col))

    for row in table.rows:
        tr = row._tr
        tcs = tr.findall(qn("w:tc"))
        nova_tc = deepcopy(tcs[-1])
        for t in nova_tc.findall(".//" + qn("w:t")):
            t.text = ""
        tr.append(nova_tc)

    _substituir_texto_celula(table.rows[0].cells[-1], header_texto)


def _preencher_ativos_detalhados(table, ativos, coluna_extra=None):
    """coluna_extra: função opcional item -> valor, pra uma 11ª coluna já inserida na
    tabela via _adicionar_coluna (usado hoje só na tabela de Fraude à Execução, coluna
    'Transmissões')."""
    itens = list(ativos or [])
    _ajustar_linhas_com_total(table, len(itens))
    for indice, row in enumerate(table.rows[1:-1]):
        item = itens[indice] if indice < len(itens) else {}
        valores = [
            item.get("matricula", ""),
            item.get("comarca", ""),
            item.get("proprietario_atual", ""),
            item.get("descricao") or item.get("tipo_ativo", ""),
            item.get("onus_vigentes", ""),
            item.get("fracao_atingivel", ""),
            item.get("vm", ""),
            item.get("vp", ""),
            item.get("onus_total", ""),
            item.get("saldo", ""),
        ]
        if coluna_extra is not None:
            valores.append(coluna_extra(item))
        for coluna, valor in enumerate(valores):
            _substituir_texto_celula(row.cells[coluna], valor)
            _compactar_celula(
                row.cells[coluna],
                5.3 if coluna >= 6 else 6.2,
                centralizar=coluna >= 5,
            )
        _nao_dividir_linha(row)
    total = table.rows[-1]
    totais = [
        "TOTAL", "", "", "", "", "",
        _somar_moeda(itens, "vm"),
        _somar_moeda(itens, "vp"),
        _somar_moeda(itens, "onus_total"),
        _somar_moeda(itens, "saldo"),
    ]
    if coluna_extra is not None:
        totais.append("")
    for coluna, valor in enumerate(totais):
        _substituir_texto_celula(total.cells[coluna], valor)
        _compactar_celula(
            total.cells[coluna],
            5.3 if coluna >= 6 else 6.2,
            centralizar=coluna >= 5,
        )
    _nao_dividir_linha(total)


def _preencher_ativos_teses_narrativas(doc, ativos: list):
    """Preenche as tabelas de imóveis das subseções narrativas de '2. TESES DE
    RECUPERAÇÃO' (Penhora Direta / IDPJ / Fraude à Execução) a partir da planilha de Coleta
    de Informações, casando por palavra-chave no texto da Tese de cada item — um mesmo
    imóvel entra em mais de uma tabela quando a Tese cita mais de uma palavra-chave (ex.:
    "Fraude à Execução/IDPJ"), já que _filtrar_ativos roda uma vez por tabela/seção.
    Localiza as tabelas pelo texto do subtítulo (robusto a deslocamento de índice causado
    pelos blocos dinâmicos de Tese em "VISÃO GERAL DOS ATINGÍVEIS", que mudam de tamanho
    conforme o caso) — nunca por índice fixo de tabela."""
    if not ativos:
        return
    mapa = [
        ("matrículas mapeadas", "penhora direta", False),
        ("ativos atingíveis via idpj", "idpj", False),
        ("ativos atingíveis via fraude à execução", "fraude à execução", True),
    ]
    for heading, table in _iter_headings_tables(doc):
        h = heading.lower()
        for chave_heading, chave_tese, com_transmissoes in mapa:
            if chave_heading in h:
                itens = _filtrar_ativos(ativos, chave_tese)
                if com_transmissoes:
                    if len(table.columns) == 10:
                        _adicionar_coluna(table, "Transmissões")
                    _preencher_ativos_detalhados(table, itens, coluna_extra=lambda item: item.get("transmissoes", ""))
                else:
                    _preencher_ativos_detalhados(table, itens)
                break


def _preencher_paragrafo_apos_titulo(doc, titulo, texto):
    if not str(texto or "").strip():
        return
    paragraphs = doc.paragraphs
    inicio = next(
        (
            indice for indice, paragraph in enumerate(paragraphs)
            if paragraph.text.strip().casefold().startswith(titulo.casefold())
        ),
        None,
    )
    if inicio is None:
        return
    candidatos = []
    for paragraph in paragraphs[inicio + 1:]:
        atual = paragraph.text.strip()
        if atual and re.match(r"^(?:\d+(?:\.\d+)*\.?|[a-z]\.)\s", atual, re.IGNORECASE):
            break
        candidatos.append(paragraph)
    if not candidatos:
        return
    destino = next((p for p in candidatos if p.text.strip()), candidatos[0])
    _substituir_texto_paragrafo(destino, texto, justificar=True)


def _atualizar_titulo_valores(doc, prefixo, resumo, tese):
    item = next(
        (
            registro for registro in (resumo or [])
            if tese.casefold() in str(registro.get("tese") or "").casefold()
        ),
        {},
    )
    vm, vp = item.get("vm", ""), item.get("vp", "")
    if not (vm or vp):
        return
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().casefold().startswith(prefixo.casefold()):
            _substituir_texto_paragrafo(
                paragraph,
                f"{prefixo} — {vm or 'R$ [___]'} (VM) / {vp or 'R$ [___]'} (VP)",
            )
            paragraph.paragraph_format.keep_with_next = True
            break


def _preencher_visao_ativos_template(doc, dados):
    resumo = dados.get("visao_consolidada_ativos") or []
    ativos = dados.get("ativos") or []
    if resumo:
        _preencher_resumo_ativos(doc.tables[2], resumo, dados)
    if ativos:
        _preencher_ativos_resumidos(
            doc.tables[3], _filtrar_ativos(ativos, "Penhora Direta")
        )
        _preencher_ativos_resumidos(
            doc.tables[4], _filtrar_ativos(ativos, "IDPJ")
        )
        _preencher_ativos_resumidos(
            doc.tables[5], _filtrar_ativos(ativos, "Fraude à Execução")
        )


def _preencher_teses_template(doc, dados):
    if len(doc.tables) < 25:
        raise ValueError("Template oficial incompleto: quadros das teses não localizados.")
    teses = dados.get("teses_recuperacao") or {}
    ativos = dados.get("ativos") or []
    resumo_ativos = dados.get("visao_consolidada_ativos") or []
    penhora = teses.get("penhora_direta") or {}
    idpj = teses.get("idpj") or {}
    fraude = teses.get("fraude_execucao") or {}

    texto = _texto_analise(penhora.get("analise"))
    if texto:
        _substituir_texto_celula(doc.tables[16].cell(0, 0), texto, justificar=True)
    penhora_ativos = _filtrar_ativos(ativos, "Penhora Direta")
    if penhora_ativos:
        _preencher_ativos_detalhados(doc.tables[17], penhora_ativos)

    texto = _texto_analise(idpj.get("resumo"))
    if texto:
        _substituir_texto_celula(doc.tables[18].cell(0, 0), texto, justificar=True)
    empresa = idpj.get("empresa_alvo") or {}
    if empresa:
        _preencher_tabela_chave_valor(doc.tables[19], {
            "Razão social da empresa-alvo": empresa.get("razao_social", ""),
            "CNPJ": empresa.get("cnpj", ""),
            "CNAE principal": empresa.get("cnae_principal", ""),
            "Atividades secundárias": empresa.get("atividades_secundarias", ""),
            "Sócio atual": empresa.get("socio_atual", ""),
            "Endereço fiscal": empresa.get("endereco_fiscal", ""),
            "Fundamentação da tese": empresa.get("fundamentacao", ""),
            "Crédito mais adequado para propositura": empresa.get("credito_propositura", ""),
        })
    cronologia = idpj.get("cronologia") or []
    if cronologia:
        _preencher_tabela_grade(
            doc.tables[20],
            [[
                item.get("data", ""),
                item.get("ato", ""),
                " — ".join(filter(None, [
                    item.get("detalhamento", ""),
                    item.get("referencia", ""),
                ])),
            ] for item in cronologia],
            minimo=1,
        )
    _preencher_paragrafo_apos_titulo(
        doc, "d. Evidências de Controle Informal", _texto_analise(idpj.get("evidencias"))
    )
    idpj_ativos = _filtrar_ativos(ativos, "IDPJ")
    if idpj_ativos:
        _preencher_ativos_detalhados(doc.tables[21], idpj_ativos)
    _atualizar_titulo_valores(
        doc, "e. Ativos Atingíveis via IDPJ", resumo_ativos, "IDPJ"
    )
    texto = _texto_analise(idpj.get("upside"))
    if texto:
        _substituir_texto_celula(doc.tables[22].cell(0, 0), texto, justificar=True)

    texto = _texto_analise(fraude.get("resumo"))
    if texto:
        _substituir_texto_celula(doc.tables[23].cell(0, 0), texto, justificar=True)
    _preencher_paragrafo_apos_titulo(
        doc, "b. Má-fé e Insolvência",
        _texto_analise(fraude.get("ma_fe_insolvencia")),
    )
    fraude_ativos = _filtrar_ativos(ativos, "Fraude à Execução")
    if fraude_ativos:
        _preencher_ativos_detalhados(doc.tables[24], fraude_ativos)
    _atualizar_titulo_valores(
        doc, "c. Ativos Atingíveis via Fraude à Execução",
        resumo_ativos,
        "Fraude à Execução",
    )
    _preencher_paragrafo_apos_titulo(
        doc, "2.4 Outras teses", _texto_analise(teses.get("outras"))
    )


def _build_doc(dados: dict) -> str:
    """Preenche uma cópia do template oficial, sem reconstruir seu design."""
    if not _TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Template oficial não localizado: {_TEMPLATE_PATH}"
        )
    dados = _normalizar_dados(dados or {})
    nome = re.sub(
        r"[^\w\s-]", "", dados.get("nome_caso", "") or "caso"
    ).strip().replace(" ", "_")[:40] or "caso"
    caminho = os.path.join(tempfile.gettempdir(), f"Dossie_PPA_{nome}.docx")
    shutil.copy2(_TEMPLATE_PATH, caminho)
    doc = Document(caminho)

    hoje = _date.today().strftime("%d/%m/%Y")
    _preencher_tabela_chave_valor(doc.tables[0], {
        "Nome do Caso": dados.get("nome_caso", ""),
        "Data da Análise": dados.get("data_analise") or hoje,
        "Advogada Responsável": dados.get("advogada_responsavel", ""),
    })

    pf = dados.get("passivo_fiscal", "")
    pt = dados.get("passivo_trabalhista", "")
    pc = dados.get("passivo_civel", "")
    _preencher_tabela_chave_valor(doc.tables[1], {
        "Exequente(s)": dados.get("exequentes", ""),
        "Executado(s)": dados.get("executados", ""),
        "SAT": dados.get("sat_total", ""),
        "Total atingível mapeado — VM": dados.get("total_atingivel_vm", ""),
        "Total atingível mapeado — VP": dados.get("total_atingivel_vp", ""),
        "Tese(s) principal(is)": dados.get("teses_principais", ""),
        "Risco Jurídico": dados.get("risco_juridico", ""),
    })
    if len(doc.tables[1].rows) >= 10:
        _substituir_texto_celula(
            doc.tables[1].rows[6].cells[1], f"Fiscal: {pf}" if pf else "Fiscal:"
        )
        _substituir_texto_celula(
            doc.tables[1].rows[7].cells[1],
            f"Trabalhista: {pt}" if pt else "Trabalhista:",
        )
        _substituir_texto_celula(
            doc.tables[1].rows[8].cells[1], f"Cível: {pc}" if pc else "Cível:"
        )
        total_passivo = dados.get("passivo_total", "")
        _substituir_texto_celula(
            doc.tables[1].rows[9].cells[1],
            f"TOTAL: {total_passivo}" if total_passivo else "TOTAL:",
        )

    consideracoes = (dados.get("consideracoes_gerais") or "").strip()
    if consideracoes:
        _substituir_texto_celula(doc.tables[6].cell(0, 0), consideracoes, justificar=True)
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith("Campo livre para registro"):
                _substituir_texto_paragrafo(paragraph, "")
                break

    _preencher_visao_ativos_template(doc, dados)
    _preencher_teses_template(doc, dados)
    _preencher_creditos_template(doc, dados.get("creditos") or [])

    doc.save(caminho)
    return caminho


# ══════════════════════════════════════════════════════════════════════════
# Entry point
# ══════════════════════════════════════════════════════════════════════════

def gerar_dossie_word(texto_completo: str, relatorio: str, client, model: str) -> str:
    """Extrai dados do texto OCR completo (+ relatório de apoio) e gera o Word PPA v3.0."""
    dados = _extrair_dados(texto_completo, relatorio, client, model)
    # Passada final: confere campos vazios contra o relatório e completa o que faltou
    dados = _completar_com_relatorio(dados, relatorio, client, model)
    return _build_doc(dados)


# ══════════════════════════════════════════════════════════════════════════
# Preenchimento do PASSIVO (Seção 3) a partir de dados da Predictus
# ══════════════════════════════════════════════════════════════════════════

def _fmt_valor_br(v) -> str:
    if v is None or v == "":
        return ""
    try:
        return "R$ " + f"{float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return str(v)


def _proc_to_row(p: dict) -> list:
    """Mapeia um processo (Predictus ou Coleta de Informações) para as 7 colunas do
    passivo do dossiê. "sat" só vem preenchido no fluxo da Coleta (Saldo Atualizado
    estimado, já calculado na própria planilha de origem) — a Predictus não tem essa
    informação, e _fmt_valor_br(None) devolve "" nesse caso."""
    return [
        p.get("cnj", ""),                 # Nº CNJ
        p.get("vinc", ""),                # Executado
        p.get("data", "") or "",          # Distribuição
        p.get("ativo", ""),               # Exequente / Autor (polo ativo)
        _fmt_valor_br(p.get("valor")),    # Valor Causa (R$)
        _fmt_valor_br(p.get("sat")),      # SAT Est. (R$)
        p.get("status", ""),              # Status / Obs.
    ]


def _fill_passivo_table(table, registros: list):
    """Substitui as linhas de dados de uma tabela do passivo (mantém o header)."""
    ncols = len(table.columns)
    larguras = [c.width for c in table.rows[0].cells]
    _repetir_cabecalho(table.rows[0])
    _nao_dividir_linha(table.rows[0])
    for header_cell in table.rows[0].cells:
        if _chave_rotulo(header_cell.text) == "vinculado a":
            _substituir_texto_celula(header_cell, "Executado")
    # remove linhas de dados (mantém a primeira = header laranja)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    if not registros:
        # deixa uma linha vazia para não colar o próximo bloco no header
        registros = [[""] * ncols]
    for reg in registros:
        r = table.add_row()
        _nao_dividir_linha(r)
        for i in range(ncols):
            c = r.cells[i]
            _set_bg(c, _BRANCO); _set_borders(c); _cell_pad(c)
            valor = reg[i] if i < len(reg) else ""
            _write(
                c,
                _normalizar_caixa_alta(valor),
                align=WD_ALIGN_PARAGRAPH.JUSTIFY
                if len(str(valor).strip()) >= 80 else None,
            )
            if i < len(larguras) and larguras[i]:
                c.width = larguras[i]


def _iter_headings_tables(doc):
    """Itera o corpo em ordem, associando cada tabela ao último subtítulo não-vazio."""
    last_heading = ""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            txt = _DocxParagraph(child, doc).text.strip()
            if txt:
                last_heading = txt
        elif child.tag == qn("w:tbl"):
            yield last_heading, _DocxTable(child, doc)


def preencher_passivo_dossie(dossie_path, fiscais: list, trabalhistas: list, civeis: list) -> str:
    """
    Preenche as tabelas da Seção 3 (Passivo) do dossiê — Execuções Fiscais,
    Trabalhista e Cível — com os processos consolidados da Predictus.
    Não toca em 'Principais Credores' nem no e-CAC.
    Se dossie_path for None, gera um dossiê novo (esqueleto) e preenche nele.
    """
    if dossie_path:
        doc = Document(dossie_path)
    else:
        doc = Document(_build_doc({}))

    for heading, table in _iter_headings_tables(doc):
        h = heading.lower()
        if "execuções fiscais" in h or "execucoes fiscais" in h:
            _fill_passivo_table(table, [_proc_to_row(p) for p in fiscais])
        elif "trabalhista" in h:
            _fill_passivo_table(table, [_proc_to_row(p) for p in trabalhistas])
        elif "cível" in h or "civel" in h:
            _fill_passivo_table(table, [_proc_to_row(p) for p in civeis])

    caminho = os.path.join(tempfile.gettempdir(), "Dossie_PPA_atualizado.docx")
    doc.save(caminho)
    return caminho


# ══════════════════════════════════════════════════════════════════════════
# Ativos Atingíveis (por Tese, dinâmico) + Passivo — a partir da Coleta de Informações
# ══════════════════════════════════════════════════════════════════════════

def _preencher_visao_geral_atingiveis_dinamica(doc, ativos_visao_geral: list):
    """Substitui os blocos fixos título+tabela de 'VISÃO GERAL DOS ATINGÍVEIS' (hoje Penhora
    Direta/IDPJ/Fraude à Execução) por um bloco por Tese distinta encontrada no Excel de
    Coleta de Informações — títulos 100% dinâmicos, batendo com o texto exato da Tese.

    Segue o mesmo padrão de clonagem já usado em _preencher_creditos_template: acha o
    intervalo do bloco original, usa o 1º par (título+tabela) como modelo, remove os
    originais do corpo do documento e insere uma cópia preenchida por tese antes da âncora
    (o próximo heading, "PRINCIPAIS CONSIDERAÇÕES")."""
    if not ativos_visao_geral:
        return
    body = doc.element.body
    elementos = list(body.iterchildren())
    inicio = _indice_paragrafo(elementos, doc, "VISÃO GERAL DOS ATINGÍVEIS") + 1
    fim = _indice_paragrafo(elementos, doc, "PRINCIPAIS CONSIDERAÇÕES", inicio)

    padrao = [deepcopy(el) for el in elementos[inicio:inicio + 2]]  # 1 título + 1 tabela

    ancora = elementos[fim]
    for elemento in elementos[inicio:fim]:
        body.remove(elemento)

    for item in ativos_visao_geral:
        bloco = [deepcopy(el) for el in padrao]
        titulo_par = _DocxParagraph(bloco[0], doc)
        _substituir_texto_paragrafo(titulo_par, item.get("tese", ""))
        titulo_par.paragraph_format.keep_with_next = True
        tabela = _DocxTable(bloco[1], doc)
        _preencher_tabela_grade(tabela, item.get("linhas") or [], minimo=1)
        for elemento in bloco:
            ancora.addprevious(elemento)


def preencher_ativos_e_passivo_dossie(dossie_path, resumo_ativos: list, ativos_visao_geral: list,
                                       fiscais: list, trabalhistas: list, civeis: list, ecac: list,
                                       ativos_detalhados: list = None) -> str:
    """
    Preenche Ativos Atingíveis (Visão Consolidada + Visão Geral, dinâmico por Tese), as
    tabelas de imóveis das subseções narrativas de "2. TESES DE RECUPERAÇÃO" (Penhora
    Direta/IDPJ/Fraude à Execução, se ativos_detalhados for informado) e o Passivo
    (Fiscal/Trabalhista/Cível/e-CAC) de um dossiê PPA a partir dos dados já extraídos do
    Excel de Coleta de Informações. Se dossie_path for None, gera um dossiê novo (esqueleto).

    Não mexe no restante do conteúdo narrativo de "2. TESES DE RECUPERAÇÃO" (Resumo da
    Tese, Empresa-Alvo, Cronologia Societária etc.) nem em "Principais Credores" — são
    conteúdo de análise jurídica manual, não derivável do Excel.
    """
    if dossie_path:
        doc = Document(dossie_path)
    else:
        doc = Document(_build_doc({}))

    if resumo_ativos:
        _preencher_resumo_ativos(doc.tables[2], resumo_ativos, {})
    if ativos_visao_geral:
        _preencher_visao_geral_atingiveis_dinamica(doc, ativos_visao_geral)
    if ativos_detalhados:
        _preencher_ativos_teses_narrativas(doc, ativos_detalhados)

    for heading, table in _iter_headings_tables(doc):
        h = heading.lower()
        if "e-cac" in h:
            _fill_passivo_table(table, [
                [e.get("nome", ""), e.get("cpf_cnpj", ""), _fmt_valor_br(e.get("saldo"))]
                for e in (ecac or [])
            ])
        elif "execuções fiscais" in h or "execucoes fiscais" in h:
            _fill_passivo_table(table, [_proc_to_row(p) for p in (fiscais or [])])
        elif "trabalhista" in h:
            _fill_passivo_table(table, [_proc_to_row(p) for p in (trabalhistas or [])])
        elif "cível" in h or "civel" in h:
            _fill_passivo_table(table, [_proc_to_row(p) for p in (civeis or [])])

    caminho = os.path.join(tempfile.gettempdir(), "Dossie_PPA_com_coleta.docx")
    doc.save(caminho)
    return caminho
