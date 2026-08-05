# -*- coding: utf-8 -*-
import os
import re
import json
import math
import time
import tempfile
import threading
import traceback
import concurrent.futures
from contextlib import nullcontext
from pathlib import Path
from typing import Optional
from collections import defaultdict

import fitz
import gradio as gr
from docx import Document
from google.genai import types

from report_template_rj import REPORT_TEMPLATE_RJ, SYSTEM_PROMPT_RJ
from utils import (
    _retry, _gerar_docx, _responder_pergunta_generica, _get_gemini_clients,
    _executar_com_failover_gemini,
    _barra_progresso, _filtrar_arquivos_existentes, _paginas_digitalizadas_pdf,
    GEMINI_MODEL_EXTRACAO, GEMINI_MODEL_OCR, GEMINI_MODEL_RELATORIO,
    GEMINI_MODEL_ESTRUTURADO, GEMINI_MODEL_QA,
)
from analysis_runtime import (
    ANALYSIS_MANAGER, cleanup_chunk, file_sha256, iter_pdf_chunks,
    pdf_preparation_slot, load_chunk_cache, queue_message, record_status,
    save_chunk_cache,
)
from checklist_rj import gerar_checklist_rj, gerar_checklist_creditos, _montar_fonte_rj
import rj_cache
from legal_prompts import (
    REGRA_EXTRACAO_POR_PAGINA,
    REGRA_FIDELIDADE_PROCESSUAL,
    REGRA_INDICES_E_ADITAMENTOS,
    REGRA_CRONOLOGIA_PROCESSUAL,
    REGRAS_CONSOLIDACAO_PROCESSUAL,
    contexto_fonte_pdf,
    normalizar_referencias_relatorio,
)

CHUNK_MAX_PAGES_RJ    = 400
AVG_MIN_POR_CHUNK_RJ  = 3
MIN_CONSOLIDACAO_RJ   = 8
LIMITE_TRUNCAMENTO_CHECKLIST = 900_000   # mesmo limite usado em checklist_rj._extrair
LIMITE_TRUNCAMENTO_CREDORES  = 150_000   # mesmo limite usado em _extrair_credores_json
LIMITE_CONSOLIDACAO_RJ       = 3_000_000  # teto de seguranca p/ nao estourar contexto do modelo
LIMITE_LOTE_FIDELIDADE_RJ    = 900_000


def _aviso_truncamento(tamanho: int, limite: int) -> str:
    if tamanho <= limite:
        return ""
    return (
        f" ⚠️ Fonte com {tamanho:,} caracteres — só os primeiros {limite:,} foram "
        "considerados; pode haver informação não coberta."
    )
# MÓDULO — ANÁLISE DE RECUPERAÇÃO JUDICIAL (Teste B)
# ══════════════════════════════════════════════════════════════════════════════

_rj_cache_nome: dict[str, str] = {}
_rj_cache_lock = threading.Lock()

PROMPT_RJ = (
    "Voce esta analisando um trecho de um processo de Recuperacao Judicial brasileiro.\n"
    "IMPORTANTE: este arquivo pode ser um fragmento de um processo maior dividido em multiplos PDFs "
    "pelo tribunal, OU pode ser um processo independente. Extraia o numero do processo se disponivel — "
    "isso sera usado para detectar continuidade na consolidacao. Trate cada fragmento como parte de "
    "um documento continuo; a consolidacao final verificara se sao o mesmo processo ou processos distintos.\n"
    "Extraia COM MAXIMO DETALHE todas as informacoes presentes nestas paginas.\n"
    "Para paginas escaneadas: aplique OCR visual completo.\n\n"
    "REFERENCIA OBRIGATORIA EM TODA INFORMACAO EXTRAIDA: cada pagina do processo tem um identificador "
    "proprio do sistema do tribunal, estampado no cabecalho/rodape/lateral da pagina — 'fls. NN' (fisico/"
    "eSAJ), 'Mov. NN' (PJe), 'Evento NN' (Eproc) ou 'ID NNNNNN' (Projudi/outro). Ao extrair QUALQUER dado "
    "(valor, data, indice, nome, garantia, andamento, clausula contratual etc.), registre junto o "
    "identificador exato da pagina/movimento onde aquele dado especifico aparece — nao apenas uma vez no "
    "topo do trecho, mas ao lado de cada informacao. Se essa referencia nao for capturada agora, ela se "
    "perde e nao pode ser reconstruida depois. Nunca invente ou estime uma referencia — se nao conseguir "
    "ler o identificador da pagina (ex.: digitalizacao ilegivel), registre isso explicitamente em vez de "
    "adivinhar.\n\n"
    "Cubra TODOS os itens encontrados:\n"
    "- Recuperandos (nome, CPF/CNPJ, papel)\n"
    "- Administrador Judicial\n"
    "- Advogados e escritorios (nome, OAB)\n"
    "- Datas relevantes: pedido, deferimento, AGC, PRJ, RMA\n"
    "- Status do PRJ, AGC, RMA, QGC\n"
    "- Stay period (prazo, prorrogacoes)\n"
    "- Endividamento fiscal e tributario\n"
    "- Imoveis, contratos, CCBs (numero, valor, garantia, avalistas, indices)\n"
    "- Creditos extraconcursais\n"
    "- Execucoes relacionadas (numero, partes, valor, status, andamentos)\n"
    "- Divergencias, impugnacoes e habilitacoes de credito\n"
    "- Andamentos processuais com datas (TODOS)\n"
    "Nao omita nada. Nao invente. Se nao encontrar, diga explicitamente.\n\n"
    + REGRA_FIDELIDADE_PROCESSUAL
    + "\n"
    + REGRA_EXTRACAO_POR_PAGINA
    + "\n"
    + REGRA_INDICES_E_ADITAMENTOS
    + "\n"
    + REGRA_CRONOLOGIA_PROCESSUAL
)


def _rj_dividir_pdf(path: str) -> list:
    chunks = []
    try:
        with pdf_preparation_slot():
            for chunk in iter_pdf_chunks(path, max_pages=CHUNK_MAX_PAGES_RJ):
                chunks.append(chunk)
        return chunks
    except Exception:
        for chunk in chunks:
            cleanup_chunk(chunk)
        raise


def _rj_extrair_chunk(args) -> tuple:
    idx, chunk_path, offset, total_pg, n_total, client, *extras = args
    model_extracao = extras[0] if len(extras) > 0 else GEMINI_MODEL_EXTRACAO
    nome_origem = extras[1] if len(extras) > 1 else Path(chunk_path).name
    paginas_digitalizadas = extras[2] if len(extras) > 2 else []
    pg_ini = offset + 1
    try:
        with fitz.open(chunk_path) as _chunk_doc:
            pg_fim = offset + len(_chunk_doc)
    except Exception:
        pg_fim = min(offset + CHUNK_MAX_PAGES_RJ, total_pg) if total_pg else "?"

    textos_pag = []
    imagens_pag = []
    paginas_omitidas = []
    total_img_bytes = 0
    MAX_IMG_BYTES = 19 * 1024 * 1024

    try:
        doc = fitz.open(chunk_path)
        for i, page in enumerate(doc):
            txt = page.get_text().strip()
            pg = offset + i + 1
            alfa = sum(c.isalpha() for c in txt)
            if len(txt) >= 50 and alfa / len(txt) >= 0.4 and len(txt.split()) >= 30:
                textos_pag.append(f"[p.{pg}]\n{txt}")
            else:
                if total_img_bytes < MAX_IMG_BYTES:
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.2, 1.2))
                    img = pix.tobytes("jpeg", jpg_quality=50)
                    total_img_bytes += len(img)
                    imagens_pag.append((pg, img))
                else:
                    paginas_omitidas.append(pg)
        doc.close()
    except Exception:
        pass

    n_scan = len(imagens_pag)
    if paginas_omitidas:
        raise RuntimeError(
            "Fallback visual excedeu o limite seguro e não enviaria todas as páginas "
            f"digitalizadas ({paginas_omitidas[0]}-{paginas_omitidas[-1]}). "
            "A análise foi interrompida para não gerar relatório incompleto."
        )

    cabecalho = (
        f"[PARTE {idx+1}/{n_total} — paginas {pg_ini}-{pg_fim}]\n"
        "Nao omita nenhum dado mesmo que pareca repetitivo.\n"
        + contexto_fonte_pdf(nome_origem, pg_ini, pg_fim, paginas_digitalizadas)
    )
    prompt_txt = cabecalho + PROMPT_RJ
    if textos_pag:
        prompt_txt += "\n\nTEXTO PESQUISAVEL EXTRAIDO:\n" + "\n\n".join(textos_pag)
    if n_scan:
        prompt_txt += f"\n\n{n_scan} pagina(s) escaneada(s) enviadas como imagem — aplique OCR visual completo."

    all_parts: list = [types.Part(text=prompt_txt)]
    for pg_num, img_bytes in imagens_pag:
        all_parts.append(types.Part(text=f"[Pagina {pg_num} — escaneada]"))
        all_parts.append(types.Part(inline_data=types.Blob(mime_type="image/jpeg", data=img_bytes)))

    def _call():
        contents = [types.Content(role="user", parts=all_parts)]
        resp = client.models.generate_content(
            model=model_extracao,
            contents=contents,
            config=types.GenerateContentConfig(temperature=0, max_output_tokens=65536),
        )
        return resp.text

    resultado = _retry(_call)
    if not resultado or not resultado.strip():
        raise RuntimeError(f"Extração vazia nas páginas {pg_ini}-{pg_fim} de {nome_origem}.")
    nota = f"{n_scan} pag. escaneadas via imagem · {model_extracao}" if n_scan else model_extracao
    return idx, resultado, nota


def _rj_extrair_chunk_fileapi(args) -> tuple:
    """Faz upload do chunk preservado para o File API e extrai o PDF nativamente."""
    idx, chunk_path, offset, total_pg, n_total, client, *extras = args
    model_extracao = extras[0] if len(extras) > 0 else GEMINI_MODEL_EXTRACAO
    nome_origem = extras[1] if len(extras) > 1 else Path(chunk_path).name
    paginas_digitalizadas = extras[2] if len(extras) > 2 else []
    pg_ini = offset + 1
    try:
        with fitz.open(chunk_path) as _chunk_doc:
            pg_fim = offset + len(_chunk_doc)
    except Exception:
        pg_fim = min(offset + CHUNK_MAX_PAGES_RJ, total_pg) if total_pg else "?"
    source_for_upload = chunk_path
    ascii_tmp = None
    try:
        source_for_upload.encode("ascii")
        upload_path = source_for_upload
    except (UnicodeEncodeError, AttributeError):
        import shutil as _shutil
        t = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        t.close()
        _shutil.copy2(source_for_upload, t.name)
        upload_path = ascii_tmp = t.name

    try:
        arq = client.files.upload(file=upload_path)
    finally:
        if ascii_tmp:
            try: os.remove(ascii_tmp)
            except OSError: pass
    total_wait, wait_time = 0, 1
    while total_wait < 120:
        _st = getattr(arq, "state", None)
        state_name = getattr(_st, "name", None) or str(_st or "")
        if state_name in ("ACTIVE", "FAILED"):
            break
        time.sleep(wait_time)
        total_wait += wait_time
        arq = client.files.get(name=arq.name)
        wait_time = min(wait_time + 1, 4)

    _st = getattr(arq, "state", None)
    state_name = getattr(_st, "name", None) or str(_st or "")
    try:
        if state_name == "FAILED":
            raise RuntimeError(f"File API: upload do chunk {idx+1} falhou (FAILED)")

        time.sleep(0.5)
        prompt = (
            f"[PARTE {idx+1}/{n_total} — paginas {pg_ini}-{pg_fim}]\n"
            "Nao omita nenhum dado mesmo que pareca repetitivo.\n\n"
            + contexto_fonte_pdf(nome_origem, pg_ini, pg_fim, paginas_digitalizadas)
            + "\n"
            + PROMPT_RJ
        )
        mime = getattr(arq, "mime_type", None) or "application/pdf"
        contents = [types.Content(role="user", parts=[
            types.Part(text=prompt),
            types.Part(file_data=types.FileData(file_uri=arq.uri, mime_type=mime)),
        ])]

        def _call():
            return client.models.generate_content(
                model=model_extracao,
                contents=contents,
                config=types.GenerateContentConfig(temperature=0, max_output_tokens=65536),
            ).text

        resultado = _retry(_call, tentativas=2, espera_base=5)
        if not resultado or not resultado.strip():
            raise RuntimeError(f"Extração vazia nas páginas {pg_ini}-{pg_fim} de {nome_origem}.")
    finally:
        try: client.files.delete(name=arq.name)
        except Exception: pass

    modo = "OCR visual forte" if paginas_digitalizadas else "texto pesquisável"
    return idx, resultado, f"File API · {modo} · {model_extracao}"


def _rj_obter_cache(client, model_cons: str) -> Optional[str]:
    global _rj_cache_nome
    with _rj_cache_lock:
        if model_cons in _rj_cache_nome:
            # Valida se o cache ainda existe na API (TTL de 1h pode ter expirado)
            try:
                client.caches.get(name=_rj_cache_nome[model_cons])
                return _rj_cache_nome[model_cons]
            except Exception:
                _rj_cache_nome.pop(model_cons, None)
        try:
            cache = client.caches.create(
                model=model_cons,
                config=types.CreateCachedContentConfig(
                    display_name="invista_rj_template_v1",
                    system_instruction=SYSTEM_PROMPT_RJ,
                    contents=[types.Content(
                        role="user",
                        parts=[types.Part(text=f"TEMPLATE DO RELATORIO:\n\n{REPORT_TEMPLATE_RJ}")]
                    )],
                    ttl="3600s",
                )
            )
            _rj_cache_nome[model_cons] = cache.name
            return cache.name
        except Exception:
            return None


def _rj_merge_textos(extracoes: list, nomes_arquivos: list = None) -> str:
    """Junta os textos extraidos de cada chunk. Se houver mais de 1 PDF de origem
    (nomes_arquivos com mais de um nome distinto), marca de qual arquivo veio cada parte —
    sem isso, o modelo de consolidacao nao tem como indicar "de qual pdf" numa referencia
    quando mais de um processo foi enviado junto."""
    multi_arquivo = bool(nomes_arquivos) and len(set(nomes_arquivos)) > 1
    partes = []
    for i, txt in enumerate(extracoes):
        if txt:
            cabecalho = f"PARTE {i+1}/{len(extracoes)}"
            if multi_arquivo:
                cabecalho += f" — arquivo: {nomes_arquivos[i]}"
            partes.append(f"{'='*60}\n{cabecalho}\n{'='*60}\n{txt}")
    return "\n\n".join(partes)


def _rj_dividir_fonte_em_lotes(texto: str, limite: int = LIMITE_LOTE_FIDELIDADE_RJ) -> list[str]:
    """Agrupa partes completas sem cortar uma extração no meio."""
    if len(texto or "") <= limite:
        return [texto] if texto else []
    partes = re.split(r"(?=^={40,}\nPARTE\s+\d+)", texto, flags=re.MULTILINE)
    partes = [p for p in partes if p and p.strip()]
    lotes, atual = [], ""
    for parte in partes:
        if atual and len(atual) + len(parte) + 2 > limite:
            lotes.append(atual)
            atual = parte
        else:
            atual = (atual + "\n\n" + parte) if atual else parte
    if atual:
        lotes.append(atual)
    return lotes


def _rj_preservar_fonte_longa(clients: list, texto: str, model_cons: str,
                              log: list = None, rotulo: str = "RJ") -> str:
    """Substitui truncamento por inventários completos de todos os lotes.

    Cada lote vira um inventário de evidências, não um relatório final. Assim o
    consolidator recebe começo, meio e fim da fonte em vez de só um prefixo.
    """
    if len(texto or "") <= LIMITE_CONSOLIDACAO_RJ:
        return texto
    lotes = _rj_dividir_fonte_em_lotes(texto)
    inventarios = []
    for indice, lote in enumerate(lotes, 1):
        prompt = (
            REGRAS_CONSOLIDACAO_PROCESSUAL
            + "\n\nVocê está na etapa intermediária de cobertura documental. Produza um INVENTÁRIO "
              "FACTUAL COMPLETO, não o relatório final, de todas as evidências do lote abaixo. "
              "Preserve cada data, valor, índice, cláusula, aditamento, parte, andamento e referência "
              "exatamente como consta. Não elimine fatos por parecerem repetidos ou pouco relevantes; "
              "a deduplicação ocorrerá depois. Mantenha os blocos FONTE_PDF/PDF_PAGE/"
              "IDENTIFICADOR_TRIBUNAL.\n\n"
            + f"LOTE {indice}/{len(lotes)} — {rotulo}:\n{lote}"
        )

        def _gerar(client, _indice):
            return _retry(lambda: client.models.generate_content(
                model=model_cons,
                contents=[prompt],
                config=types.GenerateContentConfig(temperature=0, max_output_tokens=65536),
            ).text)

        inventario = _executar_com_failover_gemini(clients, _gerar)
        inventarios.append(f"INVENTÁRIO DOCUMENTAL {indice}/{len(lotes)}\n{inventario}")
        if log is not None:
            log.append(f"   Cobertura intermediária {indice}/{len(lotes)} concluída ({rotulo}).")
    return "\n\n".join(inventarios)


TEMPLATE_RESUMIDO_RJ = """\
MODO RESUMO PROCESSUAL — siga rigorosamente este formato curto:

A. Recuperação Judicial nº [número completo] - [Vara] - [Tribunal]

• Partes:
   ∘ Recuperando(s): [Nome] — [CPF/CNPJ]
   ∘ Administrador Judicial: [Nome] — [OAB]

• Data de distribuição: DD/MM/AAAA

• Data de deferimento do processamento: DD/MM/AAAA

• Principais andamentos: (lista cronológica completa do mais antigo ao mais recente, com data, descrição objetiva e referência Mov./fls.)
   ▪ DD/MM/AAAA — [Descrição objetiva] (Mov. X | fls. XX/XX)
   ▪ DD/MM/AAAA — [Próximo ato] (Mov. X | fls. XX/XX)
   ...

REGRAS RÍGIDAS:
- NÃO gere análise de créditos, PRJ, AGC, RMA, QGC, stay period detalhado, endividamento fiscal, garantias, contratos, índices ou qualquer outra seção.
- NÃO inclua nenhum conteúdo além das 5 seções acima.
- Use os marcadores • ∘ ▪ conforme o exemplo.
- Mantenha objetividade nos andamentos.
"""


def _rj_consolidar_secao_a(client, texto_merged: str, instrucoes: str, cache_name, model_cons: str, versao_resumida: bool = False) -> str:
    if versao_resumida:
        instrucoes_extras = f"\n\nINSTRUCOES ADICIONAIS: {instrucoes.strip()}" if instrucoes.strip() else ""
        n_partes = texto_merged.count("PARTE ") if texto_merged else 0
        prompt_full = (
            SYSTEM_PROMPT_RJ + "\n\n"
            + REGRAS_CONSOLIDACAO_PROCESSUAL + "\n\n"
            + "Você está em MODO RESUMO. Ignore qualquer template detalhado e siga ESTRITAMENTE o formato abaixo.\n\n"
            + TEMPLATE_RESUMIDO_RJ
            + instrucoes_extras
            + f"\n\nA seguir estão {n_partes} PARTE(S) de informações extraídas do processo.\n"
            + "INSTRUÇÃO OBRIGATÓRIA DE LEITURA:\n"
            + "1. Leia CADA PARTE individualmente, do começo ao fim, sem pular nenhuma.\n"
            + "2. Extraia todos os andamentos de CADA UMA das partes (PARTE 1, PARTE 2, ..., PARTE N).\n"
            + "3. Ao final, mescle todos os andamentos em UMA única lista cronológica global.\n"
            + "4. Omitir andamentos de qualquer PARTE é ERRO GRAVE — antes de finalizar, confira se a quantidade de andamentos extraídos é compatível com a totalidade das partes.\n\n"
            + texto_merged
        )
        def _fn():
            return client.models.generate_content(
                model=model_cons,
                contents=[prompt_full],
                config=types.GenerateContentConfig(temperature=0, max_output_tokens=65536),
            ).text
        return _retry(_fn)

    n_partes_cons = texto_merged.count("PARTE ") if texto_merged else 0
    _aviso_multiplos_rj = (
        f"ATENCAO: foram carregados fragmentos de {n_partes_cons} parte(s). "
        "Verifique se compartilham o mesmo numero de processo de Recuperacao Judicial — "
        "se sim, trate como UM UNICO PROCESSO CONTINUO. "
        "Se forem processos distintos, gere secoes separadas para cada um.\n\n"
    ) if n_partes_cons > 1 else ""
    prompt = (
        REGRAS_CONSOLIDACAO_PROCESSUAL
        + "\n\nCom base nas informacoes extraidas abaixo (de multiplas partes do processo), "
        "gere APENAS a Secao A do relatorio de Recuperacao Judicial, "
        "seguindo RIGOROSAMENTE o template fornecido.\n\n"
        + _aviso_multiplos_rj
        + "Consolide informacoes duplicadas — priorize a mais completa e recente.\n\n"
        + f"INFORMACOES EXTRAIDAS:\n{texto_merged}\n\n"
    )
    if instrucoes.strip():
        prompt += f"INSTRUCOES ADICIONAIS: {instrucoes.strip()}\n\n"
    prompt += (
        "Gere APENAS a Secao A — de '1. VISAO JURIDICA' ate antes de qualquer "
        "secao B.1 ou B.2. Use todos os marcadores e sub-niveis do template."
    )
    if cache_name:
        contents = [types.Content(role="user", parts=[types.Part(text=prompt)])]
        config = types.GenerateContentConfig(
            cached_content=cache_name, temperature=0, max_output_tokens=65536
        )
        def _fn():
            return client.models.generate_content(
                model=model_cons, contents=contents, config=config
            ).text
        try:
            return _retry(_fn)
        except Exception as e:
            msg = str(e)
            if "PERMISSION_DENIED" in msg or "CachedContent not found" in msg or "403" in msg:
                with _rj_cache_lock:
                    _rj_cache_nome.pop(model_cons, None)
                conteudo = SYSTEM_PROMPT_RJ + "\n\n" + REPORT_TEMPLATE_RJ + "\n\n" + prompt
                def _fn_fallback():
                    return client.models.generate_content(
                        model=model_cons,
                        contents=[conteudo],
                        config=types.GenerateContentConfig(temperature=0, max_output_tokens=65536),
                    ).text
                return _retry(_fn_fallback)
            raise
    else:
        conteudo = SYSTEM_PROMPT_RJ + "\n\n" + REPORT_TEMPLATE_RJ + "\n\n" + prompt
        def _fn():
            return client.models.generate_content(
                model=model_cons,
                contents=[conteudo],
                config=types.GenerateContentConfig(temperature=0, max_output_tokens=65536),
            ).text
        return _retry(_fn)


def _rj_consolidar_secao_a_com_failover(
    clients: list,
    texto_merged: str,
    instrucoes: str,
    cache_name,
    model_cons: str,
    versao_resumida: bool = False,
    log: list = None,
) -> str:
    def _ao_falhar(indice, proximo, _exc):
        if log is not None:
            log.append(
                f"   Credencial Gemini {indice + 1} recusada ou indisponível; "
                f"tentando a credencial {proximo + 1}."
            )

    return _executar_com_failover_gemini(
        clients,
        lambda client, indice: _rj_consolidar_secao_a(
            client,
            texto_merged,
            instrucoes,
            cache_name if indice == 0 else None,
            model_cons,
            versao_resumida,
        ),
        ao_falhar=_ao_falhar,
    )


def _rj_processar_relacionados(pdf_paths: list, clients: list, instrucoes: str, cache_name, model_cons: str,
                                standalone: bool = False, job_id: str = None, runtime_job=None):
    """Extrai e consolida os 'processos relacionados'. GERADOR — durante a extração, produz
    ("progress", linha_de_log) a cada trecho concluído (qual processo, X/N trechos); ao final,
    produz exatamente um ("done", (secao_texto, texto_bruto, avisos)).

    texto_bruto é o texto extraído página-a-página (antes da consolidação) — quem chama deve
    somá-lo à fonte usada pelos checklists/Excel de credores, para que dados de execuções que só
    aparecem nos relacionados (não repetidos no PDF principal de RJ) não se percam.

    standalone=True é usado quando NÃO há PDF principal de RJ: os processos abaixo são tratados
    como processos PRINCIPAIS (não subordinados a nenhuma RJ), reaproveitando o mesmo template de
    'Processos Relacionados' já embutido em SYSTEM_PROMPT_RJ/REPORT_TEMPLATE_RJ.

    job_id, se informado, ativa o cache em disco (storage persistente) dos trechos já extraídos
    e da seção final já consolidada — uma tentativa anterior interrompida no meio não precisa
    reprocessar o que já foi obtido do Gemini.
    """
    avisos: list = []
    manifest = rj_cache.carregar_manifest(job_id) if job_id else None
    secao_cache_nome = "secao_rel_standalone" if standalone else "secao_rel"
    texto_bruto_cache_nome = "texto_bruto_rel_standalone" if standalone else "texto_bruto_rel"
    try:
        if job_id:
            secao_existente = rj_cache.carregar_secao(job_id, secao_cache_nome)
            texto_bruto_existente = rj_cache.carregar_secao(job_id, texto_bruto_cache_nome)
            if secao_existente and texto_bruto_existente:
                yield "progress", "   Processos relacionados ja haviam sido processados nesta analise (reaproveitado do cache)."
                yield "done", (secao_existente, texto_bruto_existente, avisos)
                return

        todos_chunks_rel = []
        hashes_arquivos_rel = {}
        scans_por_arquivo_rel = {}
        for path in pdf_paths:
            hashes_arquivos_rel[path] = file_sha256(path)
            scans_por_arquivo_rel[path] = _paginas_digitalizadas_pdf(path)
            for chunk in _rj_dividir_pdf(path):
                todos_chunks_rel.append((
                    chunk.path, chunk.start, chunk.total_pages, path,
                    chunk.preparation_note,
                ))

        n_rel = len(todos_chunks_rel)
        if n_rel == 0:
            yield "done", ("", "", avisos)
            return

        nomes_arquivos = sorted({Path(p).name for _, _, _, p, _ in todos_chunks_rel})
        yield "progress", (
            f"   {n_rel} trecho(s) a extrair de {len(pdf_paths)} processo(s): "
            + ", ".join(nomes_arquivos)
        )

        resultados: dict = {}
        erros_extracao: list[str] = []

        def _worker_rel(i):
            cp, offset, total_pg, original, preparation_note = todos_chunks_rel[i]
            try:
                with fitz.open(cp) as _chunk_doc:
                    chunk_end = offset + len(_chunk_doc)
            except Exception:
                chunk_end = min(offset + CHUNK_MAX_PAGES_RJ, total_pg)
            source_hash = hashes_arquivos_rel[original]
            paginas_scan = [p for p in scans_por_arquivo_rel[original] if offset < p <= chunk_end]
            model_extracao = GEMINI_MODEL_OCR if paginas_scan else GEMINI_MODEL_EXTRACAO
            cache_namespace = "rj_rel_standalone" if standalone else "rj_relacionados"
            cached = load_chunk_cache(
                source_hash, offset, chunk_end, model_extracao, cache_namespace
            )
            if cached is not None:
                return i, cached, "cache por arquivo/páginas"
            if job_id and manifest and manifest["chunks_status_relacionados"].get(str(i)) == "ok":
                return i, rj_cache.carregar_chunk(job_id, "relacionados", i), "cache"
            slot = runtime_job.worker_slot() if runtime_job else nullcontext()
            with slot:
                record_status(
                    runtime_job, "extraindo_chunk", arquivo=Path(original).name,
                    paginas=f"{offset + 1}-{chunk_end}", chunk=i + 1,
                )
                trocas = []

                def _extrair_com(client, _indice):
                    try:
                        return _retry(
                            lambda: _rj_extrair_chunk_fileapi(
                                (i, cp, offset, total_pg, n_rel, client, model_extracao,
                                 Path(original).name, paginas_scan)
                            ),
                            tentativas=2, espera_base=15,
                        )
                    except Exception as e:
                        msg_e = str(e)
                        if any(c in msg_e for c in ["400", "INVALID_ARGUMENT", "403", "PERMISSION_DENIED"]):
                            ri, res, nota = _rj_extrair_chunk(
                                (i, cp, offset, total_pg, n_rel, client, model_extracao,
                                 Path(original).name, paginas_scan)
                            )
                            return ri, res, (nota + " | " if nota else "") + "fallback inline"
                        raise

                result = _executar_com_failover_gemini(
                    clients,
                    _extrair_com,
                    indice_inicial=i % len(clients),
                    ao_falhar=lambda atual, proximo, _exc: trocas.append(
                        f"credencial Gemini {atual + 1}→{proximo + 1}"
                    ),
                )
                result = (
                    result[0], result[1],
                    " | ".join(filter(None, [result[2], preparation_note, *trocas])),
                )
            save_chunk_cache(
                source_hash, offset, chunk_end, model_extracao, cache_namespace, result[1]
            )
            return result

        concluidos = 0
        with concurrent.futures.ThreadPoolExecutor(max_workers=6) as ex:
            futures = {ex.submit(_worker_rel, i): i for i in range(n_rel)}
            for future in concurrent.futures.as_completed(futures):
                i_f = futures[future]
                nome_arq = Path(todos_chunks_rel[i_f][3]).name
                concluidos += 1
                try:
                    idx, res, nota = future.result()
                    resultados[idx] = res
                    if job_id and manifest and not nota.startswith("cache"):
                        try:
                            rj_cache.salvar_chunk(job_id, "relacionados", idx, res)
                            manifest["chunks_status_relacionados"][str(idx)] = "ok"
                            rj_cache.salvar_manifest(job_id, manifest)
                        except Exception:
                            pass
                    acao = "reaproveitado do cache" if nota.startswith("cache") else "extraido"
                    yield "progress", (
                        f"   {nome_arq} — trecho {i_f+1}/{n_rel} {acao}"
                        + (f" [{nota}]" if nota and not nota.startswith("cache") else "")
                        + f" | {concluidos}/{n_rel} prontos"
                    )
                except Exception as e:
                    cp_erro, offset_erro, total_erro, _original_erro, _ = todos_chunks_rel[i_f]
                    try:
                        with fitz.open(cp_erro) as doc_erro:
                            fim_erro = offset_erro + len(doc_erro)
                    except Exception:
                        fim_erro = min(offset_erro + CHUNK_MAX_PAGES_RJ, total_erro)
                    erros_extracao.append(
                        f"{nome_arq}, páginas {offset_erro + 1}-{fim_erro}: {e}"
                    )
                    yield "progress", f"   {nome_arq} — trecho {i_f+1}/{n_rel}: erro ({e}) | {concluidos}/{n_rel} concluidos"
                finally:
                    # Libera o chunk do disco assim que termina, em vez de so no final de
                    # todos os trechos — reduz o pico de uso de disco com PDFs grandes.
                    cp_done, _, _, orig_done, _preparation_note = todos_chunks_rel[i_f]
                    if cp_done != orig_done:
                        try: os.remove(cp_done)
                        except Exception: pass

        for cp, _off, _tot, orig, _preparation_note in todos_chunks_rel:
            if cp != orig:
                try: os.remove(cp)
                except: pass

        if erros_extracao:
            raise RuntimeError(
                "Extração incompleta dos processos relacionados; nenhum relatório parcial foi gerado. "
                + " | ".join(erros_extracao)
            )

        texto_relacionados = ""
        for i, (_cp, _offset, _total_pg, orig_path, _preparation_note) in enumerate(todos_chunks_rel):
            res = resultados.get(i, "")
            if res:
                texto_relacionados += f"\n--- {Path(orig_path).name} ---\n{res}"

        if not texto_relacionados.strip():
            yield "done", ("", "", avisos)
            return

        yield "progress", f"   Consolidando {n_rel} trecho(s) extraido(s) dos processos relacionados..."

        texto_relacionados_prompt = texto_relacionados
        if len(texto_relacionados_prompt) > LIMITE_CONSOLIDACAO_RJ:
            yield "progress", (
                f"   Fonte dos relacionados tem {len(texto_relacionados_prompt):,} caracteres — "
                "gerando inventários intermediários de todos os lotes, sem descartar o final."
            )
            texto_relacionados_prompt = _rj_preservar_fonte_longa(
                clients, texto_relacionados_prompt, model_cons,
                log=avisos, rotulo="processos relacionados",
            )

        instrucao_anti_omissao = (
            REGRAS_CONSOLIDACAO_PROCESSUAL
            + "\n\n"
            "INSTRUCAO OBRIGATORIA DE LEITURA:\n"
            "1. Leia cada processo/parte individualmente, do comeco ao fim, sem pular nenhum.\n"
            "2. Nao omita nenhuma execucao, valor, garantia, parte ou andamento relevante de qualquer processo.\n"
            f"3. Antes de finalizar, confira se voce cobriu os {len(pdf_paths)} processo(s) fornecidos abaixo.\n\n"
        )

        if standalone:
            prompt = (
                "Os processos abaixo NAO tem nenhuma Recuperacao Judicial principal associada — cada um "
                "deve ser tratado como um processo PRINCIPAL, nao subordinado a nenhum outro. Gere a Secao "
                "1 (VISAO JURIDICA) do relatorio, com uma subsecao por processo principal (A., B., C., ...) "
                "e por incidente/recurso relacionado a ele (A.1., A.2., ...), seguindo RIGOROSAMENTE o "
                "template de 'Processos Relacionados' fornecido (mesma estrutura, marcadores e nivel de "
                "detalhe usados para execucoes, embargos, IDPJs, acoes paulianas etc.).\n\n"
                + instrucao_anti_omissao
                + f"PROCESSOS ({len(pdf_paths)} arquivo(s)):\n{texto_relacionados_prompt}\n\n"
            )
        else:
            prompt = (
                "Com base nos processos relacionados abaixo, "
                "gere secoes B.1, B.2, etc., conforme o template.\n\n"
                "REGRAS:\n"
                "- Recursos e Embargos a Execucao: integre na Secao A do processo principal\n"
                "- IDPJs, Paulianas, Embargos de Terceiro: crie secoes B proprias\n"
                "Use: 'B.1. Incidente de Desconsideracao da Personalidade n. XXX'\n\n"
                + instrucao_anti_omissao
                + f"PROCESSOS RELACIONADOS ({len(pdf_paths)} arquivo(s)):\n{texto_relacionados_prompt}\n\n"
            )
        if instrucoes.strip():
            prompt += f"INSTRUCOES: {instrucoes.strip()}\n\n"
        prompt += (
            "Gere o relatorio completo (Secao 1 apenas), sem inventar informacao."
            if standalone else
            "Gere APENAS as secoes de processos relacionados (B.1, B.2, etc.)."
        )

        def _consolidar_com(client, indice):
            if cache_name and indice == 0:
                contents = [types.Content(role="user", parts=[types.Part(text=prompt)])]
                config = types.GenerateContentConfig(
                    cached_content=cache_name, temperature=0, max_output_tokens=65536
                )
                try:
                    return _retry(lambda: client.models.generate_content(
                        model=model_cons, contents=contents, config=config
                    ).text)
                except Exception as e:
                    msg = str(e)
                    if "PERMISSION_DENIED" in msg or "CachedContent not found" in msg or "403" in msg:
                        with _rj_cache_lock:
                            _rj_cache_nome.pop(model_cons, None)
                    else:
                        raise
            conteudo = SYSTEM_PROMPT_RJ + "\n\n" + REPORT_TEMPLATE_RJ + "\n\n" + prompt
            return _retry(lambda: client.models.generate_content(
                model=model_cons,
                contents=[conteudo],
                config=types.GenerateContentConfig(temperature=0, max_output_tokens=65536),
            ).text)

        secao = _executar_com_failover_gemini(
            clients,
            _consolidar_com,
            ao_falhar=lambda atual, proximo, _exc: avisos.append(
                f"Credencial Gemini {atual + 1} recusada ou indisponível; "
                f"usada a credencial {proximo + 1}."
            ),
        )
        secao = normalizar_referencias_relatorio(secao, len(nomes_arquivos) > 1)

        if job_id:
            try:
                rj_cache.salvar_secao(job_id, secao_cache_nome, secao)
                rj_cache.salvar_secao(job_id, texto_bruto_cache_nome, texto_relacionados)
            except Exception:
                pass

        yield "done", (secao, texto_relacionados, avisos)
    except Exception:
        raise


def _erro_completo_relatorio(exc: Exception) -> str:
    """Texto de erro pra aba Relatorio: traceback completo, nao so str(exc) — o Gradio (sem
    show_error=True) nao mostra nenhum detalhe do erro, entao a unica forma da usuaria ver
    o motivo real de um "Erro" vermelho e a propria funcao escrever isso num output visivel."""
    return "❌ ERRO — a análise foi interrompida.\n\n" + traceback.format_exc()


def rj_analisar(pdf_files, pdf_relacionados, instrucoes: str = "", versao_resumida: bool = False):
    runtime_job = ANALYSIS_MANAGER.create("rj")
    try:
        while True:
            active, position = runtime_job.try_activate()
            if active:
                break
            yield queue_message(position), "", "", ""
            runtime_job.wait_for_change(2.0)
        record_status(runtime_job, "iniciando")
        yield from _rj_analisar_impl(
            pdf_files, pdf_relacionados, instrucoes, versao_resumida,
            runtime_job=runtime_job,
        )
    except Exception as exc:
        yield "Erro inesperado — veja o traceback completo na aba Relatório.", _erro_completo_relatorio(exc), "", ""
    finally:
        record_status(runtime_job, "finalizada")
        runtime_job.close()


def _rj_analisar_somente_relacionados(pdf_relacionados, instrucoes: str, runtime_job=None):
    try:
        yield from _rj_analisar_somente_relacionados_impl(pdf_relacionados, instrucoes, runtime_job)
    except Exception as exc:
        yield "Erro inesperado — veja o traceback completo na aba Relatório.", _erro_completo_relatorio(exc), "", ""


def _rj_analisar_somente_relacionados_impl(pdf_relacionados, instrucoes: str, runtime_job=None):
    """Analisa só 'processos relacionados' (execuções/ações avulsas), sem PDF principal de RJ.

    Gera um relatório objetivo por processo (reaproveitando o template de 'Processos
    Relacionados') e devolve o texto bruto consolidado como fonte para os checklists/Excel de
    credores — permite gerar o Checklist de Créditos sem exigir upload de uma RJ.
    """
    yield "Iniciando analise dos processos relacionados (sem RJ principal)...", "", "", ""

    try:
        clients = _get_gemini_clients()
    except Exception as e:
        yield f"Erro de configuracao: {e}", "", "", ""
        return

    model_cons = GEMINI_MODEL_RELATORIO
    pdf_paths_rel = [f.name if hasattr(f, "name") else str(f) for f in pdf_relacionados]

    log: list = []
    pdf_paths_rel = _filtrar_arquivos_existentes(pdf_paths_rel, log)
    if not pdf_paths_rel:
        log.append("\nNenhum dos arquivos enviados pôde ser lido — tente reenviar.")
        yield "\n".join(log), "", "", ""
        return

    try:
        rj_cache.limpar_jobs_antigos()
    except Exception:
        pass
    try:
        job_id = rj_cache.calcular_job_id(
            [], pdf_paths_rel, instrucoes, False, GEMINI_MODEL_EXTRACAO,
            GEMINI_MODEL_RELATORIO, GEMINI_MODEL_OCR
        )
    except Exception:
        job_id = None
    if job_id and not rj_cache.carregar_manifest(job_id):
        rj_cache.salvar_manifest(job_id, rj_cache.novo_manifest(job_id))

    log.append(f"Processo(s) relacionado(s) recebido(s): {len(pdf_paths_rel)}")
    for p in pdf_paths_rel:
        mb = Path(p).stat().st_size / 1_048_576
        log.append(f"   · {Path(p).name} ({mb:.1f} MB)")
    log.append(f"Modelo de consolidacao: {model_cons}")
    yield "\n".join(log), "", "", ""

    t_inicio = time.time()
    log.append("\nExtraindo processos relacionados (paralelo · File API)...")
    yield "\n".join(log), "", "", ""

    try:
        secao_rel, texto_bruto, avisos = "", "", []
        for kind, payload in _rj_processar_relacionados(
            pdf_paths_rel, clients, instrucoes, cache_name=None, model_cons=model_cons,
            standalone=True, job_id=job_id, runtime_job=runtime_job,
        ):
            if kind == "progress":
                log.append(payload)
                yield "\n".join(log), "", "", ""
            else:
                secao_rel, texto_bruto, avisos = payload

        for aviso in avisos:
            log.append(f"   {aviso}")
        if avisos:
            yield "\n".join(log), "", "", ""

        if not texto_bruto.strip():
            log.append("\nNenhuma informacao relevante extraida dos processos relacionados.")
            yield "\n".join(log), "", "", ""
            return

        relatorio = secao_rel or "Nenhuma informacao relevante extraida dos processos relacionados."
        t_total = int(time.time() - t_inicio)
        log.append(f"\nAnalise concluida em {t_total//60}min{t_total%60:02d}s | {len(relatorio):,} chars")
        if job_id:
            manifest_final = rj_cache.carregar_manifest(job_id)
            if manifest_final:
                manifest_final["concluido"] = True
                try:
                    rj_cache.salvar_manifest(job_id, manifest_final)
                except Exception:
                    pass
        yield "\n".join(log), relatorio, relatorio, texto_bruto

    except Exception as exc:
        log.append("\nErro — veja o traceback completo na aba Relatório.")
        yield "\n".join(log), _erro_completo_relatorio(exc), "", ""


def _extrair_texto_docx(path: str) -> str:
    """Extrai todo o texto (parágrafos + tabelas) de um .docx — usado para reaproveitar um
    relatório de RJ já gerado anteriormente, sem precisar reprocessar o PDF original."""
    doc = Document(path)
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    partes.append(cell.text.strip())
    return "\n".join(partes)


def _rj_analisar_com_relatorio_existente(docx_paths: list, pdf_relacionados, instrucoes: str,
                                          runtime_job=None):
    try:
        yield from _rj_analisar_com_relatorio_existente_impl(
            docx_paths, pdf_relacionados, instrucoes, runtime_job
        )
    except Exception as exc:
        yield "Erro inesperado — veja o traceback completo na aba Relatório.", _erro_completo_relatorio(exc), "", ""


def _rj_analisar_com_relatorio_existente_impl(docx_paths: list, pdf_relacionados, instrucoes: str,
                                               runtime_job=None):
    """Reaproveita um relatório de RJ (.docx) já gerado numa análise anterior, em vez de
    reprocessar o PDF original — só os processos relacionados (se houver) são extraídos de
    verdade. Permite gerar o Checklist de Créditos com processos relacionados novos sem
    esperar a extração inteira da RJ de novo."""
    yield "Lendo relatório de RJ já gerado (.docx)...", "", "", ""

    try:
        relatorio_existente = "\n\n".join(_extrair_texto_docx(p) for p in docx_paths)
    except Exception as e:
        yield f"Erro ao ler o(s) .docx enviado(s): {e}", "", "", ""
        return
    if not relatorio_existente.strip():
        yield "Não foi possível extrair texto do(s) .docx enviado(s).", "", "", ""
        return

    log = [f"Relatório de RJ já gerado carregado: {len(relatorio_existente):,} caracteres ({len(docx_paths)} arquivo(s))."]
    yield "\n".join(log), "", "", ""

    if not pdf_relacionados:
        log.append("\nNenhum processo relacionado novo enviado — usando só o relatório existente.")
        yield "\n".join(log), relatorio_existente, relatorio_existente, relatorio_existente
        return

    try:
        clients = _get_gemini_clients()
    except Exception as e:
        yield f"Erro de configuracao: {e}", "", "", ""
        return

    model_cons = GEMINI_MODEL_RELATORIO
    pdf_paths_rel = [f.name if hasattr(f, "name") else str(f) for f in pdf_relacionados]
    pdf_paths_rel = _filtrar_arquivos_existentes(pdf_paths_rel, log)
    if not pdf_paths_rel:
        log.append("\nNenhum dos processos relacionados novos pôde ser lido — usando só o relatório existente.")
        yield "\n".join(log), relatorio_existente, relatorio_existente, relatorio_existente
        return

    try:
        rj_cache.limpar_jobs_antigos()
    except Exception:
        pass
    try:
        job_id = rj_cache.calcular_job_id(
            docx_paths, pdf_paths_rel, instrucoes, False, GEMINI_MODEL_EXTRACAO,
            GEMINI_MODEL_RELATORIO, GEMINI_MODEL_OCR
        )
    except Exception:
        job_id = None
    if job_id and not rj_cache.carregar_manifest(job_id):
        rj_cache.salvar_manifest(job_id, rj_cache.novo_manifest(job_id))

    t_inicio = time.time()
    log.append(f"\nProcessando {len(pdf_paths_rel)} processo(s) relacionado(s) novo(s)...")
    yield "\n".join(log), "", "", ""

    secao_rel, texto_bruto_rel, avisos = "", "", []
    for kind, payload in _rj_processar_relacionados(
        pdf_paths_rel, clients, instrucoes, cache_name=None, model_cons=model_cons,
        standalone=False, job_id=job_id, runtime_job=runtime_job,
    ):
        if kind == "progress":
            log.append(payload)
            yield "\n".join(log), "", "", ""
        else:
            secao_rel, texto_bruto_rel, avisos = payload

    for aviso in avisos:
        log.append(f"   {aviso}")

    texto_bruto_final = relatorio_existente
    if texto_bruto_rel:
        texto_bruto_final = texto_bruto_final + "\n\n" + texto_bruto_rel

    relatorio_final = relatorio_existente
    if secao_rel and "nenhum" not in secao_rel.lower():
        relatorio_final = relatorio_existente + "\n\n" + secao_rel
        log.append("   Processos relacionados novos analisados e somados ao relatório existente.")
    else:
        log.append("   Nenhuma informação relevante nos processos relacionados novos.")

    if job_id:
        manifest_final = rj_cache.carregar_manifest(job_id)
        if manifest_final:
            manifest_final["concluido"] = True
            try:
                rj_cache.salvar_manifest(job_id, manifest_final)
            except Exception:
                pass

    t_total = int(time.time() - t_inicio)
    log.append(f"\nConcluído em {t_total//60}min{t_total%60:02d}s.")
    yield "\n".join(log), relatorio_final, relatorio_final, texto_bruto_final


def _rj_analisar_impl(pdf_files, pdf_relacionados, instrucoes: str = "", versao_resumida: bool = False,
                      runtime_job=None):
    if not pdf_files and not pdf_relacionados:
        yield "Nenhum arquivo enviado.", "", "", ""
        return

    if pdf_relacionados:
        pdf_paths_rel_check = [f.name if hasattr(f, "name") else str(f) for f in pdf_relacionados]
        docx_indices = {i for i, p in enumerate(pdf_paths_rel_check) if p.lower().endswith(".docx")}
        if docx_indices:
            docx_paths = [pdf_paths_rel_check[i] for i in docx_indices]
            pdf_relacionados = [f for i, f in enumerate(pdf_relacionados) if i not in docx_indices]
            if not pdf_files:
                # Relatório de RJ já gerado (Word) enviado junto na caixa de "Processos
                # relacionados" — pula a extração da RJ inteira, só processa os PDFs novos
                # que vieram junto (se houver).
                yield from _rj_analisar_com_relatorio_existente(
                    docx_paths, pdf_relacionados, instrucoes, runtime_job
                )
                return
            # Combinação não suportada (PDF de RJ + .docx em relacionados): o .docx é
            # ignorado aqui pra não quebrar a extração normal (que só lê PDF).

    if not pdf_files:
        # Sem PDF principal de RJ: analisa só os processos relacionados (execuções/ações
        # avulsas), num fluxo mais leve — sem Seção A (que pressupõe uma RJ principal).
        yield from _rj_analisar_somente_relacionados(pdf_relacionados, instrucoes, runtime_job)
        return

    yield "Iniciando analise de Recuperacao Judicial...", "", "", ""

    try:
        clients = _get_gemini_clients()
    except Exception as e:
        yield f"Erro de configuracao: {e}", "", "", ""
        return

    def _nat_key(p):
        return [int(x) if x.isdigit() else x.lower()
                for x in re.split(r"(\d+)", Path(p).name)]

    pdf_paths = sorted(
        [f.name if hasattr(f, "name") else str(f) for f in pdf_files],
        key=_nat_key
    )
    pdf_paths_rel_all = [f.name if hasattr(f, "name") else str(f) for f in (pdf_relacionados or [])]

    log: list = []
    pdf_paths = _filtrar_arquivos_existentes(pdf_paths, log)
    pdf_paths_rel_all = _filtrar_arquivos_existentes(pdf_paths_rel_all, log)
    if not pdf_paths:
        log.append("\nNenhum dos arquivos de RJ enviados pôde ser lido — tente reenviar.")
        yield "\n".join(log), "", "", ""
        return

    try:
        rj_cache.limpar_jobs_antigos()
    except Exception:
        pass
    try:
        job_id = rj_cache.calcular_job_id(
            pdf_paths, pdf_paths_rel_all, instrucoes, versao_resumida,
            GEMINI_MODEL_EXTRACAO, GEMINI_MODEL_RELATORIO, GEMINI_MODEL_OCR,
        )
    except Exception:
        job_id = None
    manifest = rj_cache.carregar_manifest(job_id) if job_id else None
    if job_id and not manifest:
        manifest = rj_cache.novo_manifest(job_id)
        rj_cache.salvar_manifest(job_id, manifest)

    model_cons = GEMINI_MODEL_RELATORIO

    log.append(f"Arquivo(s) recebido(s): {len(pdf_paths)}")
    for p in pdf_paths:
        mb = Path(p).stat().st_size / 1_048_576
        log.append(f"   · {Path(p).name} ({mb:.1f} MB)")
    if len(pdf_paths) > 1:
        log.append(f"   → {len(pdf_paths)} arquivos — o modelo verificara se sao fragmentos do mesmo processo ou processos distintos")
    log.append(f"Modelo de consolidacao: {model_cons}")
    yield "\n".join(log), "", "", ""

    log.append("\nInspecionando e dividindo PDFs sem recompressão de imagens...")
    record_status(runtime_job, "dividindo_pdfs")
    yield "\n".join(log), "", "", ""

    todos_chunks = []
    hashes_arquivos = {}
    scans_por_arquivo = {}
    for path in pdf_paths:
        nome = Path(path).name
        mb_orig = Path(path).stat().st_size / 1_048_576
        hashes_arquivos[path] = file_sha256(path)
        record_status(runtime_job, "dividindo_pdf", arquivo=nome, tamanho_mb=round(mb_orig, 1))
        chunks = _rj_dividir_pdf(path)
        total_pg = chunks[0].total_pages if chunks else 0
        esc = _paginas_digitalizadas_pdf(path)
        scans_por_arquivo[path] = esc
        pct = int(len(esc) / total_pg * 100) if total_pg else 0
        info_scan = (
            f"{len(esc)} pag. digitalizadas ({pct}%) — extração visual com {GEMINI_MODEL_OCR}"
            if esc else f"todas pesquisáveis — extração com {GEMINI_MODEL_EXTRACAO}"
        )
        log.append(f"   · {nome}: {total_pg} pag. — {info_scan}")
        if len(chunks) > 1:
            log.append(f"     Dividido em {len(chunks)} partes de até {CHUNK_MAX_PAGES_RJ} pag. e 45 MB")
        for chunk in chunks:
            todos_chunks.append((
                chunk.path, chunk.start, chunk.total_pages, path,
                chunk.preparation_note,
            ))
            log.append(
                f"     Páginas {chunk.page_start}-{chunk.page_end}: "
                f"{chunk.preparation_note}"
            )

    n = len(todos_chunks)
    workers_iniciais = ANALYSIS_MANAGER.worker_cap() if runtime_job else 6
    n_rodadas = math.ceil(n / max(1, workers_iniciais))
    tempo_est = n_rodadas * AVG_MIN_POR_CHUNK_RJ + MIN_CONSOLIDACAO_RJ
    log.append(
        f"\nEstimativa: ~{tempo_est} min | {n} chunk(s) · até {workers_iniciais} worker(s) "
        f"nesta análise (6 globais) · File API · {model_cons}"
    )
    yield "\n".join(log), "", "", ""

    t_inicio = time.time()

    texto_merged = ""  # disponivel no finally para o Excel de credores

    try:
        record_status(runtime_job, "extraindo", chunks_total=n, chunks_concluidos=0)
        log.append(f"\nExtraindo {n} chunk(s) (workers compartilhados dinamicamente · File API)...")
        log.append(_barra_progresso(0, n))
        yield "\n".join(log), "", "", ""

        parciais: dict = {}
        erros_chunks: list[str] = []

        def _worker_rj(idx):
            cp, offset, total_pg, original, preparation_note = todos_chunks[idx]
            try:
                with fitz.open(cp) as _chunk_doc:
                    chunk_end = offset + len(_chunk_doc)
            except Exception:
                chunk_end = min(offset + CHUNK_MAX_PAGES_RJ, total_pg)
            source_hash = hashes_arquivos[original]
            paginas_scan = [p for p in scans_por_arquivo[original] if offset < p <= chunk_end]
            model_extracao = GEMINI_MODEL_OCR if paginas_scan else GEMINI_MODEL_EXTRACAO
            cached = load_chunk_cache(
                source_hash, offset, chunk_end, model_extracao, "rj_principal"
            )
            if cached is not None:
                return idx, cached, "cache por arquivo/páginas"
            if job_id and manifest and manifest["chunks_status_principal"].get(str(idx)) == "ok":
                return idx, rj_cache.carregar_chunk(job_id, "principal", idx), "cache"
            slot = runtime_job.worker_slot() if runtime_job else nullcontext()
            with slot:
                record_status(
                    runtime_job, "extraindo_chunk", arquivo=Path(original).name,
                    paginas=f"{offset + 1}-{chunk_end}", chunk=idx + 1,
                )
                trocas = []

                def _extrair_com(client, _indice):
                    try:
                        return _retry(
                            lambda: _rj_extrair_chunk_fileapi(
                                (idx, cp, offset, total_pg, n, client, model_extracao,
                                 Path(original).name, paginas_scan)
                            ),
                            tentativas=2, espera_base=15,
                        )
                    except Exception as e:
                        msg_e = str(e)
                        if any(c in msg_e for c in ["400", "INVALID_ARGUMENT", "403", "PERMISSION_DENIED"]):
                            ri, res, nota = _rj_extrair_chunk(
                                (idx, cp, offset, total_pg, n, client, model_extracao,
                                 Path(original).name, paginas_scan)
                            )
                            return ri, res, (nota + " | " if nota else "") + "fallback inline"
                        raise

                result = _executar_com_failover_gemini(
                    clients,
                    _extrair_com,
                    indice_inicial=idx % len(clients),
                    ao_falhar=lambda atual, proximo, _exc: trocas.append(
                        f"credencial Gemini {atual + 1}→{proximo + 1}"
                    ),
                )
                result = (
                    result[0], result[1],
                    " | ".join(filter(None, [result[2], preparation_note, *trocas])),
                )
            save_chunk_cache(
                source_hash, offset, chunk_end, model_extracao, "rj_principal", result[1]
            )
            return result

        t_extr = time.time()
        with concurrent.futures.ThreadPoolExecutor(max_workers=6) as ex:
            futures = {ex.submit(_worker_rj, i): i for i in range(n)}
            for future in concurrent.futures.as_completed(futures):
                try:
                    idx, res, nota = future.result()
                    parciais[idx] = res
                    if job_id and manifest and not nota.startswith("cache"):
                        try:
                            rj_cache.salvar_chunk(job_id, "principal", idx, res)
                            manifest["chunks_status_principal"][str(idx)] = "ok"
                            rj_cache.salvar_manifest(job_id, manifest)
                        except Exception:
                            pass
                    c = len(parciais)
                    record_status(
                        runtime_job, "extraindo", chunks_total=n, chunks_concluidos=c,
                        arquivo=Path(todos_chunks[idx][3]).name,
                    )
                    log[-1] = _barra_progresso(c, n)
                    current_workers = ANALYSIS_MANAGER.worker_cap() if runtime_job else 6
                    rounds_left = math.ceil((n - c) / max(1, current_workers))
                    est_rest = rounds_left * AVG_MIN_POR_CHUNK_RJ
                    acao = "reaproveitado do cache" if nota.startswith("cache") else "extraido"
                    log.append(
                        f"   Chunk {idx+1}/{n} {acao}"
                        + (f" [{nota}]" if nota and not nota.startswith("cache") else "")
                        + f" | {c}/{n} prontos | ~{est_rest}min restantes"
                    )
                except Exception as e:
                    i_f = futures[future]
                    log.append(f"   Erro no chunk {i_f+1}: {e}")
                    cp_erro, offset_erro, total_erro, original_erro, _ = todos_chunks[i_f]
                    try:
                        with fitz.open(cp_erro) as doc_erro:
                            fim_erro = offset_erro + len(doc_erro)
                    except Exception:
                        fim_erro = min(offset_erro + CHUNK_MAX_PAGES_RJ, total_erro)
                    erros_chunks.append(
                        f"{Path(original_erro).name}, páginas {offset_erro + 1}-{fim_erro}: {e}"
                    )
                finally:
                    # Libera o arquivo do chunk do disco assim que ele termina (sucesso ou
                    # erro) — sem isso, todos os chunks de um PDF grande ficavam ocupando
                    # disco simultaneamente ate o fim da analise inteira (so limpos no
                    # finally geral), o que agrava a pressao de disco com varias analises
                    # rodando ao mesmo tempo.
                    i_f = futures[future]
                    cp_done, _, _, orig_done, _preparation_note = todos_chunks[i_f]
                    if cp_done != orig_done:
                        try: os.remove(cp_done)
                        except Exception: pass
                yield "\n".join(log), "", "", ""

        t_extr_s = int(time.time() - t_extr)
        log.append(f"   Extracao total: {t_extr_s//60}min{t_extr_s%60:02d}s")
        yield "\n".join(log), "", "", ""

        if erros_chunks:
            raise RuntimeError(
                "A extração não cobriu todos os trechos; o relatório parcial foi bloqueado. "
                + " | ".join(erros_chunks)
            )

        record_status(runtime_job, "consolidando", chunks_total=n, chunks_concluidos=len(parciais))
        # Merge
        log.append("\nConsolidando textos extraidos...")
        yield "\n".join(log), "", "", ""
        lista = [parciais.get(i, "") for i in range(n)]
        nomes_por_chunk = [Path(todos_chunks[i][3]).name for i in range(n)]
        texto_merged = _rj_merge_textos(lista, nomes_por_chunk)
        log.append(f"   {len(texto_merged):,} caracteres de informacao extraida")
        yield "\n".join(log), "", "", ""

        # Cache
        log.append(f"\nConfigurando context cache ({model_cons})...")
        yield "\n".join(log), "", "", ""
        cache = _rj_obter_cache(clients[0], model_cons)
        log[-1] = "Cache configurado." if cache else "Cache nao disponivel — usando prompt completo."
        yield "\n".join(log), "", "", ""

        # Secao A
        secoes = []
        log.append(f"\nGerando Secao A — Recuperacao Judicial ({model_cons})...")
        log.append(f"   Aguardando {model_cons}... (pode levar alguns minutos)")
        yield "\n".join(log), "", "", ""

        secao_a_cache = rj_cache.carregar_secao(job_id, "secao_a") if job_id else None
        if secao_a_cache:
            secao_a = secao_a_cache
            log[-1] = "   Secao A reaproveitada do cache."
        else:
            texto_para_secao_a = texto_merged
            if len(texto_para_secao_a) > LIMITE_CONSOLIDACAO_RJ:
                log.append(
                    f"   Fonte extraída tem {len(texto_para_secao_a):,} caracteres — gerando "
                    "inventários intermediários de todos os lotes, sem descartar o final."
                )
                texto_para_secao_a = _rj_preservar_fonte_longa(
                    clients, texto_para_secao_a, model_cons, log=log, rotulo="RJ principal"
                )
            secao_a = _rj_consolidar_secao_a_com_failover(
                clients, texto_para_secao_a, instrucoes, cache,
                model_cons, versao_resumida, log,
            )
            secao_a = normalizar_referencias_relatorio(
                secao_a, len(set(nomes_por_chunk)) > 1
            )
            if log and log[-1].startswith("   Credencial Gemini"):
                log.append("   Secao A gerada.")
            else:
                log[-1] = "   Secao A gerada."
            if job_id:
                try:
                    rj_cache.salvar_secao(job_id, "secao_a", secao_a)
                except Exception:
                    pass
        secoes.append(secao_a)
        yield "\n".join(log), "", "", ""

        # Processos relacionados (já filtrados no início da função — pdf_paths_rel_all)
        if pdf_paths_rel_all:
            pdf_paths_rel = pdf_paths_rel_all
            log.append(f"\nProcessando {len(pdf_paths_rel)} processo(s) relacionado(s)...")
            yield "\n".join(log), "", "", ""
            secao_rel, texto_bruto_rel, avisos_rel = "", "", []
            for kind, payload in _rj_processar_relacionados(
                pdf_paths_rel, clients, instrucoes, cache, model_cons,
                job_id=job_id, runtime_job=runtime_job,
            ):
                if kind == "progress":
                    log.append(payload)
                    yield "\n".join(log), "", "", ""
                else:
                    secao_rel, texto_bruto_rel, avisos_rel = payload
            if texto_bruto_rel:
                # Soma ao texto bruto da RJ principal — sem isso, dados de execucoes que so
                # aparecem nos relacionados (nao repetidos no PDF da RJ) nao chegavam aos
                # checklists/Excel de credores, que priorizam texto_merged como fonte.
                texto_merged = (texto_merged + "\n\n" + texto_bruto_rel) if texto_merged else texto_bruto_rel
            if secao_rel and "nenhum" not in secao_rel.lower():
                log.append("   Processos relacionados analisados.")
                secoes.append(secao_rel)
            else:
                log.append("   Nenhuma informacao relevante nos processos relacionados.")
            for aviso in avisos_rel:
                log.append(f"   {aviso}")
            yield "\n".join(log), "", "", ""

        relatorio = "\n\n".join(s for s in secoes if s)
        t_total = int(time.time() - t_inicio)
        log.append(f"\nAnalise concluida em {t_total//60}min{t_total%60:02d}s | {len(relatorio):,} chars")
        if job_id and manifest:
            manifest["concluido"] = True
            try:
                rj_cache.salvar_manifest(job_id, manifest)
            except Exception:
                pass
        yield "\n".join(log), relatorio, relatorio, texto_merged

    except Exception as exc:
        log.append("\nErro — veja o traceback completo na aba Relatório.")
        yield "\n".join(log), _erro_completo_relatorio(exc), "", ""

    finally:
        for cp, _, _, orig, _preparation_note in todos_chunks:
            if cp != orig:
                try: os.remove(cp)
                except: pass


def rj_gerar_word(relatorio: str):
    if not relatorio.strip():
        return gr.update(value=None, visible=False)
    return gr.update(value=_gerar_docx(relatorio, "Analise de Recuperacao Judicial"), visible=True)


def rj_responder(pergunta: str, relatorio: str):
    try:
        return _executar_com_failover_gemini(
            _get_gemini_clients(),
            lambda client, _indice: _responder_pergunta_generica(
                pergunta, relatorio, client, GEMINI_MODEL_QA
            ),
        )
    except Exception:
        return "❌ Erro:\n\n" + traceback.format_exc()


def rj_gerar_checklist(relatorio: str, texto_bruto: str = ""):
    # O relatório (já consolidado, com referências corretas) é a fonte principal;
    # o texto bruto extraído do processo só complementa o que o relatório não cobrir.
    relatorio = (relatorio or "").strip()
    texto_bruto = (texto_bruto or "").strip()
    if not relatorio and not texto_bruto:
        return gr.update(value=None, visible=False), "Gere uma análise primeiro."
    try:
        path = _executar_com_failover_gemini(
            _get_gemini_clients(),
            lambda client, _indice: gerar_checklist_rj(
                relatorio, texto_bruto, client, GEMINI_MODEL_ESTRUTURADO
            ),
        )
        aviso = _aviso_truncamento(len(relatorio) + len(texto_bruto), LIMITE_TRUNCAMENTO_CHECKLIST)
        return gr.update(value=path, visible=True), "✅ Checklist RJ gerado — clique no arquivo para baixar." + aviso
    except Exception:
        return gr.update(value=None, visible=False), "❌ Erro:\n\n" + traceback.format_exc()


def rj_gerar_checklist_creditos(relatorio: str, texto_bruto: str = "", *campos):
    # Mesma regra do dossiê e do Checklist RJ: relatório consolidado é a fonte
    # principal; o texto bruto só acrescenta fatos que ainda não constem nele.
    fonte = _montar_fonte_rj(relatorio, texto_bruto)
    if not fonte:
        return gr.update(value=None, visible=False), "Gere uma análise primeiro."
    # campos = [nome_1..nome_N, doc_1..doc_N] — pareia e ignora credores sem nome
    meia = len(campos) // 2
    nomes, docs = campos[:meia], campos[meia:]
    creditores = [(str(nomes[i]).strip(), str(docs[i]).strip())
                  for i in range(meia) if nomes[i] and str(nomes[i]).strip()]
    try:
        resultado = _executar_com_failover_gemini(
            _get_gemini_clients(),
            lambda client, _indice: gerar_checklist_creditos(
                fonte, client, GEMINI_MODEL_ESTRUTURADO,
                creditores=creditores or None,
            ),
        )
        paths = resultado if isinstance(resultado, list) else [resultado]
        if creditores:
            msg = f"✅ {len(paths)} checklist(s) de crédito gerado(s) — um arquivo por credor, clique para baixar."
        elif len(paths) > 1:
            msg = f"✅ {len(paths)} credor(es) exequente(s) identificado(s) automaticamente — um checklist por credor, clique para baixar."
        else:
            msg = "✅ Checklist de créditos gerado (crédito identificado automaticamente)."
        msg += _aviso_truncamento(len(fonte), LIMITE_TRUNCAMENTO_CHECKLIST)
        return gr.update(value=paths, visible=True), msg
    except Exception:
        return gr.update(value=None, visible=False), "❌ Erro:\n\n" + traceback.format_exc()


# ── Excel de Credores ─────────────────────────────────────────────────────────

_CLASSE_ORDER = [
    "I - Trabalhista",
    "II - Garantia Real",
    "III - Quirografário",
    "IV - ME/EPP",
    "Extraconcursal",
]

_CLASSE_FILL = {
    "I - Trabalhista":     "FFF2CC",
    "II - Garantia Real":  "D9E1F2",
    "III - Quirografário": "E2EFDA",
    "IV - ME/EPP":         "FCE4D6",
    "Extraconcursal":      "F2F2F2",
}


def _normalizar_classe(raw: str) -> str:
    c = (raw or "").upper()
    if re.search(r'\bCLASSE\s*I\b', c) or any(x in c for x in ("TRABALH", "ACIDENT")):
        return "I - Trabalhista"
    if re.search(r'\bCLASSE\s*II\b', c) or "GARANTIA REAL" in c:
        return "II - Garantia Real"
    if re.search(r'\bCLASSE\s*III\b', c) or "QUIROGRAF" in c:
        return "III - Quirografário"
    if re.search(r'\bCLASSE\s*IV\b', c) or any(x in c for x in ("ME/EPP", "MICROEMPRES", "PEQUENO")):
        return "IV - ME/EPP"
    if "EXTRACONCURSAL" in c:
        return "Extraconcursal"
    return raw.strip() if raw else "Não classificado"


def _extrair_credores_json(texto: str, client, model: str) -> dict:
    prompt = (
        "Você está analisando o texto extraído de um processo de Recuperação Judicial. "
        "Sua tarefa é montar o Quadro Geral de Credores (QGC) mais completo e preciso possível.\n\n"

        "═══ REGRA DE PRIORIDADE DE FONTE ═══\n"
        "1. PRIORIDADE MÁXIMA — Edital do Art. 18 (edital consolidado/definitivo, publicado APÓS a "
        "apreciação das divergências e impugnações pelo AJ e pelo juízo). Nomes comuns: 'Edital Art. 18', "
        "'Edital Consolidado', 'Edital de Habilitação Definitivo', 'Quadro Geral Consolidado'.\n"
        "   → Se encontrar, use-o como lista base.\n"
        "2. PRIORIDADE SECUNDÁRIA — Se NÃO houver Edital Art. 18: use o primeiro edital (Art. 7-A ou "
        "edital inicial do AJ) como lista base e COMPLEMENTE com:\n"
        "   a) Manifestação do AJ sobre divergências/impugnações (ajusta valores, classe ou garantias)\n"
        "   b) Decisões judiciais sobre habilitações/impugnações\n"
        "3. FALLBACK — Se não houver nenhum edital formal, extraia da lista de credores que aparecer "
        "no texto (QGC provisório, planilha, tabela, etc.)\n\n"

        "═══ PARA CADA CREDOR, EXTRAIA ═══\n"
        "- nome: nome/razão social completo\n"
        "- cpf_cnpj: apenas dígitos (sem pontuação), ou null\n"
        "- classe: 'I - Trabalhista' | 'II - Garantia Real' | 'III - Quirografário' | 'IV - ME/EPP' | 'Extraconcursal'\n"
        "- natureza: trabalhista, bancário, fornecedor, tributário, debenturista, CRI, CRA, CCB, etc.\n"
        "- instrumento: CCB nº X, Contrato nº Y, NP, duplicata, etc.\n"
        "- garantia_tipo: alienação fiduciária de imóvel, hipoteca, penhor de ações, fiança, etc.\n"
        "- bem_garantia: descrição do bem ou imóvel dado em garantia\n"
        "- valor_original: número sem formatação (reais), ou null\n"
        "- data_base: DD/MM/AAAA, ou null\n"
        "- valor_atualizado: valor habilitado/atualizado em reais, ou null\n"
        "- status: 'Habilitado' | 'Impugnado' | 'Divergente' | 'Extraconcursal' | 'Incluído' | 'Excluído'\n"
        "- processo_habilitacao: nº do incidente de habilitação ou impugnação, ou null\n"
        "- pagina_referencia: página(s) onde a informação aparece, ex: 'p. 234', 'p. 45-46'\n"
        "- conflitos: descreva divergências entre AJ e credor, ou entre credores, sobre valor, classe "
        "ou garantia (com referência de página). null se não houver.\n"
        "- questoes_controversas: impugnações pendentes, questões não decididas, pontos que ainda "
        "dependem de decisão judicial (com página). null se não houver.\n"
        "- observacoes: outras informações relevantes\n\n"

        "═══ REGRAS CRÍTICAS ═══\n"
        "- TABELAS ESCANEADAS: mesmo que a OCR seja imperfeita, extraia TODOS os campos visíveis. "
        "Não omita nenhuma linha da tabela, mesmo que incompleta. Informe 'dados parciais (pág. X)' "
        "em observacoes quando a leitura for incerta.\n"
        "- Não omita nenhum credor, mesmo que incompleto.\n"
        "- Se o valor no edital divergir da manifestação do AJ, registre AMBOS em conflitos.\n"
        "- Se houver impugnação pendente (credor ou recuperando impugnou o crédito de outro), "
        "registre em questoes_controversas.\n"
        "- Informações de garantia e instrumento podem vir de manifestações do AJ ou das impugnações — "
        "complemente o que estiver no edital.\n\n"

        "Retorne APENAS JSON válido (sem markdown) com esta estrutura:\n"
        "{\n"
        "  \"edital_utilizado\": \"Art. 18 (definitivo)\" | \"Primeiro edital + manifestação AJ\" | \"Primeiro edital\" | \"Lista informal\",\n"
        "  \"paginas_edital\": \"ex: p. 234-267\" ou null,\n"
        "  \"observacao_fontes\": \"explique quais seções/páginas foram usadas e se há lacunas\",\n"
        "  \"credores\": [...]\n"
        "}\n\n"
        f"TEXTO EXTRAÍDO:\n{texto[:150_000]}"
    )
    config = types.GenerateContentConfig(response_mime_type="application/json")

    def _fn():
        return client.models.generate_content(model=model, contents=[prompt], config=config).text

    raw = _retry(_fn)
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"```(?:json)?\n?", "", raw).strip().rstrip("`").strip()
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            return {"edital_utilizado": "desconhecido", "paginas_edital": None,
                    "observacao_fontes": None, "credores": data}
        return data
    except json.JSONDecodeError:
        return {"edital_utilizado": "erro", "paginas_edital": None,
                "observacao_fontes": None, "credores": []}


def _gerar_excel_credores(resultado: dict) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    credores           = resultado.get("credores", [])
    edital_utilizado   = resultado.get("edital_utilizado", "")
    paginas_edital     = resultado.get("paginas_edital", "")
    observacao_fontes  = resultado.get("observacao_fontes", "")

    COR_HEADER    = "1F4E79"
    COR_SUB       = "BDD7EE"
    COR_TOTAL     = "2E75B6"
    COR_CONFLITO  = "FFD0D0"   # vermelho claro — conflito identificado
    COR_CONTROVER = "FFF3CD"   # amarelo claro — questão controversa pendente
    COR_META      = "EBF3FB"   # azul muito claro — linha de metadados

    thin  = Side(style="thin", color="D0D0D0")
    borda = Border(left=thin, right=thin, top=thin, bottom=thin)

    COLS = [
        ("Nome do Credor",                  48),
        ("CPF/CNPJ",                        18),
        ("Classe",                           22),
        ("Natureza do Crédito",              22),
        ("Instrumento / Lastro",             32),
        ("Garantia (tipo)",                  26),
        ("Bem dado em Garantia",             36),
        ("Valor Original (R$)",              18),
        ("Data Base",                        12),
        ("Valor Atualizado (R$)",            18),
        ("Status",                           16),
        ("Proc. Habilitação / Impugnação",   30),
        ("Página(s)",                        14),
        ("Conflitos (divergências)",         44),
        ("Questões Controversas",            44),
        ("% do Total",                       11),
        ("% da Classe",                      11),
        ("Observações",                      42),
    ]
    N = len(COLS)
    last_col = get_column_letter(N)

    wb = Workbook()
    ws = wb.active
    ws.title = "QGC"

    # ── Linha 1: título ────────────────────────────────────────────────────────
    ws.merge_cells(f"A1:{last_col}1")
    t = ws["A1"]
    t.value = "QUADRO GERAL DE CREDORES — RECUPERAÇÃO JUDICIAL"
    t.font      = Font(name="Calibri", bold=True, size=14, color="FFFFFF")
    t.fill      = PatternFill("solid", fgColor=COR_HEADER)
    t.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 32

    # ── Linha 2: metadados de fonte ───────────────────────────────────────────
    ws.merge_cells(f"A2:{last_col}2")
    meta_txt = f"Edital base: {edital_utilizado}"
    if paginas_edital:
        meta_txt += f"  |  Páginas: {paginas_edital}"
    if observacao_fontes:
        meta_txt += f"  |  {observacao_fontes}"
    m = ws["A2"]
    m.value     = meta_txt
    m.font      = Font(name="Calibri", italic=True, size=9, color="1F4E79")
    m.fill      = PatternFill("solid", fgColor=COR_META)
    m.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[2].height = 28

    # ── Linha 3: legenda de cores ─────────────────────────────────────────────
    ws.merge_cells(f"A3:{last_col}3")
    leg = ws["A3"]
    leg.value     = "Legenda:   Fundo vermelho = conflito identificado entre partes   |   Fundo amarelo = questão controversa pendente de decisão"
    leg.font      = Font(name="Calibri", italic=True, size=8, color="555555")
    leg.fill      = PatternFill("solid", fgColor="F8F8F8")
    leg.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[3].height = 18

    # ── Linha 4: cabeçalho de colunas ─────────────────────────────────────────
    HR = 4
    for ci, (h, w) in enumerate(COLS, 1):
        cell = ws.cell(row=HR, column=ci, value=h)
        cell.font      = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
        cell.fill      = PatternFill("solid", fgColor=COR_HEADER)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border    = borda
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.row_dimensions[HR].height = 34
    ws.freeze_panes = ws.cell(row=HR + 1, column=1)

    total_geral = sum(float(c.get("valor_atualizado") or 0) for c in credores)

    grupos: dict = defaultdict(list)
    for c in credores:
        grupos[_normalizar_classe(c.get("classe", ""))].append(c)

    def fv(v):
        try: return float(v) if v not in (None, "") else None
        except: return None

    cur = HR + 1
    classes_ord   = [cl for cl in _CLASSE_ORDER if cl in grupos]
    classes_extra = [cl for cl in grupos if cl not in _CLASSE_ORDER]

    for classe in classes_ord + classes_extra:
        lista        = grupos[classe]
        fill_cl      = PatternFill("solid", fgColor=_CLASSE_FILL.get(classe, "FFFFFF"))
        total_classe = sum(float(c.get("valor_atualizado") or 0) for c in lista)

        for credor in lista:
            va = fv(credor.get("valor_atualizado"))
            vo = fv(credor.get("valor_original"))
            pct_tot = (va / total_geral)  if (va and total_geral)  else None
            pct_cl  = (va / total_classe) if (va and total_classe) else None

            tem_conflito  = bool(credor.get("conflitos"))
            tem_controver = bool(credor.get("questoes_controversas"))

            if tem_conflito:
                fill_row = PatternFill("solid", fgColor=COR_CONFLITO)
            elif tem_controver:
                fill_row = PatternFill("solid", fgColor=COR_CONTROVER)
            else:
                fill_row = fill_cl

            row_vals = [
                credor.get("nome"),
                credor.get("cpf_cnpj"),
                classe,
                credor.get("natureza"),
                credor.get("instrumento"),
                credor.get("garantia_tipo"),
                credor.get("bem_garantia"),
                vo,
                credor.get("data_base"),
                va,
                credor.get("status"),
                credor.get("processo_habilitacao"),
                credor.get("pagina_referencia"),
                credor.get("conflitos"),
                credor.get("questoes_controversas"),
                pct_tot,
                pct_cl,
                credor.get("observacoes"),
            ]

            WRAP_COLS = {1, 5, 7, 13, 14, 15, 18}
            for ci, val in enumerate(row_vals, 1):
                cell = ws.cell(row=cur, column=ci, value=val)
                cell.fill      = fill_row
                cell.font      = Font(name="Calibri", size=10)
                cell.border    = borda
                cell.alignment = Alignment(vertical="center", wrap_text=(ci in WRAP_COLS))
                if ci in (8, 10):
                    cell.number_format = '#,##0.00'
                elif ci in (16, 17):
                    cell.number_format = '0.00%'
            cur += 1

        # Subtotal da classe
        sf = PatternFill("solid", fgColor=COR_SUB)
        sb = Font(name="Calibri", bold=True, size=10)
        for ci in range(1, N + 1):
            c = ws.cell(row=cur, column=ci)
            c.fill = sf; c.font = sb; c.border = borda
        ws.cell(row=cur, column=1).value = f"Subtotal — {classe}  ({len(lista)} credores)"
        c10 = ws.cell(row=cur, column=10, value=total_classe)
        c10.number_format = '#,##0.00'
        if total_geral:
            c16 = ws.cell(row=cur, column=16, value=total_classe / total_geral)
            c16.number_format = '0.00%'
        cur += 1

    # Total geral
    tf   = PatternFill("solid", fgColor=COR_TOTAL)
    tfnt = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    for ci in range(1, N + 1):
        c = ws.cell(row=cur, column=ci)
        c.fill = tf; c.font = tfnt; c.border = borda
    ws.cell(row=cur, column=1).value = f"TOTAL GERAL  ({len(credores)} credores)"
    c10 = ws.cell(row=cur, column=10, value=total_geral)
    c10.number_format = '#,##0.00'
    ws.row_dimensions[cur].height = 22

    ws.auto_filter.ref = f"A{HR}:{last_col}{cur - 1}"

    out = os.path.join(tempfile.gettempdir(), "qgc_credores_rj.xlsx")
    wb.save(out)
    return out


def rj_gerar_excel_credores(relatorio: str, texto_bruto: str = ""):
    fonte = texto_bruto.strip() if texto_bruto and texto_bruto.strip() else relatorio.strip()
    if not fonte:
        return gr.update(value=None, visible=False), "Gere uma análise primeiro."
    try:
        def _extrair_com(client, _indice):
            resultado_local = _extrair_credores_json(
                fonte, client, GEMINI_MODEL_ESTRUTURADO
            )
            credores_local = resultado_local.get("credores", [])
            if not credores_local and texto_bruto and relatorio.strip() and fonte != relatorio.strip():
                resultado_local = _extrair_credores_json(
                    relatorio, client, GEMINI_MODEL_ESTRUTURADO
                )
            return resultado_local

        resultado = _executar_com_failover_gemini(
            _get_gemini_clients(), _extrair_com
        )
        credores = resultado.get("credores", [])
        if not credores:
            return gr.update(value=None, visible=False), "Nenhum credor identificado. Verifique se o QGC consta no PDF."
        path = _gerar_excel_credores(resultado)
        edital = resultado.get("edital_utilizado", "")
        conflitos = sum(1 for c in credores if c.get("conflitos"))
        controver = sum(1 for c in credores if c.get("questoes_controversas"))
        msg = f"{len(credores)} credor(es) exportados"
        if edital:
            msg += f" | Edital base: {edital}"
        if conflitos:
            msg += f" | {conflitos} com conflito"
        if controver:
            msg += f" | {controver} com questão controversa"
        msg += _aviso_truncamento(len(fonte), LIMITE_TRUNCAMENTO_CREDORES)
        return gr.update(value=path, visible=True), msg
    except Exception:
        return gr.update(value=None, visible=False), "❌ Erro:\n\n" + traceback.format_exc()



