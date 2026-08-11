# -*- coding: utf-8 -*-
"""Dossiê Prévia — Análise Prévia de Viabilidade de Negócios (PPA Invista).

Triagem rápida do caso, mais enxuta que o dossiê desalinhado: dados gerais, um quadro
por processo (com constrições e recursos), os imóveis localizados e o esqueleto do
passivo. Reaproveita a MESMA extração do dossiê desalinhado — o material de origem é o
mesmo processo, então uma segunda chamada de IA só gastaria cota para reextrair o que já
foi extraído.

Campos que dependem de fonte externa ao processo (e-CAC, certidões, pesquisas de bens,
escopo do negócio) ficam em branco de propósito: preenchê-los por inferência criaria
fato onde há só o que o analista ainda vai apurar.
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as _DocxTable

from dossie_ppa import (
    _chave_rotulo,
    _completar_com_relatorio,
    _extrair_dados,
    _manter_tabela_inteira,
    _preencher_tabela_chave_valor,
    _substituir_texto_celula,
    _texto_analise,
    normalizar_referencias_objeto,
)

_TEMPLATE_PREVIA = Path(__file__).parent / "assets" / "Parecer_Previa_PPA.docx"

# Rótulos que abrem cada bloco repetível do template.
_ROTULO_PROCESSO = "dados do processo"
_ROTULO_CONSTRICOES = "constrições vigentes"
_ROTULO_RECURSO = "recurso nº"
_ROTULO_IMOVEL = "imóvel nº"


def _rotulo_tabela(tabela) -> str:
    """Primeira célula da primeira linha — é o que nomeia o quadro no template."""
    if not tabela.rows:
        return ""
    return _chave_rotulo(tabela.rows[0].cells[0].text)


def _elementos_corpo(doc) -> list:
    return list(doc.element.body.iterchildren())


def _indice_da_tabela(elementos, doc, rotulo: str, inicio: int = 0) -> int:
    for indice in range(inicio, len(elementos)):
        elemento = elementos[indice]
        if elemento.tag == qn("w:tbl") and _rotulo_tabela(_DocxTable(elemento, doc)).startswith(rotulo):
            return indice
    raise ValueError(f"Quadro não localizado no template da Prévia: {rotulo}")


def _substituir_intervalo(doc, inicio: int, fim: int, novos: list) -> None:
    """Troca os elementos [inicio, fim) do corpo pelos novos, preservando a ordem."""
    corpo = doc.element.body
    antigos = _elementos_corpo(doc)[inicio:fim]
    ancora = antigos[0]
    for elemento in novos:
        ancora.addprevious(elemento)
    for elemento in antigos:
        corpo.remove(elemento)


def _clonar(elementos: list) -> list:
    return [deepcopy(elemento) for elemento in elementos]


# ── Preenchimento de cada quadro ──────────────────────────────────────────────

def _preencher_dados_do_processo(tabela, credito: dict) -> None:
    # "Risco de Prescrição Superficial" fica fora do mapa: é caixa de marcar (☐ Sim ☐ Não)
    # e a resposta é juízo do analista, não algo que se extraia do texto do processo.
    _preencher_tabela_chave_valor(tabela, {
        "Número do processo": credito.get("numero_processo", ""),
        "Executado(s)": credito.get("executados_info", ""),
        "Advs:": credito.get("exequente_info", ""),
        "Data de distribuição": credito.get("data_distribuicao", ""),
        "SOP": credito.get("sop", ""),
        "SAT": credito.get("sat", ""),
        "Lastro / Instrumento": credito.get("lastro", ""),
        "Garantia": credito.get("garantia", ""),
        "Status": credito.get("status_processo", ""),
    })
    _manter_tabela_inteira(tabela)


def _preencher_constricoes(tabela, credito: dict) -> None:
    """TIPO | DESCRIÇÃO | VALOR (R$) | STATUS — o cabeçalho ocupa duas linhas
    (título mesclado + rótulos), então os dados começam na terceira."""
    registros = [
        [
            str(c.get("tipo", "") or ""),
            str(c.get("descricao", "") or ""),
            str(c.get("valor", "") or ""),
            str(c.get("status", "") or ""),
        ]
        for c in (credito.get("constricoes") or [])
        if isinstance(c, dict)
    ]
    _preencher_grade_com_duplo_cabecalho(tabela, registros)


def _preencher_grade_com_duplo_cabecalho(tabela, registros: list) -> None:
    dados = list(registros)
    alvo = max(len(dados), 1)
    while len(tabela.rows) - 2 < alvo:
        tabela._tbl.append(deepcopy(tabela.rows[-1]._tr))
    while len(tabela.rows) - 2 > alvo:
        tabela._tbl.remove(tabela.rows[-1]._tr)
    for indice, linha in enumerate(tabela.rows[2:]):
        valores = dados[indice] if indice < len(dados) else []
        for coluna in range(len(tabela.columns)):
            _substituir_texto_celula(
                linha.cells[coluna], valores[coluna] if coluna < len(valores) else ""
            )
    _manter_tabela_inteira(tabela)


def _preencher_recurso_previa(tabela, recurso: dict, numero: int) -> None:
    _substituir_texto_celula(tabela.rows[0].cells[0], f"Recurso nº {numero}")
    _preencher_tabela_chave_valor(tabela, {
        "Nº do Processo": recurso.get("numero_processo", "") or recurso.get("data_dist", ""),
        "Polo ativo": recurso.get("recorrente", ""),
        "Finalidade/Matéria": recurso.get("tese", "") or recurso.get("decisao_recorrida", ""),
        "Status Atual": recurso.get("status", ""),
    })
    _manter_tabela_inteira(tabela)


_RE_PREFIXO_MATRICULA = re.compile(r"^\s*matr[íi]cula\s*(?:n[ºo°]\.?)?\s*", re.IGNORECASE)


def _numero_da_matricula(valor) -> str:
    """A extração devolve "Matrícula nº 30.174" e a linha do template já se chama
    "Matrícula nº" — sem tirar o prefixo o quadro sai com o rótulo duplicado."""
    return _RE_PREFIXO_MATRICULA.sub("", str(valor or "")).strip()


def _preencher_imovel(tabela, ativo: dict, numero: int) -> None:
    # "Possível pequena propriedade rural" é caixa de marcar — fica para o analista.
    _substituir_texto_celula(tabela.rows[0].cells[0], f"Imóvel nº {numero}")
    _preencher_tabela_chave_valor(tabela, {
        "Matrícula nº": _numero_da_matricula(ativo.get("matricula", "")),
        "Comarca": ativo.get("comarca", ""),
        "Proprietário": ativo.get("proprietario_atual", ""),
        "Descrição": ativo.get("descricao", ""),
        "Ônus": ativo.get("onus_vigentes", ""),
        "Avaliação": ativo.get("vm", "") or ativo.get("vp", ""),
    })
    _manter_tabela_inteira(tabela)


# ── Blocos repetíveis ─────────────────────────────────────────────────────────

_ROTULOS_DO_BLOCO_PROCESSO = (_ROTULO_PROCESSO, _ROTULO_CONSTRICOES, _ROTULO_RECURSO)


def _fim_do_bloco(elementos, doc, inicio: int, rotulos: tuple) -> int:
    """Primeiro elemento DEPOIS do bloco: o bloco é a sequência de quadros com esses
    rótulos, mais os parágrafos vazios que os separam."""
    fim = inicio
    for indice in range(inicio, len(elementos)):
        elemento = elementos[indice]
        if elemento.tag == qn("w:tbl"):
            if not _rotulo_tabela(_DocxTable(elemento, doc)).startswith(rotulos):
                break
            fim = indice + 1
        elif elemento.tag == qn("w:p"):
            if "".join(elemento.itertext()).strip():
                break
            fim = indice + 1
        else:
            break
    return fim


def _expandir_processos(doc, creditos: list) -> None:
    """O template traz o quadro do processo três vezes; vira um por crédito."""
    elementos = _elementos_corpo(doc)
    inicio = _indice_da_tabela(elementos, doc, _ROTULO_PROCESSO)
    fim = _fim_do_bloco(elementos, doc, inicio, _ROTULOS_DO_BLOCO_PROCESSO)

    # Uma cópia do padrão vai até o próximo "Dados do Processo" dentro da região.
    proximo = fim
    for indice in range(inicio + 1, fim):
        elemento = elementos[indice]
        if elemento.tag == qn("w:tbl") and _rotulo_tabela(_DocxTable(elemento, doc)).startswith(_ROTULO_PROCESSO):
            proximo = indice
            break
    padrao = elementos[inicio:proximo]

    novos = []
    for numero, credito in enumerate(list(creditos) or [{}], 1):
        grupo = _clonar(padrao)
        _preencher_grupo_processo(doc, grupo, credito, numero)
        novos.extend(grupo)
    _substituir_intervalo(doc, inicio, fim, novos)


def _preencher_grupo_processo(doc, grupo: list, credito: dict, numero: int) -> None:
    """Preenche um bloco clonado e replica o quadro de recurso conforme a quantidade."""
    recursos = [r for r in (credito.get("recursos") or []) if isinstance(r, dict)]

    indice_recurso = None
    for indice, elemento in enumerate(grupo):
        if elemento.tag != qn("w:tbl"):
            continue
        tabela = _DocxTable(elemento, doc)
        rotulo = _rotulo_tabela(tabela)
        if rotulo.startswith(_ROTULO_PROCESSO):
            _preencher_dados_do_processo(tabela, credito)
        elif rotulo.startswith(_ROTULO_CONSTRICOES):
            _preencher_constricoes(tabela, credito)
        elif rotulo.startswith(_ROTULO_RECURSO):
            indice_recurso = indice

    if indice_recurso is None:
        return

    fim_recurso = indice_recurso + 1
    while fim_recurso < len(grupo) and grupo[fim_recurso].tag == qn("w:p"):
        fim_recurso += 1
    padrao_recurso = grupo[indice_recurso:fim_recurso]

    clones = []
    for ordem, recurso in enumerate(recursos or [{}], 1):
        copia = _clonar(padrao_recurso)
        for elemento in copia:
            if elemento.tag == qn("w:tbl"):
                _preencher_recurso_previa(_DocxTable(elemento, doc), recurso, ordem)
                break
        clones.extend(copia)
    grupo[indice_recurso:fim_recurso] = clones


def _devedores(dados: dict) -> list:
    """A extração devolve os executados como "Nome1 · CPF ...; Nome2 · CNPJ ..."."""
    bruto = str(dados.get("executados", "") or "")
    return [parte.strip() for parte in bruto.split(";") if parte.strip()]


def _preencher_passivo_fiscal(tabela, devedores: list) -> None:
    """Nomeia um bloco por devedor no quadro fiscal.

    Os VALORES (e-CAC, dívida ativa, certidões) não saem do processo — vêm do e-CAC e
    das certidões, que entram depois. O que dá para adiantar é a lista de quem precisa
    ser consultado, e o template só traz dois blocos fixos: com três devedores, um
    ficaria de fora sem a replicação.
    """
    nomes = list(devedores) or [""]
    bloco = [deepcopy(tabela.rows[indice]._tr) for indice in range(3)]  # cabeçalho + 2 linhas

    total_tr = tabela.rows[6]._tr  # "TOTAL CONSOLIDADO DO GRUPO"
    for linha in list(tabela.rows)[:6]:
        tabela._tbl.remove(linha._tr)

    for numero, nome in enumerate(nomes, 1):
        copia = [deepcopy(tr) for tr in bloco]
        for tr in copia:
            total_tr.addprevious(tr)
        cabecalho = tabela.rows[(numero - 1) * 3]
        _substituir_texto_celula(
            cabecalho.cells[0], nome or f"DEVEDOR {numero} — [Nome / CPF-CNPJ]"
        )
        for linha in tabela.rows[(numero - 1) * 3 + 1:(numero - 1) * 3 + 3]:
            _substituir_texto_celula(linha.cells[1], "")


def _expandir_imoveis(doc, ativos: list) -> None:
    elementos = _elementos_corpo(doc)
    inicio = _indice_da_tabela(elementos, doc, _ROTULO_IMOVEL)
    fim = inicio + 1
    while fim < len(elementos) and elementos[fim].tag == qn("w:p"):
        if "".join(elementos[fim].itertext()).strip():
            break
        fim += 1
    padrao = elementos[inicio:fim]

    novos = []
    for numero, ativo in enumerate(list(ativos) or [{}], 1):
        grupo = _clonar(padrao)
        for elemento in grupo:
            if elemento.tag == qn("w:tbl"):
                _preencher_imovel(_DocxTable(elemento, doc), ativo, numero)
                break
        novos.extend(grupo)
    _substituir_intervalo(doc, inicio, fim, novos)


# ── Montagem ──────────────────────────────────────────────────────────────────

def _quadro_fiscal(doc):
    """Localiza o quadro do passivo fiscal pelas linhas de endividamento.

    Não usa o cabeçalho do devedor como âncora: depois do preenchimento ele deixa de
    ser "DEVEDOR 1 — ..." e passa a ser o nome real, e o quadro sumiria da busca.
    """
    for tabela in doc.tables:
        if any(_chave_rotulo(linha.cells[0].text).startswith("endividamento fiscal")
               for linha in tabela.rows):
            return tabela
    raise ValueError("Quadro do passivo fiscal não localizado no template da Prévia.")


def _primeiro_paragrafo(texto: str) -> str:
    partes = [p.strip() for p in str(texto or "").split("\n") if p.strip()]
    return partes[0] if partes else ""


def _build_previa(dados: dict) -> str:
    if not _TEMPLATE_PREVIA.exists():
        raise FileNotFoundError(f"Template da Prévia não localizado: {_TEMPLATE_PREVIA}")

    caminho = os.path.join(tempfile.gettempdir(), "Dossie_Previa_PPA.docx")
    shutil.copy2(_TEMPLATE_PREVIA, caminho)
    doc = Document(caminho)

    _preencher_tabela_chave_valor(doc.tables[1], {
        "Nome do Caso": dados.get("nome_caso", ""),
        "Data da Análise": dados.get("data_analise", ""),
        "Advogada Responsável": dados.get("advogada_responsavel", ""),
    })

    # "Escopo da análise" é a tese do negócio, decidida fora do processo — fica em branco.
    _preencher_tabela_chave_valor(doc.tables[2], {
        "Credor": dados.get("exequentes", ""),
        "Devedor": dados.get("executados", ""),
        "SAT": dados.get("sat_total", ""),
        "Resumo do caso": _primeiro_paragrafo(dados.get("consideracoes_gerais", ""))
                          or _texto_analise(dados.get("teses_principais", "")),
        "Riscos Jurídicos": _texto_analise(dados.get("risco_juridico", "")),
    })

    creditos = [c for c in (dados.get("creditos") or []) if isinstance(c, dict)]
    _expandir_processos(doc, creditos)

    ativos = [a for a in (dados.get("ativos") or []) if isinstance(a, dict)]
    _expandir_imoveis(doc, ativos)

    _preencher_passivo_fiscal(_quadro_fiscal(doc), _devedores(dados))

    doc.save(caminho)
    return caminho


def gerar_previa_word(texto_completo: str, relatorio: str, client, model: str) -> str:
    """Extrai os dados do processo e devolve o caminho do Dossiê Prévia em Word."""
    dados = _extrair_dados(texto_completo, relatorio, client, model)
    dados = _completar_com_relatorio(dados, relatorio, client, model)
    nomes_pdf = set(re.findall(r"(?:arquivo:\s*|^---\s*)([^\n—]+)", texto_completo or "", re.MULTILINE))
    dados = normalizar_referencias_objeto(dados, multiplos_pdfs=len(nomes_pdf) > 1)
    return _build_previa(dados)
