"""Dossiê Prévia: quadros repetíveis escalam com o caso e nada do exemplo vaza.

O template nasceu de um parecer real preenchido (caso Watt Tecnologia) e traz três
cópias do quadro de processo. Um dossiê novo não pode herdar nem os dados daquele caso
nem as cópias sobrando.
"""
from __future__ import annotations

import json
import sys
import types as pytypes
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

import dossie_previa  # noqa: E402


DADOS = {
    "nome_caso": "Banco do Brasil x Agropecuária Teste",
    "data_analise": "10/08/2026",
    "advogada_responsavel": "Julia de Oliveira Bernardo da Silva",
    "exequentes": "Banco do Brasil S.A. · CNPJ 00.000.000/0001-91",
    "executados": "Agropecuária Teste Ltda · CNPJ 12.345.678/0001-90",
    "sat_total": "R$ 4.200.000,00 (fls. 210)",
    "risco_juridico": "Prescrição intercorrente em discussão (Mov. 44).",
    "consideracoes_gerais": "Execução com penhora deferida sobre a futura sede.\n\nSegundo parágrafo.",
    "creditos": [
        {
            "numero_processo": "0001234-56.2019.8.16.0014",
            "executados_info": "Agropecuária Teste Ltda · Adv: Silva Advogados | OAB/PR 1234",
            "exequente_info": "Banco do Brasil · Adv: Jurídico BB | OAB/DF 999",
            "data_distribuicao": "10/03/2019 (fls. 2)",
            "sop": "R$ 2.000.000,00", "sat": "R$ 4.200.000,00",
            "lastro": "CCB nº 40/00123-4 (fls. 12)",
            "garantia": "Hipoteca cedular sobre a matrícula 30.174",
            "status_processo": "Em fase de penhora (Mov. 51)",
            "constricoes": [
                {"tipo": "Penhora", "descricao": "Imóvel matrícula 30.174", "valor": "R$ 900.000,00",
                 "status": "Deferida em 12/05/2024 (Mov. 44)"},
                {"tipo": "Sisbajud", "descricao": "Bloqueio de ativos", "valor": "R$ 12.500,00",
                 "status": "Ativa (Mov. 47)"},
            ],
            "recursos": [
                {"recorrente": "Agropecuária Teste Ltda", "tese": "Excesso de execução",
                 "status": "Aguardando julgamento", "decisao_recorrida": "Decisão de penhora"},
                {"recorrente": "Banco do Brasil", "tese": "Majoração de honorários",
                 "status": "Provido em parte"},
            ],
        },
        {
            "numero_processo": "0009876-54.2021.8.16.0014",
            "executados_info": "João da Silva · CPF 111.222.333-44",
            "data_distribuicao": "05/06/2021",
            "status_processo": "Suspensa (art. 921)",
            "constricoes": [],
            "recursos": [],
        },
    ],
    "ativos": [
        {"matricula": "Matrícula nº 30.174", "comarca": "Montes Claros/MG",
         "proprietario_atual": "Agropecuária Teste Ltda", "descricao": "Lotes 08 a 13 da quadra 07",
         "onus_vigentes": "R-3-30174: hipoteca cedular", "vm": "R$ 1.500.000,00"},
        {"matricula": "Matrícula nº 9.596", "comarca": "Curitiba/PR",
         "proprietario_atual": "João da Silva", "descricao": "Apartamento 402",
         "onus_vigentes": "Não há", "vm": "R$ 480.000,00"},
    ],
}

VESTIGIOS_DO_EXEMPLO = ("Watt", "WATT", "Bormanas", "Montes Claros/MG — Watt",
                        "97.525.889", "31/07/2026")


def _texto_inteiro(doc) -> str:
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            partes.extend(c.text for c in linha.cells)
    return "\n".join(partes)


def _quadros(doc, rotulo: str) -> list:
    return [t for t in doc.tables
            if t.rows and dossie_previa._chave_rotulo(t.rows[0].cells[0].text).startswith(rotulo)]


@pytest.fixture(scope="module")
def documento():
    return Document(dossie_previa._build_previa(DADOS))


def testar_template_existe_e_esta_limpo():
    assert dossie_previa._TEMPLATE_PREVIA.exists(), "template da Prévia precisa ir junto no deploy"
    bruto = _texto_inteiro(Document(str(dossie_previa._TEMPLATE_PREVIA)))
    assert "Watt" not in bruto and "Bormanas" not in bruto


def testar_capa_e_dados_gerais(documento):
    texto = _texto_inteiro(documento)
    assert "Banco do Brasil x Agropecuária Teste" in texto
    assert "Julia de Oliveira Bernardo da Silva" in texto
    assert "R$ 4.200.000,00" in texto
    assert "Prescrição intercorrente em discussão" in texto
    # Resumo do caso usa o primeiro parágrafo das considerações gerais.
    assert "Execução com penhora deferida sobre a futura sede." in texto
    assert "Segundo parágrafo." not in texto


def testar_um_quadro_de_processo_por_credito(documento):
    quadros = _quadros(documento, dossie_previa._ROTULO_PROCESSO)
    assert len(quadros) == len(DADOS["creditos"]), "o template tem 3 cópias fixas; tem que virar uma por crédito"
    texto = _texto_inteiro(documento)
    assert "0001234-56.2019.8.16.0014" in texto
    assert "0009876-54.2021.8.16.0014" in texto


def testar_constricoes_viram_linhas_do_quadro(documento):
    constricoes = _quadros(documento, dossie_previa._ROTULO_CONSTRICOES)
    assert len(constricoes) == 2
    # Duas linhas de cabeçalho (título mesclado + rótulos) e duas de dados.
    assert len(constricoes[0].rows) == 4
    primeira = [c.text.strip() for c in constricoes[0].rows[2].cells]
    assert primeira[0] == "Penhora"
    assert "30.174" in primeira[1]
    assert "Deferida em 12/05/2024" in primeira[3]
    # Crédito sem constrição mantém uma linha vazia, sem herdar a do crédito anterior.
    assert len(constricoes[1].rows) == 3
    assert all(not c.text.strip() for c in constricoes[1].rows[2].cells)


def testar_recursos_sao_replicados_e_numerados(documento):
    recursos = _quadros(documento, dossie_previa._ROTULO_RECURSO)
    assert len(recursos) == 3, "2 recursos no 1º crédito + 1 quadro vazio no 2º"
    rotulos = [r.rows[0].cells[0].text.strip() for r in recursos]
    assert rotulos[:2] == ["Recurso nº 1", "Recurso nº 2"]
    texto = _texto_inteiro(documento)
    assert "Excesso de execução" in texto and "Majoração de honorários" in texto


def testar_um_quadro_por_imovel(documento):
    imoveis = _quadros(documento, dossie_previa._ROTULO_IMOVEL)
    assert len(imoveis) == len(DADOS["ativos"])
    assert [i.rows[0].cells[0].text.strip() for i in imoveis] == ["Imóvel nº 1", "Imóvel nº 2"]
    texto = _texto_inteiro(documento)
    assert "Apartamento 402" in texto and "R$ 1.500.000,00" in texto


def testar_matricula_nao_repete_o_rotulo_da_linha(documento):
    """A extração devolve "Matrícula nº 30.174" e a linha já se chama "Matrícula nº"."""
    primeiro = _quadros(documento, dossie_previa._ROTULO_IMOVEL)[0]
    valores = {r.cells[0].text.strip(): r.cells[1].text.strip() for r in primeiro.rows[1:]}
    assert valores["Matrícula nº"] == "30.174"


def testar_nada_do_caso_de_exemplo_sobrevive(documento):
    texto = _texto_inteiro(documento)
    for vestigio in ("Watt", "WATT", "Bormanas", "97.525.889", "31/07/2026"):
        assert vestigio not in texto, f"vazou do template de exemplo: {vestigio}"


def testar_caixas_de_marcar_ficam_para_o_analista(documento):
    """Prescrição superficial, pequena propriedade rural e pesquisas de bens são
    juízo de quem analisa — o gerador não pode responder por ela."""
    texto = _texto_inteiro(documento)
    assert "☐ Sim" in texto
    assert "Risco de Prescrição Superficial" in texto
    assert "Pesquisa Prévia Realizada" in texto
    assert "Análise por viés desalinhado?" in texto


def testar_quadro_fiscal_nomeia_um_bloco_por_devedor():
    """O template traz dois blocos fixos; com três devedores um ficaria de fora.
    Os valores continuam vazios — e-CAC e certidões não saem do processo."""
    dados = dict(DADOS, executados="Empresa A · CNPJ 1; Fulano · CPF 2; Beltrano · CPF 3")
    fiscal = dossie_previa._quadro_fiscal(Document(dossie_previa._build_previa(dados)))

    cabecalhos = [r.cells[0].text.strip() for r in fiscal.rows]
    assert "Empresa A · CNPJ 1" in cabecalhos
    assert "Fulano · CPF 2" in cabecalhos
    assert "Beltrano · CPF 3" in cabecalhos
    assert "TOTAL CONSOLIDADO DO GRUPO" in cabecalhos
    assert "DEVEDOR 2 — [Nome / CPF-CNPJ]" not in cabecalhos, "placeholder tem que ser substituído"

    linhas_ecac = [r for r in fiscal.rows if r.cells[0].text.strip().startswith("Endividamento Fiscal —")]
    assert len(linhas_ecac) == 6, "duas linhas por devedor"
    assert all(not r.cells[1].text.strip() for r in linhas_ecac), "valor fiscal não vem do processo"


def testar_caso_sem_creditos_nem_ativos_gera_esqueleto():
    doc = Document(dossie_previa._build_previa({"nome_caso": "Só a capa"}))
    assert len(_quadros(doc, dossie_previa._ROTULO_PROCESSO)) == 1
    assert len(_quadros(doc, dossie_previa._ROTULO_IMOVEL)) == 1
    assert "Só a capa" in _texto_inteiro(doc)


def testar_gerar_previa_word_liga_extracao_e_documento():
    """Caminho completo (extração → complemento → documento) com a IA dublada."""
    chamadas = []

    def _gerar(**kwargs):
        chamadas.append(kwargs.get("model"))
        return pytypes.SimpleNamespace(text=json.dumps(DADOS, ensure_ascii=False))

    cliente = pytypes.SimpleNamespace(models=pytypes.SimpleNamespace(generate_content=_gerar))

    caminho = dossie_previa.gerar_previa_word(
        "--- ato.pdf\ntexto extraído", "relatório consolidado", cliente, "modelo-x"
    )

    assert Path(caminho).name == "Dossie_Previa_PPA.docx"
    assert chamadas, "a extração precisa passar pelo cliente Gemini"
    texto = _texto_inteiro(Document(caminho))
    assert "Banco do Brasil x Agropecuária Teste" in texto
    assert "0001234-56.2019.8.16.0014" in texto
