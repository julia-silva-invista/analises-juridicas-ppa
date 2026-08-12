"""A instrução adicional deixou de ser um parágrafo solto no fim do prompt.

Antes o texto entrava como "INSTRUCOES ADICIONAIS: ..." depois de todas as regras e do
template — competindo com uma estrutura que enumera as seções e manda segui-la à risca.
Pedir "abra uma seção nova" era ignorado, e pedir "crie uma tabela no dossiê" não tinha
como funcionar: o dossiê nem recebia o texto.
"""
from __future__ import annotations

import inspect
import json
import sys
import types as pytypes
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

import dossie_ppa  # noqa: E402
import processos  # noqa: E402
from legal_prompts import bloco_instrucao_adicional  # noqa: E402
from report_template_processos import REPORT_TEMPLATE_INSTRUCTIONS  # noqa: E402


PEDIDO = "Verificar se há bem de família alegado e montar uma tabela dos leilões."


# ── O bloco em si ────────────────────────────────────────────────────────────

def testar_bloco_vazio_nao_polui_o_prompt():
    for vazio in ("", "   ", None):
        assert bloco_instrucao_adicional(vazio) == ""


def testar_bloco_e_rotulado_e_prioritario():
    bloco = bloco_instrucao_adicional(PEDIDO)
    assert "INSTRUÇÃO ADICIONAL DA ANALISTA (PRIORITÁRIA)" in bloco
    assert PEDIDO in bloco
    assert "ACRESCENTE um item novo" in bloco


def testar_bloco_nao_afrouxa_a_regra_de_nao_inventar():
    """A precedência é sobre a estrutura, não sobre a exigência de prova."""
    bloco = bloco_instrucao_adicional(PEDIDO)
    assert "NÃO autoriza inventar" in bloco
    assert "não consta" in bloco
    assert "referência" in bloco


def testar_secao_nova_pode_ser_desligada():
    """Nos processos relacionados e no dossiê a estrutura é fixa — lá o modelo não
    deve inventar seção."""
    bloco = bloco_instrucao_adicional(PEDIDO, permite_secao_nova=False)
    assert "ACRESCENTE um item novo" not in bloco
    assert "NÃO autoriza inventar" in bloco


# ── Onde é injetado ──────────────────────────────────────────────────────────

def testar_consolidacao_usa_o_bloco_nos_tres_pontos():
    fonte = inspect.getsource(processos)
    assert "INSTRUCOES ADICIONAIS: {instrucoes" not in fonte, "formato antigo ainda no código"
    assert fonte.count("bloco_instrucao_adicional(") >= 3


def testar_extracao_por_chunk_continua_sem_a_instrucao():
    """Se entrasse na extração, invalidaria o cache de chunk a cada texto novo."""
    assert "instrucoes" not in inspect.signature(processos._proc_extrair_chunk).parameters
    assert "bloco_instrucao_adicional" not in processos.PROMPT_EXTR_PROC


def testar_template_autoriza_item_novo():
    assert "INSTRUÇÃO ADICIONAL DA ANALISTA (PRIORITÁRIA)" in REPORT_TEMPLATE_INSTRUCTIONS
    assert "ACRESCENTE um item novo ao fim da subseção" in REPORT_TEMPLATE_INSTRUCTIONS


def testar_cache_do_template_foi_versionado():
    """O template fica cacheado 1h no Gemini; sem bumpar o nome, a versão velha
    continuaria sendo servida depois da mudança."""
    fonte = inspect.getsource(processos)
    assert 'display_name="invista_proc_template_v2"' in fonte
    assert "invista_proc_template_v1" not in fonte


def testar_dossies_recebem_a_instrucao():
    for funcao in (processos.proc_gerar_dossie, processos.proc_gerar_previa):
        assert "instrucoes" in inspect.signature(funcao).parameters, funcao.__name__
    assert "instrucoes" in inspect.signature(dossie_ppa.gerar_dossie_word).parameters
    assert "instrucoes" in inspect.signature(dossie_ppa._extrair_dados).parameters


def testar_prompt_do_dossie_so_pede_quadro_extra_quando_ha_instrucao():
    capturados = []

    def _gerar(**kwargs):
        capturados.append(kwargs["contents"][0].parts[0].text)
        return pytypes.SimpleNamespace(text="{}")

    cliente = pytypes.SimpleNamespace(models=pytypes.SimpleNamespace(generate_content=_gerar))

    dossie_ppa._extrair_dados("texto", "relatório", cliente, "modelo-x")
    assert "quadros_extras" not in capturados[-1]

    dossie_ppa._extrair_dados("texto", "relatório", cliente, "modelo-x", PEDIDO)
    assert "quadros_extras" in capturados[-1]
    assert PEDIDO in capturados[-1]


# ── Quadro extra no documento ────────────────────────────────────────────────

QUADRO = {
    "titulo": "Leilões designados",
    "colunas": ["DATA", "PRAÇA", "LANCE MÍNIMO"],
    "linhas": [["10/03/2026", "1ª", "R$ 900.000,00"],
               ["24/03/2026", "2ª", "R$ 450.000,00"]],
}


def _tabela_por_cabecalho(doc, colunas: list):
    """O template já tem uma tabela que começa por DATA (andamentos) — casar pelo
    cabeçalho inteiro, não pela primeira célula."""
    alvo = [c.upper() for c in colunas]
    for tabela in doc.tables:
        if tabela.rows and [c.text.strip().upper() for c in tabela.rows[0].cells] == alvo:
            return tabela
    return None


def testar_quadro_extra_entra_no_fim_do_dossie():
    caminho = dossie_ppa._build_doc({"nome_caso": "Teste", "quadros_extras": [QUADRO]})
    doc = Document(caminho)

    tabela = _tabela_por_cabecalho(doc, QUADRO["colunas"])
    assert tabela is not None, "o quadro pedido pela instrução não foi criado"
    assert len(tabela.rows) == 3
    assert [c.text.strip() for c in tabela.rows[1].cells] == ["10/03/2026", "1ª", "R$ 900.000,00"]

    titulos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert any("Leilões Designados" in t or "Leilões designados" in t for t in titulos)


def testar_sem_quadros_extras_o_documento_nao_muda():
    base = Document(dossie_ppa._build_doc({"nome_caso": "Teste"}))
    com_chave_vazia = Document(dossie_ppa._build_doc({"nome_caso": "Teste", "quadros_extras": []}))
    assert len(base.tables) == len(com_chave_vazia.tables)


@pytest.mark.parametrize("quadro", [
    {"titulo": "Sem colunas", "colunas": [], "linhas": [["a"]]},
    {"titulo": "Sem linhas", "colunas": ["A"], "linhas": []},
    {"titulo": "Linhas vazias", "colunas": ["A"], "linhas": [["", "  "]]},
    "não é dict",
])
def testar_quadro_malformado_nao_quebra_o_documento(quadro):
    """A chave vem de JSON gerado por IA; formato torto não pode derrubar a geração."""
    base = Document(dossie_ppa._build_doc({"nome_caso": "Teste"}))
    doc = Document(dossie_ppa._build_doc({"nome_caso": "Teste", "quadros_extras": [quadro]}))
    assert len(doc.tables) == len(base.tables)
