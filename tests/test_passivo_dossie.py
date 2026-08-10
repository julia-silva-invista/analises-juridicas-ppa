"""Ordem das colunas do passivo do dossiê: Nº CNJ → Distribuição → Executado → Exequente.

"Executado" tem que sair do POLO PASSIVO da planilha (quem figura como réu/executado),
não de "Vinculado à" — essa coluna diz apenas a quem a pesquisa estava atrelada.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

import coleta  # noqa: E402
import dossie_ppa  # noqa: E402


LINHA_PLANILHA = {
    "Número CNJ": "0001234-56.2019.8.16.0014",
    "Vinculado à": "JOÃO DA SILVA (CPF nº 111.222.333-44)",
    "Polo passivo": "AGROPECUÁRIA TESTE LTDA (CNPJ nº 12.345.678/0001-90); JOÃO DA SILVA",
    "Polo ativo": "FAZENDA PÚBLICA DO ESTADO DO PARANÁ",
    "Data da distribuição": "10/03/2019",
    "Valor da causa": 250000.0,
    "Situação atual": "Em andamento",
    "Saldo Atualizado estimado": 310000.0,
}


def _tabelas_de_passivo(doc):
    """Tabelas de processos do passivo: 7 colunas começando por Nº CNJ."""
    for heading, table in dossie_ppa._iter_headings_tables(doc):
        cabecalho = table.rows[0].cells
        if len(cabecalho) == 7 and "cnj" in dossie_ppa._chave_rotulo(cabecalho[0].text):
            yield heading, table


def testar_executado_vem_do_polo_passivo():
    proc = coleta._linha_processo_coleta(LINHA_PLANILHA)
    assert proc["passivo"] == "AGROPECUÁRIA TESTE LTDA; JOÃO DA SILVA"
    assert proc["vinc"] == "JOÃO DA SILVA"

    linha = dossie_ppa._proc_to_row(proc)
    assert linha[0] == "0001234-56.2019.8.16.0014"   # Nº CNJ
    assert linha[1] == "10/03/2019"                   # Distribuição
    assert linha[2] == "AGROPECUÁRIA TESTE LTDA; JOÃO DA SILVA"  # Executado
    assert linha[3] == "FAZENDA PÚBLICA DO ESTADO DO PARANÁ"     # Exequente


def testar_executado_cai_para_vinculado_sem_polo_passivo():
    """Predictus e planilhas antigas podem não trazer o polo passivo — melhor exibir
    "Vinculado à" do que deixar a coluna Executado vazia."""
    sem_passivo = dict(LINHA_PLANILHA, **{"Polo passivo": ""})
    linha = dossie_ppa._proc_to_row(coleta._linha_processo_coleta(sem_passivo))
    assert linha[2] == "JOÃO DA SILVA"


@pytest.mark.parametrize(
    "cabecalhos",
    [
        {"Polo passivo": "P", "Polo ativo": "A", "Situação atual": "S"},
        {"POLO PASSIVO ": "P", "Polo Ativo": "A", "Situacao atual": "S"},
        {"Partes passivas": "P", "Partes ativas": "A", "Status": "S"},
    ],
)
def testar_colunas_toleram_acento_caixa_e_sinonimo(cabecalhos):
    proc = coleta._linha_processo_coleta({"Número CNJ": "X", **cabecalhos})
    assert (proc["passivo"], proc["ativo"], proc["status"]) == ("P", "A", "S")


def testar_cabecalho_e_linhas_do_dossie_na_ordem_nova():
    caminho = dossie_ppa._build_doc({"nome_caso": "Teste"})
    doc = Document(caminho)

    tabelas = list(_tabelas_de_passivo(doc))
    assert len(tabelas) == 3, "esperado Fiscal + Trabalhista + Cível"
    for _, table in tabelas:
        rotulos = [dossie_ppa._chave_rotulo(c.text) for c in table.rows[0].cells]
        assert rotulos[1] == "distribuição"
        assert rotulos[2] == "executado"
        assert "vinculado" not in " ".join(rotulos)

    proc = coleta._linha_processo_coleta(LINHA_PLANILHA)
    preenchido = dossie_ppa.preencher_passivo_dossie(caminho, [proc], [proc], [proc])
    doc_final = Document(preenchido)

    tabelas_finais = list(_tabelas_de_passivo(doc_final))
    assert len(tabelas_finais) == 3
    for _, table in tabelas_finais:
        dados = [c.text.strip() for c in table.rows[1].cells]
        assert dados[0] == "0001234-56.2019.8.16.0014"
        assert dados[1] == "10/03/2019"
        assert dados[2].upper().startswith("AGROPECUÁRIA TESTE LTDA")
        assert dados[3].upper() == "FAZENDA PÚBLICA DO ESTADO DO PARANÁ"


def testar_reordenacao_do_cabecalho_e_idempotente():
    """Rodar de novo sobre um dossiê já convertido não pode embaralhar o cabeçalho."""
    caminho = dossie_ppa._build_doc({"nome_caso": "Teste"})
    proc = coleta._linha_processo_coleta(LINHA_PLANILHA)
    primeiro = dossie_ppa.preencher_passivo_dossie(caminho, [proc], [], [])
    segundo = dossie_ppa.preencher_passivo_dossie(primeiro, [proc], [], [])

    for _, table in _tabelas_de_passivo(Document(segundo)):
        rotulos = [dossie_ppa._chave_rotulo(c.text) for c in table.rows[0].cells]
        assert rotulos[1] == "distribuição"
        assert rotulos[2] == "executado"
