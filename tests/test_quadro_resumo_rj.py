"""Quadro-resumo de lastros e garantias no fim da Análise de Créditos em RJ.

A seção 5 detalha cada execução num quadro próprio — bom para ler caso a caso, ruim
para comparar. Este quadro põe lastro, garantia e valor lado a lado, uma linha por
processo, no mesmo estilo das demais tabelas do modelo.

Arquivo separado do test_checklist_rj.py de propósito: aquela bateria roda em ~1 min e
tem guarda de tempo próprio; não vale a pena arrastá-la para um teste de render.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

import checklist_rj as cr  # noqa: E402
import dossie_ppa as dp  # noqa: E402


EXECUCOES = [
    {"numero": "5219163-18.2026.8.09.0014",
     "lastro_fls": "Cédula de Crédito Bancário (CCB) nº 301.906.503 (fls. 12)",
     "garantia": "Penhor Rural; Intervenientes Garantes",
     "valor_causa": "R$ 241.263,81"},
    {"numero": "5548551-24.2025.8.09.0014",
     "lastro_fls": "Cédula de Crédito Bancário (CCB) nº 301.906.772",
     "garantia": "Penhor Cedular de Primeiro Grau; Hipotecas Cedulares",
     "valor_causa": "R$ 2.226.491,03"},
]

DADOS = {
    "credor": "Banco do Brasil S.A. · CNPJ 00.000.000/0001-91",
    "rj_numero": "1234567-89.2024.8.11.0000",
    "vara": "1ª Vara Empresarial",
    "requerentes": "Empresa Teste Ltda",
    "execucoes": EXECUCOES,
}


def _quadro(doc):
    """A tabela cujo cabeçalho começa por Nº e tem as cinco colunas do quadro."""
    for tabela in doc.tables:
        cabecalho = [c.text.strip().upper() for c in tabela.rows[0].cells]
        if cabecalho[:2] == ["Nº", "PROCESSO"]:
            return tabela
    return None


def _fill_da_celula(celula) -> str | None:
    tc_pr = celula._tc.tcPr
    shd = tc_pr.find(qn("w:shd")) if tc_pr is not None else None
    return shd.get(qn("w:fill")) if shd is not None else None


@pytest.fixture(scope="module")
def documento():
    return Document(cr._build_checklist_creditos(DADOS))


def testar_uma_linha_por_execucao(documento):
    tabela = _quadro(documento)
    assert tabela is not None, "o quadro-resumo tem que existir"
    assert len(tabela.rows) == 1 + len(EXECUCOES), "cabeçalho + uma linha por processo"

    primeira = [c.text.strip() for c in tabela.rows[1].cells]
    assert primeira[0] == "1"
    assert primeira[1] == "5219163-18.2026.8.09.0014"
    assert "301.906.503" in primeira[2]
    assert "Penhor Rural" in primeira[3]
    assert primeira[4] == "R$ 241.263,81"
    assert [c.text.strip() for c in tabela.rows[2].cells][0] == "2"


def testar_usa_o_estilo_do_modelo(documento):
    """Mesmo layout e mesmas cores das outras tabelas: header laranja, linhas brancas."""
    tabela = _quadro(documento)
    assert all(_fill_da_celula(c) == dp._LARANJA for c in tabela.rows[0].cells)
    assert all(_fill_da_celula(c) == dp._BRANCO for c in tabela.rows[1].cells)

    # O _write deixa um run vazio na frente; só os runs com texto carregam formatação.
    escritos = [r for c in tabela.rows[0].cells
                for r in c.paragraphs[0].runs if r.text.strip()]
    assert escritos
    assert all(str(r.font.color.rgb) == "FFFFFF" and r.bold for r in escritos), \
        "header do modelo é texto branco em negrito"

    larguras = [round(c.width.cm, 1) for c in tabela.rows[0].cells]
    assert larguras == [round(w, 1) for w in cr._WS_QUADRO_RESUMO]
    assert sum(larguras) < 17.0, "não pode estourar a área útil do A4"


def testar_fica_no_fim_do_documento(documento):
    """Depois da seção 5 e dos recursos órfãos. Os títulos saem em caixa normalizada
    pelo _sec_title ("6. Quadro-Resumo — ..."), então a comparação ignora a caixa."""
    titulos = [p.text.strip().upper() for p in documento.paragraphs if p.text.strip()]
    resumo = [t for t in titulos if t.startswith("6. QUADRO-RESUMO")]
    assert resumo, f"título da seção 6 não encontrado em {titulos[-6:]}"
    assert titulos.index(resumo[0]) > titulos.index(
        next(t for t in titulos if t.startswith("5. AÇÕES JUDICIAIS"))
    )


def testar_execucao_vazia_nao_vira_linha():
    """_creditor_sections usa [{}] como placeholder quando não há execução; isso não
    pode virar uma linha em branco no quadro."""
    doc = Document(cr._build_checklist_creditos(dict(DADOS, execucoes=[])))
    assert _quadro(doc) is None

    doc_parcial = Document(cr._build_checklist_creditos(
        dict(DADOS, execucoes=[EXECUCOES[0], {}, {"numero": "", "garantia": ""}])
    ))
    assert len(_quadro(doc_parcial).rows) == 2, "só a execução com conteúdo entra"
