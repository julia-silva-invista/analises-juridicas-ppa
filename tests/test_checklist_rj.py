# -*- coding: utf-8 -*-
"""Bateria de testes de fidelidade do Checklist RJ (sem chamar o Gemini de verdade).

Roda com:
    python tests/test_checklist_rj.py
ou:
    pytest tests/test_checklist_rj.py -v
"""
from __future__ import annotations

import copy
import json
import os
import re
import sys
import types as pytypes

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

# Console do Windows às vezes usa cp1252 e não imprime ☑/☐ — evita crash no print.
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(encoding="utf-8", errors="replace")

# ── Stubs de dependências externas (não precisamos de rede/API key) ────────
try:
    import gradio  # noqa: F401
except ImportError:
    sys.modules["gradio"] = pytypes.ModuleType("gradio")

try:
    from google import genai  # noqa: F401
    from google.genai import types as genai_types  # noqa: F401
except ImportError:
    google_mod = pytypes.ModuleType("google")
    genai_mod = pytypes.ModuleType("google.genai")
    genai_types_mod = pytypes.ModuleType("google.genai.types")
    genai_types_mod.GenerateContentConfig = lambda **k: k
    genai_types_mod.Content = lambda **k: k
    genai_types_mod.Part = lambda **k: k
    genai_mod.types = genai_types_mod
    genai_mod.Client = object
    google_mod.genai = genai_mod
    sys.modules["google"] = google_mod
    sys.modules["google.genai"] = genai_mod
    sys.modules["google.genai.types"] = genai_types_mod

if "utils" not in sys.modules:
    utils_mod = pytypes.ModuleType("utils")
    utils_mod._retry = lambda fn, tentativas=3, espera_base=10: fn()
    utils_mod._erro_gemini_permite_failover = lambda exc: False
    sys.modules["utils"] = utils_mod

import checklist_rj as cr  # noqa: E402
import dossie_ppa as dp  # noqa: E402
from docx import Document  # noqa: E402


# ══════════════════════════════════════════════════════════════════════════
# Infraestrutura de teste
# ══════════════════════════════════════════════════════════════════════════

class FakeClient:
    """Simula client.models.generate_content(...).text sem chamar a API."""

    def __init__(self, resposta: str):
        self._resposta = resposta
        self.chamadas = 0
        self.ultima_chamada_kwargs = None
        self.models = self

    def generate_content(self, **kwargs):
        self.chamadas += 1
        self.ultima_chamada_kwargs = kwargs
        resp = pytypes.SimpleNamespace()
        resp.text = self._resposta
        return resp


# Padrões que NUNCA deveriam aparecer no documento final — indicam vazamento
# de instrução de formato (do prompt ou do molde oficial) para o texto visível.
_PADROES_VAZAMENTO = [
    r"\(fls\.\)",
    r"\(referência\)",
    r"\(referencia\)",
    r"\(Mov\.\s*/\s*ID\s*/\s*fls\.\s*/\s*Evento\)",
    r"\[caso haja",
    r"indicar páginas",
    r"indicar paginas",
    r"\[Nome",
    r"\(referência não localizada\)\s*$",  # aceitável isolado, mas nunca em excesso — ver nota abaixo
]
# "(referência não localizada)" é uma saída LEGÍTIMA do prompt (não um vazamento
# de placeholder) — removido da lista de proibidos; mantido comentado para
# referência. Lista efetiva de vazamento:
_PADROES_VAZAMENTO = [
    r"\(fls\.\)",
    r"\(referência\)",
    r"\(referencia\)",
    r"\(Mov\.\s*/\s*ID\s*/\s*fls\.\s*/\s*Evento\)",
    r"\[caso haja",
    r"indicar páginas",
    r"indicar paginas",
    r"\[Nome",
]
_REGEX_VAZAMENTO = re.compile("|".join(_PADROES_VAZAMENTO), re.IGNORECASE)

# Ruído estrutural: marcadores de separação de chunk (rj._rj_merge_textos) que,
# num processo gigante dividido em muitas partes, podem acabar ecoados pelo
# modelo dentro de um campo em vez de filtrados — não é uma instrução de
# formato, mas também não deveria aparecer no documento final.
_REGEX_RUIDO_ESTRUTURAL = re.compile(r"={10,}|PARTE\s+\d+\s*/\s*\d+", re.IGNORECASE)

# Referência "de verdade": rótulo do sistema seguido de um número real ou de
# "s/n" (sem número — comum em processos físicos antigos). Exige \d ou "s/n"
# para não contar como válida uma referência vazia/garbled tipo "(Mov. )" ou
# "(fls. -)" que uma digitalização ruim poderia produzir.
_REGEX_REFERENCIA_REAL = re.compile(
    r"\((?:Mov\.|fls\.|Evento|ID)\s+(?:\d[\d./-]*|s\.?\s*/\s*n\.?)\b", re.IGNORECASE
)


def _iter_celulas(doc: Document):
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                yield ti, ri, ci, cell.text


def achar_vazamentos(docx_path: str):
    doc = Document(docx_path)
    achados = []
    for ti, ri, ci, texto in _iter_celulas(doc):
        if texto and _REGEX_VAZAMENTO.search(texto):
            achados.append((ti, ri, ci, texto.strip()[:120]))
    return achados


def contar_referencias(docx_path: str) -> int:
    doc = Document(docx_path)
    total = 0
    for _, _, _, texto in _iter_celulas(doc):
        if texto and _REGEX_REFERENCIA_REAL.search(texto):
            total += 1
    return total


def achar_ruido_estrutural(docx_path: str):
    doc = Document(docx_path)
    achados = []
    for ti, ri, ci, texto in _iter_celulas(doc):
        if texto and _REGEX_RUIDO_ESTRUTURAL.search(texto):
            achados.append((ti, ri, ci, texto.strip()[:120]))
    return achados


# ══════════════════════════════════════════════════════════════════════════
# Fixtures sintéticas de `dados` (simulando extrações boas e ruins)
# ══════════════════════════════════════════════════════════════════════════

def _base_limpa(rotulo_ref: str) -> dict:
    ref = lambda n: f"({rotulo_ref} {n})"  # noqa: E731
    return {
        "rj_numero": "1234567-89.2024.8.11.0000",
        "vara": f"1ª Vara Empresarial {ref(1)}",
        "data_analise": "31/07/2026",
        "requerentes": f"Empresa Teste Ltda · 12.345.678/0001-90 {ref(10)}",
        "advogados_requerentes": f"Dr. Fulano de Tal {ref(10)}",
        "administrador_judicial": f"AJ Consultoria {ref(15)}",
        "data_pedido": f"01/01/2024 {ref(1)}",
        "data_deferimento": f"15/01/2024 {ref(20)}",
        "consolidacao_substancial": f"Deferido {ref(340)}",
        "periodo_blindagem": f"Ativo {ref(350)}",
        "previsao_encerramento_stay": f"15/01/2025 {ref(350)}",
        "stay_prorrogavel": f"Sim {ref(400)}",
        "recursos_relevantes": [
            {"recurso": f"Agravo de Instrumento nº 111 {ref(50)}", "status": f"Pendente {ref(50)}"},
        ],
        "imoveis_requerentes": [
            {"matricula": f"12.454 {ref(88)}", "cartorio": f"1º RI de Cuiabá {ref(88)}",
             "descricao": f"Fazenda Boa Vista {ref(88)}", "proprietario": f"Empresa Teste {ref(88)}"},
        ],
        "imoveis_essenciais": [
            {"matricula": f"99.001 {ref(90)}", "cartorio": f"2º RI {ref(90)}",
             "descricao": f"Sede da empresa {ref(90)}", "proprietario": f"Empresa Teste {ref(90)}"},
        ],
        "prj_classe_ii": {"desagio": f"30% {ref(200)}", "carencia": f"12 meses {ref(200)}",
                          "parcelas": f"60 {ref(200)}", "juros": f"2% a.a. {ref(200)}",
                          "correcao": f"IPCA {ref(200)}"},
        "prj_classe_iii": {"desagio": f"40% {ref(210)}", "carencia": f"24 meses {ref(210)}",
                           "parcelas": f"72 {ref(210)}", "juros": f"1% a.a. {ref(210)}",
                           "correcao": f"INPC {ref(210)}"},
        "qgc": {"classe_i": f"R$ 500.000,00 {ref(300)}", "classe_ii": f"R$ 2.000.000,00 {ref(300)}",
                "classe_iii": f"R$ 8.000.000,00 {ref(300)}", "classe_iv": f"R$ 100.000,00 {ref(300)}",
                "total": f"R$ 10.600.000,00 {ref(300)}"},
        "agc_situacao": f"Convocada {ref(320)}",
        "agc_1a": f"10/03/2025 {ref(320)}", "agc_2a": f"20/03/2025 {ref(320)}",
        "agc_continuacao": f"27/03/2025 {ref(320)}",
        "documentos_salvos": {
            "peticao_inicial": ref(1).strip("()"),
            "quadro_ativos": ref(88).strip("()"),
            "pericia_previa": "Não consta",
            "laudo_imoveis": ref(90).strip("()"),
            "ultimo_rma": ref(95).strip("()"),
            "qgc_recuperando": ref(300).strip("()"),
            "qgc_aj": ref(305).strip("()"),
            "relatorio_divergencia": "Não consta",
            "prj_aditivos": ref(200).strip("()"),
            "atas_agc": ref(320).strip("()"),
        },
    }


def fixture_processo_longo() -> dict:
    dados = _base_limpa("Mov.")
    dados["imoveis_requerentes"] = [
        {"matricula": f"{10_000 + i} (Mov. {400 + i})", "cartorio": f"{i}º RI (Mov. {400 + i})",
         "descricao": f"Imóvel rural {i} (Mov. {400 + i})", "proprietario": f"Recuperando {i} (Mov. {400 + i})"}
        for i in range(1, 21)
    ]
    dados["recursos_relevantes"] = [
        {"recurso": f"Recurso nº {i} (Mov. {500 + i})", "status": f"Pendente (Mov. {500 + i})"}
        for i in range(1, 8)
    ]
    return dados


def fixture_ocr_ruim() -> dict:
    dados = _base_limpa("fls.")
    dados["requerentes"] = "Empresa   Teste\nLtda .  12.345.678/0001-90   (fls. 10 )"
    dados["administrador_judicial"] = "AJ  Consultoria\n(fls . 15)"
    dados["consolidacao_substancial"] = "Deferido (fls 340)"  # sem ponto — ambíguo de propósito
    dados["previsao_encerramento_stay"] = "15/01/2025(fls.350)"  # sem espaço
    dados["qgc"]["total"] = "R$  10.600.000,00\n(fls. 300)"
    return dados


def fixture_vazio_total() -> dict:
    return {}


def fixture_tipos_malformados() -> dict:
    return {
        "rj_numero": None,
        "requerentes": None,
        "recursos_relevantes": None,
        "imoveis_requerentes": "não é uma lista",
        "imoveis_essenciais": [{"matricula": None, "cartorio": None}],
        "prj_classe_ii": None,
        "qgc": None,
        "documentos_salvos": None,
    }


def fixture_checkbox_ambiguo() -> dict:
    dados = _base_limpa("Mov.")
    dados["consolidacao_substancial"] = "Deferido e Indeferido em parte (Mov. 12)"
    dados["periodo_blindagem"] = "Favorável"  # não é opção válida do campo — deve ficar tudo desmarcado
    return dados


def fixture_processo_gigante() -> dict:
    """Processo enorme: muitos imóveis/recursos — checa volume e desempenho."""
    dados = _base_limpa("Mov.")
    dados["imoveis_requerentes"] = [
        {"matricula": f"{50_000 + i} (Mov. {4000 + i})", "cartorio": f"{(i % 12) + 1}º RI (Mov. {4000 + i})",
         "descricao": f"Fazenda / gleba rural nº {i}, com benfeitorias (Mov. {4000 + i})",
         "proprietario": f"Recuperando {i} Participações S.A. (Mov. {4000 + i})"}
        for i in range(1, 121)
    ]
    dados["imoveis_essenciais"] = [
        {"matricula": f"{90_000 + i} (Mov. {5000 + i})", "cartorio": f"{(i % 5) + 1}º RI (Mov. {5000 + i})",
         "descricao": f"Unidade industrial {i} (Mov. {5000 + i})",
         "proprietario": f"Recuperando {i} Participações S.A. (Mov. {5000 + i})"}
        for i in range(1, 31)
    ]
    dados["recursos_relevantes"] = [
        {"recurso": f"Agravo de Instrumento nº {i} (Mov. {6000 + i})", "status": f"Pendente de julgamento (Mov. {6000 + i})"}
        for i in range(1, 41)
    ]
    return dados


def fixture_mal_digitalizado_antigo() -> dict:
    """Processo físico antigo, digitalização ruim: texto quebrado, refs sem número,
    moeda em formato antigo, campo enorme (parágrafo inteiro colado), muitos "Não consta"."""
    dados = _base_limpa("fls.")
    texto_colado = (
        "Em cumprimento ao despacho de fls. 12, certifico e dou fé que procedi a  ju nta da"
        " da  do cumen taçao acostada aos autos, referen te ao pedido de recuperaçã0 judicial"
        " formulado pel a empresa r equer ente, cujos t ermos passo a transcr ever na integra "
        "para os devidos fins de direi to, ficando consig nado que a numeraçao das folhas pode"
        " apresentar falhas em razão do estado de conservação dos autos fisicos digitalizados "
        "neste ato (fls. s/n).  " * 8
    )
    dados["requerentes"] = f"Empresa Antiga Textil Ltda (em Recuperação) · CNPJ ilegível (fls. s/n)"
    dados["advogados_requerentes"] = "Não consta"
    dados["administrador_judicial"] = texto_colado
    dados["data_pedido"] = "12/03/2009 (fls. s/n)"
    dados["data_deferimento"] = "Não consta"
    dados["consolidacao_substancial"] = "Não consta"
    dados["periodo_blindagem"] = "Não consta"
    dados["previsao_encerramento_stay"] = "Não consta"
    dados["stay_prorrogavel"] = ""
    dados["agc_situacao"] = "Não consta"
    dados["agc_1a"] = "Não consta"
    dados["agc_2a"] = "Não consta"
    dados["agc_continuacao"] = "Não consta"
    dados["qgc"] = {
        "classe_i": "Cr$ 8.000.000,00 (fls. s/n)", "classe_ii": "Não consta",
        "classe_iii": "Cr$ 45.000.000,00 (fls. 340/342)", "classe_iv": "Não consta",
        "total": "Não consta",
    }
    dados["imoveis_requerentes"] = [
        {"matricula": "12.454-B (fls. 88/90)", "cartorio": "1º  R I  de   Cuiabá (fls. 88)",
         "descricao": "Ga lp ão indus trial (fls. 88)", "proprietario": "Emp resa Antiga Textil (fls. 88)"},
    ]
    dados["documentos_salvos"] = {
        "peticao_inicial": "s/n",
        "quadro_ativos": "Não consta",
        "pericia_previa": "Não consta",
        "laudo_imoveis": "Não consta",
        "ultimo_rma": "Não consta",
        "qgc_recuperando": "340/342",
        "qgc_aj": "Não consta",
        "relatorio_divergencia": "Não consta",
        "prj_aditivos": "Não consta",
        "atas_agc": "Não consta",
    }
    return dados


def fixture_referencias_quebradas() -> dict:
    """Referências presentes mas vazias/garbled — número perdido na digitalização.
    Não é um vazamento de placeholder, mas também não é uma referência válida:
    serve pra medir a precisão do próprio checador de referências."""
    dados = _base_limpa("Mov.")
    dados["requerentes"] = "Empresa Teste Ltda (Mov. )"
    dados["consolidacao_substancial"] = "Deferido (fls. -)"
    dados["qgc"]["total"] = "R$ 10.600.000,00 (Evento)"
    return dados


def fixture_ruido_estrutural() -> dict:
    """Simula um processo gigante dividido em muitas partes (rj._rj_merge_textos)
    cujo separador de chunk vazou para dentro de um campo extraído."""
    dados = _base_limpa("Mov.")
    dados["administrador_judicial"] = (
        "AJ Consultoria (Mov. 15)\n============================================================\n"
        "PARTE 7/12\n============================================================\n"
        "continuação da atuação do AJ (Mov. 16)"
    )
    return dados


def fixture_vazamento_proposital() -> dict:
    dados = _base_limpa("Mov.")
    dados["requerentes"] = "Empresa Teste Ltda (fls.)"
    dados["consolidacao_substancial"] = "Deferido (referência)"
    dados["qgc"]["total"] = "R$ 10.600.000,00 (Mov./ID/fls./Evento)"
    dados["documentos_salvos"]["peticao_inicial"] = "[caso haja, indicar páginas]"
    return dados


def fixture_agc_vazio() -> dict:
    """Nenhuma informação de AGC — deve cair no fallback "Sem datas designadas",
    nunca em "Não consta" (regra específica do AGC)."""
    dados = _base_limpa("Mov.")
    dados["agc_situacao"] = "Não consta"
    dados["agc_1a"] = "Não consta"
    dados["agc_2a"] = "Não consta"
    dados["agc_continuacao"] = "Não consta"
    return dados


def fixture_formato_citacao_novo() -> dict:
    """Novo formato de referência (identificador + página juntos, com sufixo de qual pdf
    quando há mais de um) — ver REGRA_CITACAO_PADRAO em dossie_ppa.py. Confirma que o
    pipeline de renderização aceita o novo formato sem gerar vazamento nem falso-negativo
    na contagem de referências."""
    dados = _base_limpa("ID")
    dados["consolidacao_substancial"] = "Deferido (ID 188753786 | fl. 135)"
    dados["periodo_blindagem"] = "Ativo (ID 188753786 | fl. 135 do pdf 0001259-66.1996.8.11.0041)"
    return dados


FIXTURES = {
    "clean_pje": lambda: _base_limpa("Mov."),
    "clean_esaj": lambda: _base_limpa("fls."),
    "clean_eproc": lambda: _base_limpa("Evento"),
    "clean_projudi": lambda: _base_limpa("ID"),
    "processo_longo": fixture_processo_longo,
    "processo_gigante_120imoveis": fixture_processo_gigante,
    "ocr_ruim": fixture_ocr_ruim,
    "mal_digitalizado_antigo": fixture_mal_digitalizado_antigo,
    "referencias_quebradas": fixture_referencias_quebradas,
    "ruido_estrutural_chunk": fixture_ruido_estrutural,
    "vazio_total": fixture_vazio_total,
    "tipos_malformados": fixture_tipos_malformados,
    "checkbox_ambiguo": fixture_checkbox_ambiguo,
    "vazamento_proposital": fixture_vazamento_proposital,
    "agc_vazio": fixture_agc_vazio,
    "formato_citacao_novo": fixture_formato_citacao_novo,
}

# fixtures que DEVEM sair 100% limpas (zero vazamento de placeholder)
FIXTURES_ESPERAM_ZERO_VAZAMENTO = {
    "clean_pje", "clean_esaj", "clean_eproc", "clean_projudi",
    "processo_longo", "processo_gigante_120imoveis", "agc_vazio",
    "ocr_ruim", "mal_digitalizado_antigo", "referencias_quebradas",
    "ruido_estrutural_chunk", "vazio_total", "tipos_malformados", "checkbox_ambiguo",
    "formato_citacao_novo",
}
# fixture que DEVE ter vazamento detectado (prova que o checador funciona)
FIXTURES_ESPERAM_VAZAMENTO = {"vazamento_proposital"}
# fixture que DEVE ter ruído estrutural detectado (prova que esse checador funciona)
FIXTURES_ESPERAM_RUIDO_ESTRUTURAL = {"ruido_estrutural_chunk"}
# limite de tempo (segundos) pra geração do documento — pega lentidão patológica
LIMITE_SEGUNDOS_GERACAO = 8.0


# ══════════════════════════════════════════════════════════════════════════
# Testes de _cb (lógica de checkbox)
# ══════════════════════════════════════════════════════════════════════════

CASOS_CB = [
    (["Ativo", "Inativo"], "Ativo", "☑ Ativo"),
    (["Ativo", "Inativo"], "Inativo", "☑ Inativo"),
    (["Ativo", "Inativo"], "", None),  # nenhuma marcada
    (["Favorável", "Desfavorável", "Pendente"], "Desfavorável", "☑ Desfavorável"),
    (["Favorável", "Desfavorável", "Pendente"], "Favorável", "☑ Favorável"),
    (["Deferido", "Indeferido"], "Deferido (Mov. 12)", "☑ Deferido"),
    (["Deferido", "Indeferido"], "Indeferido (fls. 88)", "☑ Indeferido"),
    (["Sim", "Não"], "Não consta", None),
    (["Ativo", "Inativo"], "Deferido e Indeferido em parte (Mov. 12)", None),
    ([], "qualquer coisa", None),
]


def testar_cb():
    falhas = []
    for options, selected, esperado_marcado in CASOS_CB:
        resultado = cr._cb(options, selected)
        marcados = re.findall(r"☑ ([^☐☑]+?)(?:\s{2,}|$)", resultado)
        marcados = [m.strip() for m in marcados]
        if esperado_marcado is None:
            if marcados:
                falhas.append(f"_cb({options!r}, {selected!r}) marcou {marcados}, esperava nenhuma")
        else:
            esperado_texto = esperado_marcado.replace("☑ ", "").strip()
            if marcados != [esperado_texto]:
                falhas.append(f"_cb({options!r}, {selected!r}) => {marcados!r}, esperava [{esperado_texto!r}]")
    return falhas


# ══════════════════════════════════════════════════════════════════════════
# Testes de _extrair (parsing de resposta do "Gemini")
# ══════════════════════════════════════════════════════════════════════════

CASOS_EXTRAIR = [
    ("json_limpo", '{"rj_numero": "123"}', {"rj_numero": "123"}),
    ("com_cercas_markdown", '```json\n{"rj_numero": "123"}\n```', {"rj_numero": "123"}),
    ("com_cercas_sem_lang", '```\n{"rj_numero": "123"}\n```', {"rj_numero": "123"}),
    ("prosa_antes_e_depois", 'Aqui esta o JSON:\n{"rj_numero": "123"}\nFim.', "FALLBACK_VAZIO"),
    ("truncado", '{"rj_numero": "123", "vara":', "FALLBACK_VAZIO"),
    ("vazio", "", "FALLBACK_VAZIO"),
    ("nao_e_objeto", "[1, 2, 3]", "FALLBACK_VAZIO"),
    ("com_mojibake", '{"requerentes": "Empresa Teste �� Ltda (fls. 10)"}',
     {"requerentes": "Empresa Teste �� Ltda (fls. 10)"}),
    ("com_ruido_estrutural_no_valor",
     '{"administrador_judicial": "AJ (Mov. 15)\\n====\\nPARTE 2/5\\n====\\ncontinua (Mov. 16)"}',
     {"administrador_judicial": "AJ (Mov. 15)\n====\nPARTE 2/5\n====\ncontinua (Mov. 16)"}),
]


def testar_extrair():
    falhas = []
    for nome, resposta, esperado in CASOS_EXTRAIR:
        client = FakeClient(resposta)
        try:
            resultado = cr._extrair("PROMPT ", "TEXTO", client, "modelo-fake")
        except Exception as e:  # noqa: BLE001
            falhas.append(f"_extrair[{nome}] levantou exceção: {e!r}")
            continue
        if esperado == "FALLBACK_VAZIO":
            if resultado != {}:
                falhas.append(f"_extrair[{nome}] esperava fallback {{}}, veio {resultado!r}")
        else:
            if resultado != esperado:
                falhas.append(f"_extrair[{nome}] esperava {esperado!r}, veio {resultado!r}")
    return falhas


def testar_agc_fallback():
    """"Não consta"/vazio em agc_situacao deve virar "Sem datas designadas"
    marcado — nunca "Não consta" nem nenhuma opção marcada."""
    falhas = []
    for entrada in ("", "Não consta", "não consta ", None):
        dados = copy.deepcopy(fixture_agc_vazio())
        dados["agc_situacao"] = entrada
        caminho = cr._build_checklist_rj(dados)
        doc = Document(caminho)
        achou = any(
            "☑ Sem datas designadas" in texto
            for _, _, _, texto in _iter_celulas(doc)
        )
        if not achou:
            falhas.append(f"agc_situacao={entrada!r} não marcou 'Sem datas designadas'")
    return falhas


def testar_montar_fonte_rj():
    """Relatório é a fonte principal; texto bruto só complementa."""
    falhas = []
    fonte = cr._montar_fonte_rj("RELATORIO XYZ", "TEXTO BRUTO ABC")
    if "RELATORIO XYZ" not in fonte:
        falhas.append("_montar_fonte_rj não incluiu o relatório")
    if "TEXTO BRUTO ABC" not in fonte:
        falhas.append("_montar_fonte_rj não incluiu o texto bruto complementar")
    if fonte.index("RELATORIO XYZ") > fonte.index("TEXTO BRUTO ABC"):
        falhas.append("_montar_fonte_rj não colocou o relatório antes do texto bruto")
    # sem texto bruto — só o relatório, sem quebrar
    so_relatorio = cr._montar_fonte_rj("RELATORIO XYZ", "")
    if "RELATORIO XYZ" not in so_relatorio:
        falhas.append("_montar_fonte_rj sem texto bruto perdeu o relatório")
    return falhas


def testar_completude_dossie():
    """Relatório deve ser fonte principal e poder ampliar um campo não vazio."""
    falhas = []
    curto = "Não especificado expressamente no contrato original (ID 47118237 | fls. 34/35)"
    completo = (
        "Não especificado expressamente no contrato original (posteriormente definido como INPC "
        "mediante acórdão do TJMT) (ID 47118237 | fls. 34/35; ID 47119780 | fls. 314, 326)"
    )
    relatorio = "Correção monetária: " + completo
    bruto = "Contrato original: " + curto

    material = dp._montar_material_prioritario(relatorio, bruto)
    if material.index(relatorio) > material.index(bruto):
        falhas.append("material do dossiê não colocou o relatório antes do texto bruto")
    if "FONTE PRINCIPAL; PRESERVAR TEXTO COMPLETO" not in material:
        falhas.append("material do dossiê não identifica o relatório como fonte principal")

    base = {"creditos": [{"ind_cm": curto}]}
    candidato = {"creditos": [{"ind_cm": completo}]}
    client = FakeClient(json.dumps(candidato, ensure_ascii=False))
    mesclado = dp._completar_com_relatorio(base, relatorio, client, "modelo-fake")
    if mesclado.get("creditos", [{}])[0].get("ind_cm") != completo:
        falhas.append("segunda passada não substituiu a versão curta pela completa")

    preservado = dp._mesclar_mais_completo(candidato, base)
    if preservado.get("creditos", [{}])[0].get("ind_cm") != completo:
        falhas.append("mesclagem permitiu que uma versão curta apagasse a completa")

    prompt = client.ultima_chamada_kwargs["contents"][0]["parts"][0]["text"]
    if "sem resumir ou parafrasear" not in prompt:
        falhas.append("prompt de completude não proíbe resumir/parafrasear")
    return falhas


def _paragrafos_doc(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _falhas_italico_doc(doc, contexto):
    falhas = []
    refs_encontradas = 0
    for paragraph in _paragrafos_doc(doc):
        for run in paragraph.runs:
            if dp._RE_MARCADOR_REFERENCIA.search(run.text or ""):
                refs_encontradas += 1
                if run.italic is not True:
                    falhas.append(f"{contexto}: referência sem itálico: {run.text!r}")
    if not refs_encontradas:
        falhas.append(f"{contexto}: nenhuma referência encontrada para validar")
    return falhas


def testar_referencias_em_italico():
    falhas = []
    texto = (
        "Não especificado expressamente no contrato original (posteriormente definido como INPC) "
        "(ID 188753786 | fl. 1)"
    )

    doc = Document()
    cell_write = doc.add_table(rows=1, cols=1).cell(0, 0)
    dp._write(cell_write, texto)
    falhas.extend(_falhas_italico_doc(doc, "helper compartilhado dos checklists"))
    if any(run.italic is True for run in cell_write.paragraphs[0].runs if "posteriormente" in run.text):
        falhas.append("parêntese explicativo foi colocado em itálico como se fosse referência")

    doc_template = Document()
    cell_template = doc_template.add_table(rows=1, cols=1).cell(0, 0)
    dp._substituir_texto_celula(cell_template, texto)
    falhas.extend(_falhas_italico_doc(doc_template, "helper do template do dossiê"))

    dados_rj = copy.deepcopy(fixture_formato_citacao_novo())
    caminho_rj = cr._build_checklist_rj(dados_rj)
    falhas.extend(_falhas_italico_doc(Document(caminho_rj), "Checklist RJ"))

    dados_credito = {
        "rj_numero": "0000000-00.0000.0.00.0000",
        "credor": "Banco Teste",
        "lastros": [{"cedula": f"CCB nº 1 {texto}"}],
        "garantias": [],
        "execucoes": [],
    }
    caminho_credito = cr._build_checklist_creditos(dados_credito)
    falhas.extend(_falhas_italico_doc(Document(caminho_credito), "Checklist de Créditos"))
    return falhas


def testar_extrair_fonte_gigante():
    """Processo gigante: fonte >900k chars — confirma que o truncamento
    (checklist_rj.py: fonte[:900_000]) funciona e não estoura memória/exceção."""
    falhas = []
    prompt_base = "PROMPT_BASE\n"
    fonte_gigante = "A" * 1_000_000
    client = FakeClient('{"rj_numero": "OK"}')
    try:
        resultado = cr._extrair(prompt_base, fonte_gigante, client, "modelo-fake")
    except Exception as e:  # noqa: BLE001
        return [f"_extrair[fonte_gigante] levantou exceção: {e!r}"]
    if resultado != {"rj_numero": "OK"}:
        falhas.append(f"_extrair[fonte_gigante] resultado inesperado: {resultado!r}")
    enviado = client.ultima_chamada_kwargs["contents"][0]["parts"][0]["text"]
    esperado_len = len(prompt_base) + 900_000
    if len(enviado) != esperado_len:
        falhas.append(
            f"_extrair[fonte_gigante] enviou {len(enviado)} chars ao modelo, "
            f"esperava exatamente {esperado_len} (prompt_base + fonte truncada em 900_000)"
        )
    return falhas


# ══════════════════════════════════════════════════════════════════════════
# Execução principal
# ══════════════════════════════════════════════════════════════════════════

def rodar_bateria():
    import time

    resultados = []
    falhas_totais = []

    falhas_cb = testar_cb()
    resultados.append(("_cb (lógica de checkbox)", len(CASOS_CB), len(falhas_cb)))
    falhas_totais.extend(f"[_cb] {f}" for f in falhas_cb)

    falhas_extrair = testar_extrair()
    resultados.append(("_extrair (parsing da resposta)", len(CASOS_EXTRAIR), len(falhas_extrair)))
    falhas_totais.extend(f"[_extrair] {f}" for f in falhas_extrair)

    falhas_truncamento = testar_extrair_fonte_gigante()
    resultados.append(("_extrair (truncamento fonte gigante)", 1, len(falhas_truncamento)))
    falhas_totais.extend(f"[_extrair] {f}" for f in falhas_truncamento)

    falhas_fonte = testar_montar_fonte_rj()
    resultados.append(("_montar_fonte_rj (relatório > texto bruto)", 1, len(falhas_fonte)))
    falhas_totais.extend(f"[_montar_fonte_rj] {f}" for f in falhas_fonte)

    falhas_completude = testar_completude_dossie()
    resultados.append(("completude relatório > dossiê", 4, len(falhas_completude)))
    falhas_totais.extend(f"[completude] {f}" for f in falhas_completude)

    falhas_italico = testar_referencias_em_italico()
    resultados.append(("referências em itálico nos DOCX", 4, len(falhas_italico)))
    falhas_totais.extend(f"[itálico] {f}" for f in falhas_italico)

    falhas_agc = testar_agc_fallback()
    resultados.append(("AGC fallback (Sem datas designadas)", 4, len(falhas_agc)))
    falhas_totais.extend(f"[agc_fallback] {f}" for f in falhas_agc)

    print("=" * 78)
    print("BATERIA DE FIDELIDADE — CHECKLIST RJ")
    print("=" * 78)

    for nome, fabrica in FIXTURES.items():
        dados = copy.deepcopy(fabrica())
        erro_build = None
        vazamentos = []
        ruidos = []
        n_refs = 0
        inicio = time.perf_counter()
        try:
            caminho = cr._build_checklist_rj(dados)
            duracao = time.perf_counter() - inicio
            vazamentos = achar_vazamentos(caminho)
            ruidos = achar_ruido_estrutural(caminho)
            n_refs = contar_referencias(caminho)
        except Exception as e:  # noqa: BLE001
            duracao = time.perf_counter() - inicio
            erro_build = repr(e)

        status = "OK"
        if erro_build:
            status = "ERRO"
            falhas_totais.append(f"[{nome}] _build_checklist_rj levantou {erro_build}")
        elif nome in FIXTURES_ESPERAM_ZERO_VAZAMENTO and vazamentos:
            status = "FALHA (vazamento inesperado)"
            falhas_totais.append(f"[{nome}] vazamentos inesperados: {vazamentos}")
        elif nome in FIXTURES_ESPERAM_VAZAMENTO and not vazamentos:
            status = "FALHA (checador não detectou vazamento esperado)"
            falhas_totais.append(f"[{nome}] esperava vazamento detectado, não achou nenhum")
        elif nome in FIXTURES_ESPERAM_RUIDO_ESTRUTURAL and not ruidos:
            status = "FALHA (checador não detectou ruído estrutural esperado)"
            falhas_totais.append(f"[{nome}] esperava ruído estrutural detectado, não achou nenhum")
        elif nome not in FIXTURES_ESPERAM_RUIDO_ESTRUTURAL and ruidos:
            status = "FALHA (ruído estrutural inesperado)"
            falhas_totais.append(f"[{nome}] ruído estrutural inesperado: {ruidos}")
        elif not erro_build and duracao > LIMITE_SEGUNDOS_GERACAO:
            status = "FALHA (lento)"
            falhas_totais.append(f"[{nome}] geração levou {duracao:.2f}s (limite {LIMITE_SEGUNDOS_GERACAO}s)")

        print(f"{nome:38s} status={status:38s} refs={n_refs:3d} vazam={len(vazamentos)} ruido={len(ruidos)} {duracao:.2f}s")
        if vazamentos and nome in FIXTURES_ESPERAM_VAZAMENTO:
            for v in vazamentos:
                print(f"    (esperado) tabela={v[0]} linha={v[1]} col={v[2]}: {v[3]!r}")
        elif vazamentos:
            for v in vazamentos:
                print(f"    !! tabela={v[0]} linha={v[1]} col={v[2]}: {v[3]!r}")
        if ruidos and nome in FIXTURES_ESPERAM_RUIDO_ESTRUTURAL:
            for r in ruidos:
                print(f"    (esperado) ruído tabela={r[0]} linha={r[1]} col={r[2]}: {r[3]!r}")
        elif ruidos:
            for r in ruidos:
                print(f"    !! ruído tabela={r[0]} linha={r[1]} col={r[2]}: {r[3]!r}")

    print("-" * 78)
    for nome, total, n_falhas in resultados:
        print(f"{nome:38s} {total - n_falhas}/{total} OK")

    print("=" * 78)
    if falhas_totais:
        print(f"RESULTADO: {len(falhas_totais)} problema(s) encontrado(s):")
        for f in falhas_totais:
            print(f"  - {f}")
    else:
        print("RESULTADO: nenhum problema encontrado.")
    print("=" * 78)
    return falhas_totais


if __name__ == "__main__":
    falhas = rodar_bateria()
    sys.exit(1 if falhas else 0)


# ── Wrappers pytest (opcional) ──────────────────────────────────────────────
def test_bateria_completa():
    falhas = rodar_bateria()
    assert not falhas, "\n".join(falhas)
