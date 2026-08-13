# -*- coding: utf-8 -*-
"""Extração, edição e exportação da Timeline Societária."""

from __future__ import annotations

import base64
import binascii
import copy
import html
import json
import os
import re
import tempfile
import time
import unicodedata
from pathlib import Path
from typing import Any

import gradio as gr
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from google import genai
from google.genai import types

from utils import (
    _codigo_http_gemini,
    _detalhe_erro_gemini,
    _erro_gemini_e_teto_de_gasto,
    _executar_com_failover_gemini,
    _get_gemini_clients,
    _retry,
)


MODEL_TIMELINE = os.getenv("GEMINI_MODEL_TIMELINE", "gemini-2.5-pro")
TEMPLATE_PATH = Path(__file__).parent / "templates" / "cronologia_societaria.docx"

PROMPT_TIMELINE = """\
Você é um analista jurídico societário sênior. Analise TODOS os atos societários
anexados, organize-os cronologicamente e reconstrua o estado da sociedade após
cada ato. Aplique OCR visual quando o PDF estiver escaneado.

REGRAS CRÍTICAS
- Não invente. Quando o documento não informar algo, use string vazia ou lista vazia.
- Diferencie data de assinatura, data do ato e data de arquivamento. Em "data",
  priorize a data do ato; se ausente, use a data do arquivamento.
- "ato" deve trazer a denominação legível: Constituição, 1ª Alteração,
  5ª Alteração/ACS etc. "numero_arquivamento" é o número da Junta Comercial.
- Em "socios_apos", represente o quadro societário resultante DEPOIS do ato.
  Informe percentual quando constar ou puder ser calculado exatamente pelas quotas.
- Em "administradores_apos", represente a administração resultante depois do ato.
- Em "capital_social_apos" e "capital_social_anterior", use o valor em reais com
  formatação brasileira. Preencha o anterior sempre que o documento informar.
- Em "sede_apos", informe cidade/UF da sede depois do ato. Em "objeto_apos",
  resuma o objeto social vigente depois do ato (uma linha).
- Em "cessoes", registre CADA transmissão de quotas: quem cedeu ("cedente"),
  quem recebeu ("cessionario"), o percentual ou número de quotas cedidas
  ("participacao"), o valor da cessão ("valor") e observações relevantes
  ("observacao") — por exemplo valor simbólico, cessão a parente ou a pessoa
  jurídica do mesmo grupo, doação, permuta.
- Em "imoveis", registre CADA bem imóvel movimentado no ato, com "matricula"
  (número da matrícula), "cartorio" (registro de imóveis), "cidade", "valor" e
  "movimento" preenchido com uma destas palavras: "integralizacao" (conferido ao
  capital), "saida" (restituído ao sócio, cindido ou transferido), "aquisicao"
  (comprado pela sociedade). Em "descricao", detalhe a operação.
- Em "filiais_apos", liste todas as filiais existentes depois do ato; em
  "filiais_adicionadas", somente as criadas naquele ato.
- Em "detalhamento", descreva objetivamente todas as mudanças relevantes:
  entrada/saída de sócio, cessão de quotas, integralização, capital, administração,
  objeto, sede, nome empresarial, filiais, dissolução ou outras deliberações.
- Cite a fonte em linguagem humana: nome do documento e página(s).
- Ordene os eventos em ordem cronológica crescente.

Responda SOMENTE com JSON válido nesta estrutura:
{
  "empresa": "",
  "cnpj": "",
  "eventos": [
    {
      "data": "DD/MM/AAAA",
      "ato": "",
      "numero_arquivamento": "",
      "detalhamento": "",
      "categorias": ["Sócios", "Capital", "Administração", "Filial", "Imóvel", "Sede", "Objeto", "Outros"],
      "socios_apos": [{"nome": "", "participacao": "", "quotas": ""}],
      "administradores_apos": [{"nome": "", "cargo": "Administrador"}],
      "capital_social_anterior": "",
      "capital_social_apos": "",
      "sede_apos": "",
      "objeto_apos": "",
      "cessoes": [{"cedente": "", "cessionario": "", "participacao": "", "valor": "", "observacao": ""}],
      "imoveis": [{"matricula": "", "cartorio": "", "cidade": "", "valor": "", "movimento": "", "descricao": ""}],
      "filiais_apos": [{"nome": "", "local": ""}],
      "filiais_adicionadas": [{"nome": "", "local": ""}],
      "fonte": ""
    }
  ]
}
"""


def _resumo_erro_gemini(exc: Exception) -> str:
    """Uma linha por credencial: código + o que a API respondeu, truncado."""
    codigo = _codigo_http_gemini(exc)
    detalhe = _detalhe_erro_gemini(exc)
    if len(detalhe) > 220:
        detalhe = detalhe[:220].rstrip() + "…"
    return f"{codigo} — {detalhe}" if codigo else detalhe


def _modelos_disponiveis(client: genai.Client) -> list[str]:
    """Modelos com generateContent nesta credencial. Só é chamado para explicar um 404:
    dizer "o modelo não existe" sem conferir já mandou a usuária mexer na variável errada."""
    try:
        return sorted(
            m.name.replace("models/", "")
            for m in client.models.list()
            if "generateContent" in (getattr(m, "supported_actions", None) or [])
        )
    except Exception:
        return []


def _mensagem_erro_gemini(exc: Exception, clients: list | None = None) -> str:
    """Traduz a falha da API para uma mensagem acionável (sem expor a chave).

    Sempre carrega o texto cru que a API devolveu: classificar por conta própria e
    esconder o original já produziu diagnóstico errado (um 404 virou "ajuste
    GEMINI_MODEL_TIMELINE" quando o modelo existia na chave).
    """
    codigo = _codigo_http_gemini(exc)
    falhas = list(getattr(exc, "falhas_por_chave", None) or [])
    rodape = ("\n\nO que cada credencial respondeu:\n" + "\n".join(f"  • {f}" for f in falhas)) if falhas else ""

    if _erro_gemini_e_teto_de_gasto(exc):
        return (
            "A API Gemini recusou a análise: o projeto atingiu o teto de gasto mensal "
            "(429 RESOURCE_EXHAUSTED). Nenhuma chave configurada tem cota disponível. "
            "Eleve o limite em https://ai.studio/spend ou cadastre outra chave "
            "(GEMINI_API_KEY_2, _3, ...) nos secrets do Space." + rodape
        )
    if codigo == 429:
        return (
            "A API Gemini está sem cota no momento (429). Tente novamente em alguns "
            "minutos ou cadastre outra chave (GEMINI_API_KEY_2, _3, ...) nos secrets do Space." + rodape
        )
    if codigo in (401, 403):
        return (
            "A chave da API Gemini foi recusada (autenticação/permissão). Revise os secrets "
            f"do Space. Resposta da API: {_resumo_erro_gemini(exc)}" + rodape
        )
    if codigo == 404:
        modelos = _modelos_disponiveis(clients[-1]) if clients else []
        if modelos and MODEL_TIMELINE not in modelos:
            proximos = ", ".join(m for m in modelos if m.startswith("gemini-")) or ", ".join(modelos)
            return (
                f"O modelo '{MODEL_TIMELINE}' não existe nesta chave. Ajuste "
                f"GEMINI_MODEL_TIMELINE nas variáveis do Space para um destes: {proximos}." + rodape
            )
        disponivel = " (o modelo existe nesta chave, então o 404 veio de outra coisa)" if modelos else ""
        return f"A API Gemini respondeu 404{disponivel}: {_resumo_erro_gemini(exc)}" + rodape
    return f"Falha ao analisar os atos societários: {_resumo_erro_gemini(exc)}" + rodape


def _file_path(value: Any) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        return value.get("path") or value.get("name") or ""
    return getattr(value, "path", None) or getattr(value, "name", None) or ""


def _clean_json(raw: str) -> dict:
    raw = re.sub(r"^```[a-z]*\s*", "", (raw or "").strip(), flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    data = json.loads(raw)
    if not isinstance(data, dict):
        raise ValueError("A resposta da IA não contém um objeto JSON.")
    data.setdefault("empresa", "")
    data.setdefault("cnpj", "")
    data.setdefault("eventos", [])
    data["eventos"] = sorted(
        [e for e in data["eventos"] if isinstance(e, dict)],
        key=lambda e: _date_key(e.get("data", "")),
    )
    return data


def _date_key(value: str) -> tuple[int, int, int]:
    m = re.search(r"(\d{1,2})/(\d{1,2})/(\d{4})", str(value))
    if not m:
        return (9999, 12, 31)
    return (int(m.group(3)), int(m.group(2)), int(m.group(1)))


def _upload_pdf(client: genai.Client, path: str):
    upload_path = path
    tmp_ascii = None
    try:
        path.encode("ascii")
    except UnicodeEncodeError:
        import shutil

        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.close()
        shutil.copy2(path, tmp.name)
        upload_path = tmp_ascii = tmp.name

    uploaded = client.files.upload(file=upload_path)
    waited = 0
    while waited < 120:
        state = getattr(getattr(uploaded, "state", None), "name", None)
        if state in ("ACTIVE", "FAILED"):
            break
        time.sleep(2)
        waited += 2
        uploaded = client.files.get(name=uploaded.name)
    if tmp_ascii:
        try:
            os.remove(tmp_ascii)
        except OSError:
            pass
    if getattr(getattr(uploaded, "state", None), "name", None) == "FAILED":
        raise RuntimeError(f"Falha ao enviar {Path(path).name} para análise.")
    return uploaded


def _extract_com_cliente(client: genai.Client, paths: list[str]) -> dict:
    """Upload + geração na MESMA credencial — o arquivo enviado só existe no projeto
    daquela chave, então uma troca de chave no meio invalidaria os file_uri."""
    uploaded = []
    try:
        parts: list[types.Part] = [types.Part(text=PROMPT_TIMELINE)]
        for index, path in enumerate(paths, 1):
            item = _upload_pdf(client, path)
            uploaded.append(item)
            parts.append(types.Part(text=f"\nDOCUMENTO {index}: {Path(path).name}\n"))
            parts.append(
                types.Part(
                    file_data=types.FileData(
                        file_uri=item.uri,
                        mime_type=getattr(item, "mime_type", None) or "application/pdf",
                    )
                )
            )

        config = types.GenerateContentConfig(response_mime_type="application/json")

        def _call():
            return client.models.generate_content(
                model=MODEL_TIMELINE,
                contents=[types.Content(role="user", parts=parts)],
                config=config,
            ).text

        return _clean_json(_retry(_call, tentativas=3, espera_base=8))
    finally:
        for item in uploaded:
            try:
                client.files.delete(name=item.name)
            except Exception:
                pass


def _extract(files: list[Any]) -> dict:
    paths = [_file_path(f) for f in (files or [])]
    paths = [p for p in paths if p and Path(p).exists()]
    if not paths:
        raise ValueError("Envie ao menos um ato societário em PDF.")

    # Antes a timeline usava só a GEMINI_API_KEY_1 — quando o projeto dela estourava o teto
    # de gasto, a aba inteira caía com traceback. Agora percorre todas as chaves cadastradas,
    # como já fazem Processos e RJ.
    clients = _get_gemini_clients()

    # Só o erro da ÚLTIMA credencial propaga; sem registrar os anteriores, uma chave
    # inválida no começo da fila fica invisível e o diagnóstico sai pela metade.
    anteriores: list[str] = []

    def _registrar(indice: int, _proximo: int, exc: Exception) -> None:
        anteriores.append(f"chave {indice + 1}: {_resumo_erro_gemini(exc)}")

    try:
        return _executar_com_failover_gemini(
            clients,
            lambda client, _i: _extract_com_cliente(client, paths),
            ao_falhar=_registrar,
        )
    except Exception as exc:
        # Viajam junto com o erro para a mensagem: as falhas de cada credencial e os
        # próprios clients, usados para conferir o catálogo de modelos num 404.
        exc.falhas_por_chave = anteriores + [f"chave {len(clients)}: {_resumo_erro_gemini(exc)}"]
        exc.clients_gemini = clients
        raise


def _joined_people(items: list[dict], include_share: bool = False) -> str:
    values = []
    for item in items or []:
        name = str(item.get("nome", "")).strip()
        if not name:
            continue
        if include_share:
            share = str(item.get("participacao", "") or item.get("quotas", "")).strip()
            values.append(f"{name} | {share}" if share else name)
        else:
            role = str(item.get("cargo", "")).strip()
            values.append(f"{name} ({role})" if role else name)
    return "; ".join(values)


def _joined_filiais(items: list[dict]) -> str:
    return "; ".join(
        " — ".join(x for x in [str(i.get("nome", "")).strip(), str(i.get("local", "")).strip()] if x)
        for i in (items or [])
        if i.get("nome") or i.get("local")
    )





def _share(socio: dict) -> str:
    return str(socio.get("participacao", "") or socio.get("quotas", "")).strip()


def _socios_map(event: dict | None) -> dict:
    result = {}
    for socio in ((event or {}).get("socios_apos") or []):
        name = str(socio.get("nome", "")).strip()
        if name:
            result[name] = _share(socio)
    return result


def _valor_num(value: str):
    raw = re.sub(r"[^\d,.-]", "", str(value or ""))
    if not raw:
        return None
    raw = raw.replace(".", "").replace(",", ".")
    try:
        return float(raw)
    except ValueError:
        return None


def _moeda(value: float) -> str:
    texto = f"{value:,.2f}".replace(",", "@").replace(".", ",").replace("@", ".")
    return f"R$ {texto}"


def _quadro_resumo(event: dict) -> str:
    partes = []
    for nome, pct in _socios_map(event).items():
        partes.append(f"{nome} {pct}" if pct else nome)
    return " · ".join(partes) or "Não informado"


def movimentos_do_ato(event: dict, prev: dict | None) -> list[dict]:
    """Movimentações objetivas do ato, com nomes, matrículas e valores."""
    movs: list[dict] = []
    atual = _socios_map(event)
    anterior = _socios_map(prev) if prev else {}

    cessoes = event.get("cessoes") or []
    # Sócios cuja entrada/saída já é explicada por uma cessão deste mesmo ato —
    # evita mostrar "Entrada de sócio"/"Saída de sócio" redundante ao lado do
    # bloco de "Cessão de quotas" entre as mesmas pessoas.
    cedentes_que_sairam = {
        str(c.get("cedente", "")).strip()
        for c in cessoes
        if str(c.get("cedente", "")).strip() in anterior and str(c.get("cedente", "")).strip() not in atual
    }
    cessionarios_que_entraram = {
        str(c.get("cessionario", "")).strip()
        for c in cessoes
        if str(c.get("cessionario", "")).strip() in atual and str(c.get("cessionario", "")).strip() not in anterior
    }

    if prev:
        for nome, pct in atual.items():
            if nome not in anterior and nome not in cessionarios_que_entraram:
                movs.append(
                    {
                        "tipo": "entrada",
                        "rotulo": "Entrada de sócio",
                        "nome": nome,
                        "extra": f"{pct} do capital" if pct else "",
                    }
                )
        for nome, pct in anterior.items():
            if nome not in atual and nome not in cedentes_que_sairam:
                movs.append(
                    {
                        "tipo": "saida",
                        "rotulo": "Saída de sócio",
                        "nome": nome,
                        "extra": f"detinha {pct}" if pct else "",
                    }
                )

    for cessao in cessoes:
        cedente = str(cessao.get("cedente", "")).strip()
        cessionario = str(cessao.get("cessionario", "")).strip()
        if not (cedente or cessionario):
            continue
        extra_itens = [
            str(cessao.get("participacao", "")).strip(),
            str(cessao.get("valor", "")).strip(),
            str(cessao.get("observacao", "")).strip(),
        ]
        if cessionario in cessionarios_que_entraram:
            extra_itens.append(f"Entrada: {cessionario}")
        if cedente in cedentes_que_sairam:
            extra_itens.append(f"Saída: {cedente}")
        extra = " · ".join(x for x in extra_itens if x)
        movs.append(
            {
                "tipo": "cessao",
                "rotulo": "Cessão de quotas",
                "nome": f"{cedente or '—'} → {cessionario or '—'}",
                "extra": extra,
            }
        )

    for imovel in event.get("imoveis") or []:
        matricula = str(imovel.get("matricula", "")).strip()
        cartorio = str(imovel.get("cartorio", "")).strip() or str(imovel.get("cidade", "")).strip()
        movimento = _normalize(imovel.get("movimento", ""))
        if "said" in movimento or "transfer" in movimento or "restitu" in movimento or "cind" in movimento:
            tipo, rotulo = "imovel-saida", "Saída de imóvel"
        elif "aquis" in movimento or "compr" in movimento:
            tipo, rotulo = "imovel", "Aquisição de imóvel"
        else:
            tipo, rotulo = "imovel", "Integralização de imóvel"
        nome = f"Matrícula {matricula}" if matricula else "Imóvel"
        if cartorio:
            nome = f"{nome} · {cartorio}"
        extra = " — ".join(
            x
            for x in [str(imovel.get("valor", "")).strip(), str(imovel.get("descricao", "")).strip()]
            if x
        )
        movs.append({"tipo": tipo, "rotulo": rotulo, "nome": nome, "extra": extra})

    # Troca de administração e variação de capital não geram mais cartão próprio
    # na timeline — ficam só no "Ver detalhamento" (modal) e na tabela do Word,
    # onde já aparecem como parte do estado da sociedade após o ato.

    sede = str(event.get("sede_apos", "")).strip()
    sede_anterior = str((prev or {}).get("sede_apos", "")).strip() if prev else ""
    if sede and sede_anterior and _normalize(sede) != _normalize(sede_anterior):
        movs.append(
            {"tipo": "sede", "rotulo": "Mudança de sede", "nome": f"{sede_anterior} → {sede}", "extra": ""}
        )

    objeto = str(event.get("objeto_apos", "")).strip()
    objeto_anterior = str((prev or {}).get("objeto_apos", "")).strip() if prev else ""
    if objeto and objeto_anterior and _normalize(objeto) != _normalize(objeto_anterior):
        movs.append(
            {"tipo": "objeto", "rotulo": "Alteração do objeto social", "nome": objeto, "extra": ""}
        )

    filiais = _joined_filiais(event.get("filiais_adicionadas") or [])
    if filiais:
        movs.append({"tipo": "filial", "rotulo": "Filial aberta", "nome": filiais, "extra": ""})

    if not movs:
        capital_apos = str(event.get("capital_social_apos", "")).strip()
        movs.append(
            {
                "tipo": "base" if prev else "constituicao",
                "rotulo": "Constituição" if not prev else "Sem alteração relevante",
                "nome": _quadro_resumo(event),
                "extra": (f"Capital {capital_apos}" if capital_apos else ""),
            }
        )
    return movs


def _periodo(events: list[dict]) -> str:
    anos = [str(e.get("data", ""))[-4:] for e in events if str(e.get("data", ""))[-4:].isdigit()]
    if not anos:
        return ""
    return anos[0] if anos[0] == anos[-1] else f"{anos[0]}–{anos[-1]}"


def _bloco_editavel(chave: str, campos: list, rotulos: list, itens: list,
                    titulo: str, singular: str) -> str:
    """Lista do JSON como linhas editáveis, com + para acrescentar e × para remover.

    Cada linha carrega um <span data-tl-campo> por campo, na ordem de `campos`. A
    serialização lê por POSIÇÃO dentro do bloco, não por índice no atributo — é o que
    permite acrescentar e remover linha sem renumerar nada.
    """
    def _linha(item: dict) -> str:
        celulas = "".join(
            f'<span class="tl2-celula" data-tl-campo="{campo}" '
            f'data-tl-rotulo="{html.escape(rotulo)}">'
            f'{html.escape(str(item.get(campo, "") or ""))}</span>'
            for campo, rotulo in zip(campos, rotulos)
        )
        return ('<li data-tl-item>' + celulas
                + '<button type="button" class="tl2-rm" data-tl-remover '
                  'title="Remover">×</button></li>')

    linhas = "".join(_linha(i) for i in itens if isinstance(i, dict))
    modelo = html.escape(json.dumps(list(zip(campos, rotulos)), ensure_ascii=False), quote=True)
    return (
        f'<div class="tl2-bloco" data-tl-lista="{chave}" data-tl-modelo="{modelo}">'
        f"<h4>{html.escape(titulo)}</h4>"
        f'<ul class="tl2-list">{linhas}</ul>'
        f'<button type="button" class="tl2-add" data-tl-adicionar>+ {html.escape(singular)}</button>'
        "</div>"
    )


def _detalhe_modal(index: int, event: dict) -> str:
    """Painel de detalhe do ato — e, no modo de edição, o formulário do ato.

    Todos os campos de origem vivem aqui. Os cards da coluna são DERIVADOS
    (movimentos_do_ato compara um ato com o anterior), então não há como escrever de
    volta neles: quem se edita é o dado que os gera, e os cards se redesenham depois.
    """
    blocos = "".join([
        _bloco_editavel("socios_apos", ["nome", "participacao", "quotas"],
                        ["Nome", "Participação", "Quotas"],
                        event.get("socios_apos") or [],
                        "Quadro societário após o ato", "sócio"),
        _bloco_editavel("cessoes", ["cedente", "cessionario", "participacao", "valor", "observacao"],
                        ["Cedente", "Cessionário", "Participação", "Valor", "Observação"],
                        event.get("cessoes") or [],
                        "Transmissão de quotas", "cessão"),
        _bloco_editavel("imoveis", ["matricula", "cartorio", "cidade", "valor", "movimento", "descricao"],
                        ["Matrícula", "Cartório", "Cidade", "Valor",
                         "Movimento (integralizacao/saida/aquisicao)", "Descrição"],
                        event.get("imoveis") or [],
                        "Bens imóveis", "imóvel"),
        _bloco_editavel("administradores_apos", ["nome", "cargo"], ["Nome", "Cargo"],
                        event.get("administradores_apos") or [],
                        "Administração após o ato", "administrador"),
        _bloco_editavel("filiais_apos", ["nome", "local"], ["Nome", "Local"],
                        event.get("filiais_apos") or [],
                        "Filiais existentes após o ato", "filial"),
        _bloco_editavel("filiais_adicionadas", ["nome", "local"], ["Nome", "Local"],
                        event.get("filiais_adicionadas") or [],
                        "Filiais abertas NESTE ato", "filial"),
    ])

    def _campo(rotulo: str, chave: str, vazio: str = "—") -> str:
        return (f"<dt>{html.escape(rotulo)}</dt>"
                f'<dd data-tl-campo="{chave}">'
                f'{html.escape(str(event.get(chave, "") or "") or vazio)}</dd>')

    return f"""
      <input type="checkbox" id="tl2-det-{index}" class="tl2-toggle">
      <label class="tl2-btn" for="tl2-det-{index}">Ver detalhamento</label>
      <div class="tl2-modal">
        <label class="tl2-backdrop" for="tl2-det-{index}"></label>
        <div class="tl2-modal-card">
          <div class="tl2-modal-head">
            <div>
              <span class="tl2-modal-date" data-tl-campo="data">{html.escape(str(event.get("data", "")) or "—")}</span>
              <strong data-tl-campo="ato">{html.escape(str(event.get("ato", "")) or "Ato societário")}</strong>
              <small data-tl-campo="numero_arquivamento">{html.escape(str(event.get("numero_arquivamento", "")) or "—")}</small>
            </div>
            <div class="tl2-modal-acoes">
              <button type="button" class="tl2-rm tl2-rm-modal" data-tl-remover-ato
                      title="Excluir este ato inteiro">Excluir ato</button>
              <label class="tl2-close" for="tl2-det-{index}">Fechar</label>
            </div>
          </div>
          <div class="tl2-modal-body">
            <h4>Detalhamento do ato</h4>
            <p data-tl-campo="detalhamento">{html.escape(str(event.get("detalhamento", "")) or "Não informado.")}</p>
            {blocos}
            <h4>Estado da sociedade</h4>
            <dl class="tl2-dl">
              {_campo("Capital social anterior", "capital_social_anterior")}
              {_campo("Capital social", "capital_social_apos")}
              {_campo("Sede", "sede_apos")}
              {_campo("Objeto social", "objeto_apos")}
            </dl>
            <p class="tl2-fonte">Fonte: <span data-tl-campo="fonte">{html.escape(str(event.get("fonte", "")) or "não informada")}</span></p>
          </div>
        </div>
      </div>
    """



CAMPOS_CAIXINHA = ["rotulo", "nome", "extra", "tipo"]
_ROTULOS_CAIXINHA = ["Rótulo", "Texto principal", "Detalhe", "tipo (cor)"]


def _caixinha(mov: dict) -> str:
    """Uma caixa do ato — editável célula a célula, como as do Painel de Deals.

    "tipo" é a chave da cor (entrada, saida, cessao, imovel, admin, capital-up...);
    fica visível só no modo de edição, para dar de trocar a cor sem sair do painel.
    """
    tipo = str(mov.get("tipo", "") or "outros")
    return (
        f'<div class="tl2-move tl2-move-{html.escape(tipo)}" data-tl-item>'
        f'<button type="button" class="tl2-rm" data-tl-remover title="Remover caixa">×</button>'
        f'<span class="tl2-move-label" data-tl-campo="rotulo" data-tl-rotulo="Rótulo">'
        f'{html.escape(str(mov.get("rotulo", "") or ""))}</span>'
        f'<strong data-tl-campo="nome" data-tl-rotulo="Texto principal">'
        f'{html.escape(str(mov.get("nome", "") or "—"))}</strong>'
        f'<small data-tl-campo="extra" data-tl-rotulo="Detalhe">'
        f'{html.escape(str(mov.get("extra", "") or ""))}</small>'
        f'<small class="tl2-tipo" data-tl-campo="tipo" data-tl-rotulo="tipo">'
        f"{html.escape(tipo)}</small></div>"
    )


def render_timeline_html(data: dict) -> str:
    events = data.get("eventos", []) or []
    if not events:
        return '<div class="timeline-empty">A análise ainda não gerou eventos societários.</div>'

    colunas = []
    for index, event in enumerate(events):
        prev = events[index - 1] if index else None
        # As caixinhas nascem derivadas (movimentos_do_ato compara o ato com o anterior),
        # mas viram DADO na primeira vez que são desenhadas. Só assim dá para editá-las,
        # acrescentar e remover caixinha: derivado não se escreve de volta. Depois de
        # editadas, o que vale é o que está gravado.
        movimentos = event.get("movimentos")
        if not isinstance(movimentos, list) or not movimentos:
            movimentos = movimentos_do_ato(event, prev)
        blocos = "".join(_caixinha(m) for m in movimentos) + (
            '<button type="button" class="tl2-add" data-tl-adicionar>+ caixa</button>'
        )
        modelo_caixinha = html.escape(json.dumps(
            list(zip(CAMPOS_CAIXINHA, _ROTULOS_CAIXINHA)), ensure_ascii=False), quote=True)
        # Campos do JSON que o painel não desenha viajam aqui, para o round-trip não
        # perdê-los — e para o servidor não precisar do estado anterior na volta, o que
        # obrigaria a passar um gr.State pelo JavaScript.
        extra = html.escape(json.dumps(
            {c: event.get(c) for c in ("categorias",) if c in event}, ensure_ascii=False
        ), quote=True)
        colunas.append(
            f"""
        <article class="tl2-col" data-tl-evento data-tl-extra="{extra}">
          <div class="tl2-col-head">
            <strong data-tl-campo="data">{html.escape(str(event.get("data", "")) or "—")}</strong>
            <span data-tl-campo="ato">{html.escape(str(event.get("ato", "")) or "Ato societário")}</span>
            <button type="button" class="tl2-rm tl2-rm-ato" data-tl-remover-ato
                    title="Remover este ato">×</button>
          </div>
          <div class="tl2-axis"><i class="tl2-pin"></i></div>
          <div class="tl2-moves" data-tl-lista="movimentos" data-tl-modelo="{modelo_caixinha}">{blocos}</div>
          <div class="tl2-quadro">Quadro após o ato<strong>{html.escape(_quadro_resumo(event))}</strong></div>
          {_detalhe_modal(index, event)}
        </article>
        """
        )

    # O <style> vai DENTRO da seção de propósito: a exportação de imagem lê esse CSS
    # pelo próprio DOM para rasterizar a timeline exatamente como ela aparece na tela.
    return f"""
    <section class="tl2-shell" id="timeline-export-area" data-tl-raiz>
      <style id="tl2-export-style">{TL2_CSS}</style>
      <div class="tl2-head">
        <span class="tl2-kicker">Cronologia societária</span>
        <h2 data-tl-campo="empresa">{html.escape(data.get("empresa", "") or "Sociedade analisada")}</h2>
        <p>CNPJ <span data-tl-campo="cnpj">{html.escape(data.get("cnpj", "") or "—")}</span>
           · {len(events)} atos · {_periodo(events)}</p>
        <p class="tl2-dica">Modo de edição: clique em qualquer texto para reescrever —
           inclusive dentro de <b>Ver detalhamento</b>, onde ficam o quadro societário, as
           cessões e os imóveis. <b>×</b> remove o ato, a caixa ou a linha; <b>+</b>
           acrescenta. Data e ato aparecem na coluna e no pop-up: alterar um atualiza o
           outro na hora.</p>
      </div>
      <div class="tl2-scroll">
        <div class="tl2-track" style="--tl2-cols:{len(events)}">{"".join(colunas)}</div>
      </div>
      <button type="button" class="tl2-add tl2-add-ato" data-tl-adicionar-ato>+ Adicionar ato</button>
    </section>
    """


def timeline_analisar(files):
    yield (
        "Preparando os atos societários...",
        '<div class="timeline-loading">Lendo os documentos e reconstruindo a evolução societária…</div>',
        {},
    )
    try:
        data = _extract(files)
    except gr.Error:
        raise
    except ValueError as exc:
        raise gr.Error(str(exc)) from exc
    except Exception as exc:
        mensagem = _mensagem_erro_gemini(exc, getattr(exc, "clients_gemini", None))
        yield (mensagem, f'<div class="timeline-empty">{html.escape(mensagem)}</div>', {})
        raise gr.Error(mensagem) from exc
    yield (
        f"Concluído: {len(data.get('eventos', []))} evento(s) societário(s) identificado(s).",
        render_timeline_html(data),
        data,
    )


# ══════════════════════════════════════════════════════════════════════════════
# Edição no próprio HTML
#
# O painel renderizado É o editor: "Editar timeline" liga contenteditable nos campos
# marcados com data-tl-campo, você digita direto no card, e "Concluir edição" devolve a
# árvore serializada para cá. Mesmo modelo do Painel de Deals.
#
# Os cards da coluna são derivados (movimentos_do_ato compara um ato com o anterior),
# então a superfície editável é o cabeçalho da coluna e o modal de detalhe, onde vivem
# os campos de origem. Ao concluir, os cards se redesenham a partir do dado novo.
# ══════════════════════════════════════════════════════════════════════════════

# Campos escalares do evento que o HTML expõe para edição.
CAMPOS_TEXTO_DO_EVENTO = [
    "data", "ato", "numero_arquivamento", "detalhamento",
    "capital_social_anterior", "capital_social_apos", "sede_apos", "objeto_apos", "fonte",
]

# Listas do evento e os campos de cada item, na ordem em que o HTML as desenha.
LISTAS_DO_EVENTO = {
    "movimentos": CAMPOS_CAIXINHA,
    "socios_apos": ["nome", "participacao", "quotas"],
    "cessoes": ["cedente", "cessionario", "participacao", "valor", "observacao"],
    "imoveis": ["matricula", "cartorio", "cidade", "valor", "movimento", "descricao"],
    "administradores_apos": ["nome", "cargo"],
    "filiais_apos": ["nome", "local"],
    "filiais_adicionadas": ["nome", "local"],
}

MOVIMENTOS_IMOVEL = ["integralizacao", "saida", "aquisicao"]

# Texto que o render usa quando o campo está vazio. Voltando igual, é vazio de novo —
# senão o primeiro "Concluir edição" gravaria "—" e "Não informado." como conteúdo.
_PLACEHOLDERS = {"—", "-", "Não informado.", "não informada", "Ato societário",
                 "Sociedade analisada"}


def _limpar(valor) -> str:
    texto = re.sub(r"\s+", " ", str(valor or "")).strip()
    return "" if texto in _PLACEHOLDERS else texto


def _evento_do_html(recebido: dict) -> dict:
    """Um ato, montado a partir do que o navegador devolveu."""
    evento = {}
    extra = recebido.get("_extra")
    if isinstance(extra, dict):
        evento.update(extra)          # campos que o painel não desenha (categorias)
    for campo in CAMPOS_TEXTO_DO_EVENTO:
        evento[campo] = _limpar(recebido.get(campo))
    for chave, campos in LISTAS_DO_EVENTO.items():
        itens = []
        for item in recebido.get(chave) or []:
            if not isinstance(item, dict):
                continue
            valores = {campo: _limpar(item.get(campo)) for campo in campos}
            if any(valores.values()):     # linha em branco não vira item
                itens.append(valores)
        evento[chave] = itens
    return evento


def aplicar_edicao_html(bruto: str):
    """Reconstrói o JSON a partir da árvore que o navegador serializou.

    O painel carrega o schema inteiro — inclusive o que ele não desenha, que viaja em
    `data-tl-extra`. Por isso não precisa do estado anterior aqui, e o `gr.State` não
    precisa atravessar o JavaScript, que é onde essa ponte seria frágil.

    Os atos são lidos por POSIÇÃO, na ordem em que estão na tela e já com as inclusões
    e remoções feitas ali.
    """
    if not isinstance(bruto, (str, dict)) or (isinstance(bruto, str) and not bruto.strip()):
        return None                       # nada veio do navegador
    try:
        editado = json.loads(bruto) if isinstance(bruto, str) else bruto
    except (json.JSONDecodeError, TypeError):
        editado = None
    if not isinstance(editado, dict):
        return None                       # serialização torta: quem chama mantém o que tinha

    return {
        "empresa": _limpar(editado.get("empresa")),
        "cnpj": _limpar(editado.get("cnpj")),
        "eventos": [_evento_do_html(e) for e in (editado.get("eventos") or [])
                    if isinstance(e, dict)],
    }


def timeline_aplicar_html(modo: str, bruto: str):
    """Liga e desliga o modo de edição do painel.

    Entrando, só troca o rótulo do botão — quem pinta os campos é o JS, e o estado fica
    intocado via gr.skip(). Saindo, aplica o que foi digitado e redesenha, para os cards
    derivados refletirem o dado novo.
    """
    if str(modo or "0") == "1":
        return ("1", "", gr.skip(), gr.skip(),
                gr.update(value="Concluir edição", variant="primary"))

    novo = aplicar_edicao_html(bruto)
    if novo is None:   # serialização torta: não mexe no que já estava
        return ("0", "", gr.skip(), gr.skip(),
                gr.update(value="Editar timeline", variant="secondary"))
    return ("0", "", novo, render_timeline_html(novo),
            gr.update(value="Editar timeline", variant="secondary"))


def ordenar_eventos(data: dict):
    """Reordena os atos por data. Sai do botão da barra, não do painel."""
    data = copy.deepcopy(data or {})
    data.get("eventos", []).sort(key=lambda e: _date_key(e.get("data", "")))
    return data, render_timeline_html(data)






def timeline_salvar_imagem(captura: str) -> str:
    """Salva o PNG capturado do próprio HTML da timeline (rasterizado no navegador) e
    devolve o caminho para o link de download.

    A captura é a imagem exata do que está na tela — mesmo layout, mesmas cores, mesmos
    cards. Chega como data URL, produzida pelo JS que roda antes desta função no mesmo
    evento de clique. Quando vier vazia, o motivo quase sempre é não haver timeline na
    tela; avisar é melhor do que entregar um PNG diferente do que a usuária está vendo.
    """
    captura = str(captura or "")
    if not captura.startswith("data:image/png;base64,"):
        raise gr.Error(
            "Não há timeline na tela para exportar. Analise os atos societários primeiro; "
            "se a timeline já estiver visível, recarregue a página e tente de novo."
        )
    try:
        conteudo = base64.b64decode(captura.split(",", 1)[1], validate=True)
    except (binascii.Error, ValueError):
        conteudo = b""
    if not conteudo:
        raise gr.Error(
            "O navegador não conseguiu gerar a imagem da timeline. Recarregue a página e "
            "tente de novo; se persistir, use Exportar HTML."
        )
    output = Path(tempfile.gettempdir()) / "timeline_societaria.png"
    output.write_bytes(conteudo)
    return str(output)


TL2_CSS = """
.timeline-empty, .timeline-loading {
    margin: 18px 0; padding: 44px 24px; text-align: center; background: #f5f4f1;
    border: 1px dashed #cbc9c1; color: #71746f;
}
.timeline-actions { align-items: flex-start; margin-top: 18px; }
.timeline-action-stack { gap: 8px !important; }

/* ── Timeline Societária — cronologia nomeada (v2) ──────────────────────── */
.tl2-shell {
    margin-top: 18px; background: #fff; border: 1px solid #d9d9d5;
    box-shadow: 0 12px 28px rgba(36, 38, 35, .07);
}
.tl2-head { padding: 24px 28px 20px; border-bottom: 1px solid #dededb; }
.tl2-kicker {
    display: block; color: #1A56A0; font-size: 11px; font-weight: 800; letter-spacing: .14em;
    text-transform: uppercase;
}
.tl2-head h2 { margin: 6px 0 4px; font-size: 25px; font-weight: 700; color: #2c302c; }
.tl2-head p { margin: 0; font-size: 13px; color: #737670; }
.tl2-scroll { overflow-x: auto; padding: 26px 28px 30px; }
.tl2-track {
    display: grid; grid-template-columns: repeat(var(--tl2-cols), 200px);
    min-width: 100%; column-gap: 0;
}
.tl2-col { padding: 0 10px; display: flex; flex-direction: column; }
.tl2-col-head { text-align: center; }
.tl2-col-head strong { display: block; font-size: 14px; font-weight: 800; color: #1f211f; }
.tl2-col-head span {
    display: block; margin-top: 2px; font-size: 10px; font-weight: 800; letter-spacing: .06em;
    text-transform: uppercase; color: #737670;
}
.tl2-axis {
    margin-top: 12px; height: 3px; background: #e4e4e0; display: flex; align-items: center;
    justify-content: center;
}
.tl2-pin {
    width: 13px; height: 13px; border-radius: 50%; background: #1A56A0; border: 2px solid #fff;
    box-shadow: 0 0 0 1px #1A56A0;
}
.tl2-moves { margin-top: 18px; display: flex; flex-direction: column; gap: 7px; }
.tl2-move { padding: 8px 9px; border-left: 3px solid #77817a; background: #f6f6f4; }
.tl2-move-constituicao { border-left-color: #2f6b3a; background: rgba(47, 107, 58, .07); }
.tl2-move-constituicao .tl2-move-label { color: #2f6b3a; }
.tl2-move-label {
    display: block; font-size: 9px; font-weight: 800; letter-spacing: .08em; text-transform: uppercase;
    color: #77817a;
}
.tl2-move strong {
    display: block; margin-top: 3px; font-size: 12px; font-weight: 700; color: #1f211f;
    line-height: 1.35; text-wrap: pretty;
}
.tl2-move small { display: block; margin-top: 3px; font-size: 10px; color: #666963; line-height: 1.4; }
.tl2-move-entrada { border-left-color: #235472; background: rgba(35, 84, 114, .07); }
.tl2-move-entrada .tl2-move-label { color: #235472; }
.tl2-move-saida { border-left-color: #dc4405; background: rgba(220, 68, 5, .07); }
.tl2-move-saida .tl2-move-label { color: #dc4405; }
.tl2-move-cessao { border-left-color: #1c6e8c; background: rgba(28, 110, 140, .07); }
.tl2-move-cessao .tl2-move-label { color: #1c6e8c; }
.tl2-move-imovel { border-left-color: #a6486a; background: rgba(166, 72, 106, .08); }
.tl2-move-imovel .tl2-move-label { color: #a6486a; }
.tl2-move-imovel-saida { border-left-color: #dc4405; background: rgba(220, 68, 5, .07); }
.tl2-move-imovel-saida .tl2-move-label { color: #dc4405; }
.tl2-move-admin { border-left-color: #1c6e8c; background: rgba(28, 110, 140, .07); }
.tl2-move-admin .tl2-move-label { color: #1c6e8c; }
.tl2-move-capital-up { border-left-color: #2f6b3a; background: rgba(47, 107, 58, .07); }
.tl2-move-capital-up .tl2-move-label { color: #2f6b3a; }
.tl2-move-capital-down { border-left-color: #dc4405; background: rgba(220, 68, 5, .07); }
.tl2-move-capital-down .tl2-move-label { color: #dc4405; }
.tl2-move-sede { border-left-color: #dc4405; background: rgba(220, 68, 5, .07); }
.tl2-move-sede .tl2-move-label { color: #dc4405; }
.tl2-move-objeto { border-left-color: #9a7b2f; background: rgba(154, 123, 47, .08); }
.tl2-move-objeto .tl2-move-label { color: #9a7b2f; }
.tl2-move-filial { border-left-color: #77636a; background: rgba(119, 99, 106, .08); }
.tl2-move-filial .tl2-move-label { color: #77636a; }
.tl2-quadro {
    margin-top: 12px; padding-top: 10px; border-top: 1px dashed #dededa; font-size: 10px;
    color: #737670; line-height: 1.45;
}
.tl2-quadro strong { display: block; margin-top: 2px; color: #3f423e; font-weight: 700; }
.tl2-toggle { display: none !important; }
.tl2-btn {
    margin-top: 10px; padding: 7px 10px; background: #fff; border: 1px solid #dcdcd7;
    border-radius: 999px; font-size: 10px; font-weight: 800; letter-spacing: .06em;
    text-transform: uppercase; color: #4a4d48; text-align: center; cursor: pointer;
    transition: background .16s ease, color .16s ease;
}
.tl2-btn:hover { background: #f5f5f2; color: #1f211f; }
.tl2-modal {
    display: none; position: fixed; inset: 0; z-index: 90; align-items: center;
    justify-content: center; padding: 40px;
}
.tl2-toggle:checked ~ .tl2-modal { display: flex; }
.tl2-backdrop { position: absolute; inset: 0; background: rgba(31, 33, 31, .42); cursor: pointer; }
.tl2-modal-card {
    position: relative; width: 640px; max-height: 84vh; overflow-y: auto; background: #fff;
    border: 1px solid #d9d9d5; box-shadow: 0 24px 60px rgba(31, 33, 31, .28); text-align: left;
}
.tl2-modal-head {
    display: flex; align-items: flex-start; justify-content: space-between; gap: 16px;
    padding: 20px 24px 16px; border-bottom: 1px solid #ededea;
}
.tl2-modal-date {
    display: block; font-size: 11px; font-weight: 800; letter-spacing: .1em; text-transform: uppercase;
    color: #dc4405;
}
.tl2-modal-head strong { display: block; margin-top: 5px; font-size: 19px; font-weight: 800; color: #1f211f; }
.tl2-modal-head small { display: block; margin-top: 2px; font-size: 11px; color: #92948f; }
.tl2-close {
    flex: none; padding: 7px 14px; background: #f5f5f2; border: 1px solid #e4e4e0; border-radius: 999px;
    font-size: 10px; font-weight: 800; letter-spacing: .06em; text-transform: uppercase;
    color: #4a4d48; cursor: pointer;
}
.tl2-close:hover { background: #ededea; }
.tl2-modal-body { padding: 20px 24px 24px; }
.tl2-modal-body h4 {
    margin: 18px 0 8px; font-size: 10px; font-weight: 800; letter-spacing: .1em;
    text-transform: uppercase; color: #92948f;
}
.tl2-modal-body h4:first-child { margin-top: 0; }
.tl2-modal-body p { margin: 0; font-size: 13px; line-height: 1.65; color: #3f423e; text-wrap: pretty; }
.tl2-chips { display: flex; flex-wrap: wrap; gap: 8px; }
.tl2-chip {
    padding: 5px 10px; background: #f4f7f9; border: 1px solid #e2eaef; font-size: 12px; color: #235472;
}
.tl2-list { margin: 0; padding-left: 18px; font-size: 12px; line-height: 1.65; color: #3f423e; }
.tl2-dl { display: grid; grid-template-columns: auto 1fr; gap: 6px 14px; margin: 0; font-size: 12px; }
.tl2-dl dt { color: #92948f; }
.tl2-dl dd { margin: 0; color: #4a4d48; font-weight: 600; }
.tl2-fonte { margin-top: 18px !important; font-size: 10px !important; color: #92948f !important; }

/* ── Modo de edição — mesma linguagem do Painel de Deals ─────────────────── */
/* Fora do modo de edição, nada disso aparece: os botões somem e os campos não
   mostram contorno nenhum. */
.tl2-add, .tl2-rm, .tl2-dica { display: none; }
.tl2-edit-mode .tl2-dica {
    display: block; margin-top: 12px; padding: 8px 11px; background: #FBF6E7;
    border: 1px solid #E8D9A0; font-size: 11px !important; line-height: 1.6;
    color: #8A6D1F !important;
}
.tl2-edit-mode .tl2-dica b { color: #6f571a; }
.tl2-bloco > h4 { margin-top: 14px; }
.tl2-bloco .tl2-celula:not(:last-of-type)::after { content: " · "; color: #b9bbb6; }
.tl2-bloco .tl2-list { list-style: none; padding-left: 0; }
.tl2-bloco .tl2-list li { padding: 2px 0; }

.tl2-edit-mode [contenteditable="true"] {
    outline: 1px dashed #8FB2DA; outline-offset: 2px; border-radius: 3px;
    background: rgba(235, 241, 250, .38); cursor: text; min-width: 26px;
    display: inline-block;
}
.tl2-edit-mode [contenteditable="true"]:focus {
    outline: 2px solid #1A56A0; background: #fff;
}
.tl2-edit-mode p[contenteditable="true"], .tl2-edit-mode dd[contenteditable="true"] { display: block; }
.tl2-edit-mode .tl2-celula[data-tl-rotulo]:empty::before {
    content: attr(data-tl-rotulo); color: #b0b2ad; font-style: italic;
}
.tl2-edit-mode .tl2-add, .tl2-edit-mode .tl2-rm { display: inline-block; }
.tl2-add {
    margin-top: 6px; padding: 3px 9px; font-size: 11px; font-weight: 700;
    color: #1A56A0; background: #eef3fa; border: 1px solid #cfdcec; border-radius: 3px;
    cursor: pointer;
}
.tl2-add:hover { background: #e2ebf7; }
.tl2-add-ato { margin: 0 28px 22px; }
.tl2-rm {
    margin-left: 8px; padding: 0 6px; font-size: 12px; line-height: 1.5;
    color: #dc4405; background: transparent; border: 1px solid #f0d3c8; border-radius: 3px;
    cursor: pointer;
}
.tl2-rm:hover { background: #fdeee8; }
.tl2-rm-ato { float: right; }
.tl2-modal-acoes { flex: none; display: flex; align-items: center; gap: 8px; }
.tl2-rm-modal {
    float: none; margin-left: 0; padding: 6px 11px; font-size: 10px; font-weight: 800;
    letter-spacing: .06em; text-transform: uppercase; line-height: 1.4;
}

/* A chave da cor da caixa só aparece quando se está editando. */
.tl2-tipo { display: none; }
.tl2-edit-mode .tl2-tipo {
    display: inline-block; margin-top: 4px; font-size: 9px !important;
    letter-spacing: .06em; text-transform: uppercase; color: #92948f !important;
}
.tl2-move { position: relative; }
.tl2-edit-mode .tl2-move .tl2-rm {
    position: absolute; top: 4px; right: 4px; margin: 0; padding: 0 4px; font-size: 11px;
}
.tl2-edit-mode .tl2-moves { padding-bottom: 4px; }
"""


_BOTOES_DE_EDICAO = re.compile(
    r'<button[^>]*class="[^"]*\btl2-(?:add|rm)\b[^"]*"[^>]*>.*?</button>', re.DOTALL)
_ATRIBUTOS_DE_EDICAO = re.compile(r'\sdata-tl-[a-z-]+(?:="[^"]*")?')
_DICA_DE_EDICAO = re.compile(r'<p class="tl2-dica">.*?</p>', re.DOTALL)


def remover_edicao(marcacao: str) -> str:
    """Tira do HTML tudo o que só existe para editar no painel.

    O arquivo exportado é registro: os botões de + e ×, os ganchos data-tl-* e a dica do
    modo de edição não têm para onde salvar fora do Space, então não vão junto. O
    conteúdo — colunas, caixas, modais de detalhamento — fica intacto.
    """
    marcacao = _BOTOES_DE_EDICAO.sub("", marcacao or "")
    marcacao = _DICA_DE_EDICAO.sub("", marcacao)
    return _ATRIBUTOS_DE_EDICAO.sub("", marcacao)


def timeline_exportar_html(data: dict) -> str:
    """Exporta a timeline interativa (colunas + modais de detalhamento) como um HTML autocontido."""
    events = (data or {}).get("eventos", [])
    if not events:
        raise gr.Error("Gere a timeline antes de exportar.")
    empresa = html.escape(str(data.get("empresa", "") or "Sociedade analisada"))
    corpo = remover_edicao(render_timeline_html(data))
    documento = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Cronologia Societária — {empresa}</title>
<style>
body {{ margin: 0; padding: 0; background: #FFFFFF; font-family: Arial, Helvetica, sans-serif; }}
{TL2_CSS}
</style>
</head>
<body>
{corpo}
</body>
</html>
"""
    output = Path(tempfile.gettempdir()) / "timeline_societaria_interativa.html"
    output.write_text(documento, encoding="utf-8")
    return str(output)


def _normalize(value: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFD", str(value or "").lower()) if unicodedata.category(c) != "Mn"
    )


def _find_timeline_table(doc: Document):
    body = doc.element.body
    for paragraph in doc.paragraphs:
        normalized = _normalize(paragraph.text)
        if "cronologia societaria" not in normalized:
            continue
        start = list(body).index(paragraph._p)
        for child in list(body)[start + 1 :]:
            if child.tag == qn("w:tbl"):
                from docx.table import Table

                return Table(child, doc)
            if child.tag == qn("w:p") and "evidencias" in _normalize("".join(child.itertext())):
                break
    return None


def _set_cell(cell, value: str, bold: bool = False, color: str = "555555"):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value or ""))
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def detalhamento_word(event: dict, prev: dict | None) -> list[str]:
    """Tópicos do dossiê: detalhamento do ato + movimentações objetivas, sem a fonte."""
    itens = []
    detalhe = str(event.get("detalhamento", "")).strip()
    if detalhe:
        itens.append(detalhe)
    for m in movimentos_do_ato(event, prev):
        if m["tipo"] == "base":
            continue
        itens.append(f"{m['rotulo']}: {m['nome']}" + (f" ({m['extra']})" if m.get("extra") else ""))
    quadro = _quadro_resumo(event)
    if quadro and quadro != "Não informado":
        itens.append("Quadro societário após o ato: " + quadro)
    admin = _joined_people(event.get("administradores_apos", []))
    if admin:
        itens.append("Administração: " + admin)
    capital = str(event.get("capital_social_apos", "")).strip()
    if capital:
        itens.append("Capital social: " + capital)
    return itens or ["Sem alterações relevantes registradas."]


def _set_cell_bullets(cell, itens: list[str], color: str = "555555"):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for index, item in enumerate(itens):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run("• " + str(item))
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(color)


def _set_col_widths(table, widths=(Inches(1.0), Inches(1.1), Inches(4.4))):
    """Data/Ato mais estreitos; Detalhamento absorve o espaço restante."""
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = width
    for index, width in enumerate(widths):
        if index < len(table.columns):
            table.columns[index].width = width


def _fill_word_table(table, data: dict):
    if not table.rows:
        return
    template_row = copy.deepcopy(table.rows[1]._tr if len(table.rows) > 1 else table.rows[0]._tr)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    events = (data or {}).get("eventos", [])
    if not events:
        table._tbl.append(copy.deepcopy(template_row))
        _set_col_widths(table)
        return
    for index, event in enumerate(events):
        row_xml = copy.deepcopy(template_row)
        table._tbl.append(row_xml)
        row = table.rows[-1]
        prev = events[index - 1] if index else None
        _set_cell(row.cells[0], event.get("data", ""))
        _set_cell(row.cells[1], event.get("ato", ""))
        _set_cell_bullets(row.cells[2], detalhamento_word(event, prev))
    _set_col_widths(table)


def timeline_gerar_word(data: dict, dossie_file: Any = None) -> str:
    if not (data or {}).get("eventos"):
        raise gr.Error("Gere a timeline antes de exportar a tabela Word.")
    dossier_path = _file_path(dossie_file)
    if dossier_path and Path(dossier_path).exists():
        doc = Document(dossier_path)
        output_name = "dossie_com_cronologia_societaria.docx"
    elif TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        output_name = "cronologia_societaria.docx"
    else:
        doc = Document()
        title = doc.add_paragraph()
        run = title.add_run("c. Cronologia Societária — Atos Relevantes")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, ["DATA", "ATO", "DETALHAMENTO"]):
            _set_cell(cell, label, bold=True, color="FFFFFF")
        output_name = "cronologia_societaria.docx"

    table = _find_timeline_table(doc)
    if table is None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("c. Cronologia Societária — Atos Relevantes")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, ["DATA", "ATO", "DETALHAMENTO"]):
            _set_cell(cell, label, bold=True, color="FFFFFF")
    _fill_word_table(table, data)

    output = Path(tempfile.gettempdir()) / output_name
    doc.save(output)
    return str(output)
