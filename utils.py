# -*- coding: utf-8 -*-
import os
import re
import time
import tempfile
import threading
from pathlib import Path

import fitz
from google import genai
from google.genai import types
from docx import Document as DocxDoc
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Timeout (ms) por requisicao HTTP ao Gemini. Converte travamento de socket
# (conexao viva mas sem resposta) em excecao, permitindo que o _retry atue.
# 10 min cobre folgado ate a consolidacao; acima disso e hang de verdade.
GEMINI_TIMEOUT_MS = int(os.getenv("GEMINI_TIMEOUT_MS", "600000"))

# Modelos por responsabilidade. As variaveis permitem override no ambiente, mas os
# defaults seguem a arquitetura recomendada: Lite para leitura/JSON em alto volume e
# Flash 3.6 para sintese juridica e perguntas.
GEMINI_MODEL_EXTRACAO = os.getenv("GEMINI_MODEL_EXTRACAO", "gemini-3.5-flash-lite")
GEMINI_MODEL_RELATORIO = os.getenv("GEMINI_MODEL_RELATORIO", "gemini-3.6-flash")
# Página digitalizada exige leitura visual e interpretação de contratos/atos.
# Por padrão usa o mesmo modelo forte da consolidação; pode ser alterado no Space
# sem rebaixar a extração pesquisável, que continua no modelo de alto volume.
GEMINI_MODEL_OCR = os.getenv("GEMINI_MODEL_OCR", GEMINI_MODEL_RELATORIO)
GEMINI_MODEL_ESTRUTURADO = os.getenv("GEMINI_MODEL_ESTRUTURADO", "gemini-3.5-flash-lite")
GEMINI_MODEL_QA = os.getenv("GEMINI_MODEL_QA", "gemini-3.6-flash")


def _pagina_tem_texto_pesquisavel(page) -> bool:
    """Heurística única para distinguir texto pesquisável de imagem/OCR ruim."""
    texto = (page.get_text() or "").strip()
    if len(texto) < 50 or len(texto.split()) < 30:
        return False
    alfa = sum(c.isalpha() for c in texto)
    return bool(texto) and alfa / len(texto) >= 0.4


def _paginas_digitalizadas_pdf(pdf_path: str, offset: int = 0) -> list[int]:
    """Devolve páginas absolutas (base 1) sem camada textual confiável.

    A detecção é local e determinística. Ela não faz OCR; serve para escolher o
    modelo visual forte antes da chamada ao Gemini e para registrar a cobertura.
    """
    digitalizadas = []
    with fitz.open(pdf_path) as doc:
        for indice, page in enumerate(doc):
            if not _pagina_tem_texto_pesquisavel(page):
                digitalizadas.append(offset + indice + 1)
    return digitalizadas

# Trava global de processamento pesado de PDF (compressao/OCR via PyMuPDF) — compartilhada
# entre rj.py e processos.py. Nao limita quantas analises rodam ao mesmo tempo (isso e o
# concurrency_limit do Gradio); limita só quantos CHUNKS podem estar sendo comprimidos/
# renderizados em CPU/RAM ao mesmo tempo no container inteiro, independente de quantas
# analises/abas dispararam esse trabalho. A espera pela resposta da Gemini (I/O, nao CPU)
# NAO passa por essa trava — só a etapa de compressao, que é o gargalo real de recurso.
LIMITE_PROCESSAMENTO_PESADO = int(os.getenv("LIMITE_PROCESSAMENTO_PESADO_PDF", "6"))
_SEMAFORO_PROCESSAMENTO_PESADO = threading.Semaphore(LIMITE_PROCESSAMENTO_PESADO)


def _get_gemini_clients() -> list:
    """Lê GEMINI_API_KEY_1, _2, _3, ... (nesta ordem, sem pular número) e devolve um
    client por chave configurada — permite espalhar carga entre N contas/projetos Google
    distintos (cada um com sua própria cota), em vez de fixar em 2. GEMINI_API_KEY (sem
    sufixo) é o fallback legado só pra chave 1."""
    http_opts = types.HttpOptions(timeout=GEMINI_TIMEOUT_MS)
    chaves = []
    k1 = os.getenv("GEMINI_API_KEY_1") or os.getenv("GEMINI_API_KEY")
    if not k1:
        raise RuntimeError("GEMINI_API_KEY_1 não configurada.")
    chaves.append(k1)
    n = 2
    while True:
        k = os.getenv(f"GEMINI_API_KEY_{n}")
        if not k:
            break
        chaves.append(k)
        n += 1
    return [genai.Client(api_key=k, http_options=http_opts) for k in chaves]


def _erro_gemini_permite_failover(exc: Exception) -> bool:
    """Indica erros que podem mudar de uma chave/projeto Gemini para outro.

    Não inclui erros de schema, JSON ou programação: nesses casos trocar a chave
    apenas esconderia o defeito real. As chaves nunca são incluídas na mensagem.
    """
    msg = str(exc).lower()
    return any(token in msg for token in (
        "401", "unauthenticated", "access_token_type_unsupported",
        "403", "permission_denied", "api_key_invalid", "api key not valid",
        "429", "resource_exhausted", "quota", "rate limit",
        "404", "not_found", "model is no longer available",
    ))


def _executar_com_failover_gemini(
    clients: list,
    fn,
    *,
    indice_inicial: int = 0,
    ao_falhar=None,
):
    """Executa ``fn(client, indice)`` e roda para a próxima credencial elegível.

    Cada cliente é tentado no máximo uma vez por esta camada. O número de
    tentativas internas de ``_retry`` permanece sob controle de quem chama.
    """
    if not clients:
        raise RuntimeError("Nenhum cliente Gemini configurado.")

    ultimo_erro = None
    total = len(clients)
    for deslocamento in range(total):
        indice = (indice_inicial + deslocamento) % total
        try:
            return fn(clients[indice], indice)
        except Exception as exc:
            ultimo_erro = exc
            ultimo = deslocamento == total - 1
            if ultimo or not _erro_gemini_permite_failover(exc):
                raise
            proximo = (indice + 1) % total
            if ao_falhar is not None:
                ao_falhar(indice, proximo, exc)

    raise ultimo_erro or RuntimeError("Falha em todas as credenciais Gemini.")


def _get_clients_rj():
    """Mantido por compatibilidade — usa só as 2 primeiras chaves (client1, client2)."""
    clients = _get_gemini_clients()
    return clients[0], (clients[1] if len(clients) > 1 else clients[0])


def _get_clients_proc():
    """Mantido por compatibilidade — usa só as 2 primeiras chaves (client1, client2)."""
    clients = _get_gemini_clients()
    return clients[0], (clients[1] if len(clients) > 1 else clients[0])


def _filtrar_arquivos_existentes(paths: list, log: list) -> list:
    """Remove da lista qualquer arquivo que tenha sumido do disco entre o upload e o
    processamento (ex.: limpeza do /tmp do container) — evita que o processamento inteiro
    trave com um FileNotFoundError cru por causa de UM arquivo problemático."""
    existentes = []
    for p in paths:
        if Path(p).exists():
            existentes.append(p)
        else:
            log.append(f"   ⚠️ '{Path(p).name}' não encontrado após o upload (tente reenviar este arquivo).")
    return existentes


def _erro_gemini_e_teto_de_gasto(exc: Exception) -> bool:
    """429 por teto de gasto do projeto (spend cap) — diferente de rate limit.

    Insistir não resolve: a cota só volta quando o teto for elevado no AI Studio ou
    quando virar o mês. Serve para pular o backoff e ir direto para a próxima chave.
    """
    msg = str(exc).lower()
    return "spend cap" in msg or "spending cap" in msg


def _retry(fn, tentativas=5, espera_base=20):
    for t in range(1, tentativas + 1):
        try:
            return fn()
        except Exception as e:
            if _erro_gemini_e_teto_de_gasto(e):
                raise
            msg = str(e)
            msg_low = msg.lower()
            retryable = (
                any(c in msg for c in ["503", "500", "504", "429", "UNAVAILABLE", "overloaded"])
                or any(c in msg_low for c in [
                    "timeout", "timed out", "deadline", "deadlineexceeded",
                    "connection reset", "connection error", "connecterror",
                    "read timed out", "readtimeout",
                    "resource_exhausted", "rate limit", "quota",
                    "json_invalid", "eof while parsing",
                ])
            )
            if retryable and t < tentativas:
                time.sleep(espera_base * t)
            else:
                raise
    raise RuntimeError("Falha após todas as tentativas.")


def _gerar_docx(relatorio: str, titulo: str = "Análise Jurídica") -> str:
    caminho = os.path.join(tempfile.gettempdir(), f"{titulo.lower().replace(' ', '_')}.docx")
    doc = DocxDoc()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(titulo)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    doc.add_paragraph()

    REF_PAT = re.compile(
        r'\(([^)]*(?:Mov\.|fls?\.|Fls?\.|Evento\s+n[oº°]?|ID\s*\d+'
        r'|identificador\s+processual\s+n[aã]o\s+localizado|\|\s*fls?\.)[^)]*)\)',
        re.IGNORECASE,
    )
    HEADING_PAT = re.compile(
        r'^[A-Z]\.\s+[A-ZÁÉÍÓÚ]|^[A-Z]\.\d+\.|^\d+\.\s+[A-ZÁÉÍÓÚ]'
    )

    BULLET_MAP = [
        ('•',  'List Bullet',   None),
        ('∘',  'List Bullet 2', None),
        ('▪',  'List Bullet 3', None),
        ('▵',  'List Bullet 3', Pt(12)),
    ]

    def _add_runs(paragraph, text):
        pos = 0
        for m in REF_PAT.finditer(text):
            if m.start() > pos:
                run = paragraph.add_run(text[pos:m.start()])
                run.font.name = "Calibri"
            run = paragraph.add_run(f"({m.group(1)})")
            run.italic = True
            run.font.name = "Calibri"
            pos = m.end()
        if pos < len(text):
            run = paragraph.add_run(text[pos:])
            run.font.name = "Calibri"

    for linha in relatorio.split("\n"):
        s = linha.strip()
        if not s:
            doc.add_paragraph()
            continue

        if HEADING_PAT.match(s):
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
            continue

        bullet_style = None
        extra_indent = None
        for char, style, extra in BULLET_MAP:
            if s.startswith(char):
                bullet_style = style
                extra_indent = extra
                s = s[len(char):].strip()
                break

        if bullet_style:
            try:
                p = doc.add_paragraph(style=bullet_style)
            except Exception:
                p = doc.add_paragraph()
            if extra_indent:
                current = p.paragraph_format.left_indent or Pt(0)
                p.paragraph_format.left_indent = current + extra_indent
        else:
            p = doc.add_paragraph()

        _add_runs(p, s)

    doc.save(caminho)
    return caminho


def _responder_pergunta_generica(pergunta: str, relatorio: str, client, model: str) -> str:
    if not relatorio.strip():
        return "Gere uma análise primeiro."
    if not pergunta.strip():
        return "Digite uma pergunta."
    prompt = (
        f"Com base no relatório jurídico abaixo, responda de forma objetiva e precisa:\n\n"
        f"PERGUNTA: {pergunta.strip()}\n\n"
        f"RELATÓRIO:\n{relatorio[:80_000]}"
    )
    def _fn():
        return client.models.generate_content(model=model, contents=[prompt]).text
    return _retry(_fn)



def _comprimir_pdf_limite(path: str, max_mb: float = 40.0) -> tuple:
    """
    Garante que o PDF resultante fique abaixo de max_mb.
    Tenta compressão normal primeiro; se ainda acima do limite,
    re-renderiza TODAS as páginas (inclusive texto) como JPEG,
    começando em 110 DPI (suficiente para leitura por IA).
    """
    orig_mb = Path(path).stat().st_size / 1_048_576

    # Fast path — já está dentro do limite, não faz nada
    if orig_mb <= max_mb:
        return path, orig_mb, orig_mb

    comp, _, comp_mb = _comprimir_pdf(path)
    if comp_mb <= max_mb:
        return comp, orig_mb, comp_mb

    if comp != path:
        try: os.remove(comp)
        except: pass

    # 110 DPI é o ponto de partida — evita passes desnecessários em 150/130 DPI
    # que para PDFs de texto raramente cabem no limite
    for dpi, quality in [(110, 42), (96, 40), (80, 38)]:
        tmp_path = None
        try:
            doc = fitz.open(path)
            new_doc = fitz.open()
            for page in doc:
                mat = fitz.Matrix(dpi / 72, dpi / 72)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("jpeg", jpg_quality=quality)
                rect = page.rect
                new_page = new_doc.new_page(width=rect.width, height=rect.height)
                new_page.insert_image(rect, stream=img_bytes, keep_proportion=False)
            tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            tmp_path = tmp.name
            tmp.close()
            new_doc.save(tmp_path, garbage=4, deflate=True)
            new_doc.close()
            doc.close()
            result_mb = Path(tmp_path).stat().st_size / 1_048_576
            if result_mb <= max_mb:
                return tmp_path, orig_mb, result_mb
            os.remove(tmp_path)
        except Exception:
            if tmp_path and os.path.exists(tmp_path):
                try: os.remove(tmp_path)
                except: pass

    return path, orig_mb, orig_mb


def _comprimir_pdf(path: str) -> tuple:
    """
    Re-renderiza páginas escaneadas a 110 DPI / JPEG quality 48 (moderado-alto).
    Páginas de texto são copiadas sem perda. Retorna (path_comprimido, mb_original, mb_comprimido).
    Se a compressão não trouxer ganho ou falhar, devolve o path original.
    """
    original_mb = Path(path).stat().st_size / 1_048_576
    tmp_path = None
    try:
        doc = fitz.open(path)
        new_doc = fitz.open()
        for pn, page in enumerate(doc):
            txt = page.get_text().strip()
            alfa = sum(c.isalpha() for c in txt) if txt else 0
            is_text = (
                len(txt) >= 50
                and alfa / max(len(txt), 1) >= 0.4
                and len(txt.split()) >= 30
            )
            if is_text:
                new_doc.insert_pdf(doc, from_page=pn, to_page=pn)
            else:
                mat = fitz.Matrix(110 / 72, 110 / 72)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("jpeg", jpg_quality=48)
                rect = page.rect
                new_page = new_doc.new_page(width=rect.width, height=rect.height)
                new_page.insert_image(rect, stream=img_bytes, keep_proportion=False)
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp_path = tmp.name
        tmp.close()
        new_doc.save(tmp_path, garbage=4, deflate=True, clean=True)
        new_doc.close()
        doc.close()
        compressed_mb = Path(tmp_path).stat().st_size / 1_048_576
        if compressed_mb >= original_mb * 0.95:
            os.remove(tmp_path)
            return path, original_mb, original_mb
        return tmp_path, original_mb, compressed_mb
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
        return path, original_mb, original_mb


def _barra_progresso(atual: int, total: int) -> str:
    pct = atual / total if total else 0
    blocos = int(pct * 20)
    barra = "█" * blocos + "░" * (20 - blocos)
    return f"[{barra}] {int(pct*100)}% ({atual}/{total} chunks)"



